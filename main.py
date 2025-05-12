import os
import json
import wx
import wx.lib.scrolledpanel as scrolled
import wx.lib.newevent
import threading
import shutil
import requests
import sys
import traceback
import wx
import time
import subprocess
import platform
import signal
import zipfile
import tarfile
import atexit
import re
import select
import importlib
import uuid
import tempfile

# Define a function to check if a package is installed
def check_package_installed(package_name):
    try:
        importlib.import_module(package_name)
        return True
    except ImportError:
        return False

# Special function to ensure Google packages are installed correctly
def ensure_google_packages():
    """Ensure all required Google packages are installed with the correct versions"""
    try:
        # Check if google packages are already installed
        if check_package_installed("google.generativeai"):
            print("Google Generative AI package is already installed")
            return True
            
        print("Installing Google packages...")
        packages = [
            "protobuf>=4.23.0",
            "google-api-python-client",
            "google-api-core",
            "google-cloud-core",
            "google-generativeai>=0.3.0"
        ]
        
        # Install packages one by one to better handle errors
        for package in packages:
            try:
                print(f"Installing {package}...")
                subprocess.check_call([sys.executable, "-m", "pip", "install", "--upgrade", package])
                print(f"{package} installed successfully")
            except Exception as e:
                print(f"Error installing {package}: {e}")
                
        # Install LangChain integration packages
        langchain_packages = [
            "langchain-google-genai",
            "langchain_google_genai"
        ]
        
        for package in langchain_packages:
            try:
                print(f"Installing {package}...")
                subprocess.check_call([sys.executable, "-m", "pip", "install", "--upgrade", package])
                print(f"{package} installed successfully")
            except Exception as e:
                print(f"Error installing {package}: {e}")
        
        # Verify installation by importing
        try:
            import google.generativeai
            print("Google packages installed and verified successfully")
            return True
        except ImportError as e:
            print(f"Failed to import Google packages after installation: {e}")
            return False
            
    except Exception as e:
        print(f"Error ensuring Google packages: {e}")
        return False

# Try to install Google packages at startup
ensure_google_packages()

# Conditionally import fcntl as it's only available on Unix-like systems
try:
    import fcntl
    FCNTL_AVAILABLE = True
except ImportError:
    FCNTL_AVAILABLE = False
from pathlib import Path

# Add detailed startup logging
def log_message(message, is_error=False):
    """
    Logs a message to the console with optional error formatting.
    """
    if is_error:
        print(f"\033[91m[ERROR] {message}\033[0m")  # Red text for errors
    else:
        print(f"\033[92m[INFO] {message}\033[0m")   # Green text for info

# Determine application path
def get_app_path():
    """Get the application path"""
    if getattr(sys, 'frozen', False):
        # We're running in a bundle
        base_path = os.path.dirname(sys.executable)
        # For macOS bundle
        if platform.system() == "Darwin" and ".app/Contents/MacOS" in base_path:
            # If we're in a .app bundle on macOS, we may need to adjust path
            potential_paths = [
                base_path,  # /path/to/app.app/Contents/MacOS
                os.path.dirname(os.path.dirname(base_path)),  # /path/to/app.app
                os.environ.get("RA_APP_PATH", ""),  # From hook-app.py
                os.environ.get("RA_RESOURCES_PATH", ""),  # From hook-macos-paths.py
                os.getcwd()  # Last resort - current directory
            ]
            # Use the first valid path that has a Documents directory or creates it
            for path in potential_paths:
                if path and os.path.isdir(path):
                    documents_dir = os.path.join(path, "Documents")
                    os.makedirs(documents_dir, exist_ok=True)
                    return path
    else:
        # We're running in a development environment
        base_path = os.path.dirname(os.path.abspath(__file__))
    
    return base_path

# Set application path
APP_PATH = get_app_path()

# Function to load environment variables from .env file
def load_env_variables():
    """Load environment variables from .env file if it exists"""
    env_path = os.path.join(APP_PATH, ".env")
    try:
        if os.path.exists(env_path):
            log_message("Loading environment variables from .env file")
            with open(env_path, 'r') as f:
                for line in f:
                    line = line.strip()
                    if line and not line.startswith('#'):
                        key, value = line.split('=', 1)
                        os.environ[key.strip()] = value.strip()
                        log_message(f"Loaded environment variable: {key.strip()}")
    except Exception as e:
        log_message(f"Error loading environment variables: {str(e)}", True)

# Load environment variables at startup
load_env_variables()

# Add new imports for RAG/GraphRAG
import time
import hashlib
from typing import List, Dict, Optional, Tuple, Union, Any

# Neo4j imports - will be conditionally imported
NEO4J_AVAILABLE = False
try:
    from neo4j import GraphDatabase
    import neo4j
    NEO4J_AVAILABLE = True
    log_message("Neo4j package successfully imported")
except ImportError:
    log_message("Warning: neo4j package not installed. Neo4j functionality will be disabled.", True)

# LangChain imports - will be conditionally imported
LANGCHAIN_AVAILABLE = False
try:
    # Import basic components from langchain
    from langchain.prompts import ChatPromptTemplate
    from langchain.schema import Document
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    
    # Import directly from neo4j_vector module for compatibility with older versions
    from langchain_community.vectorstores.neo4j_vector import Neo4jVector
        
    LANGCHAIN_AVAILABLE = True
    log_message("LangChain packages successfully imported")
except ImportError:
    log_message("Warning: langchain packages not installed. RAG functionality will be disabled.", True)

# Embeddings imports - will be conditionally imported
EMBEDDINGS_AVAILABLE = False
try:
    from langchain.embeddings import OpenAIEmbeddings
    EMBEDDINGS_AVAILABLE = True
    log_message("Embedding packages successfully imported")
except ImportError:
    log_message(
        "Warning: embedding packages not installed. Vector functionality will be limited.", True)

# Neo4j Database Manager class - moved here to be available globally
class Neo4jDatabaseManager:
    def __init__(self, uri="bolt://localhost:7687", username="neo4j", password="neo4j", database="neo4j", show_ui=True):
        self.uri = uri
        self.username = username
        self.password = password
        self.database = database
        self.driver = None
        self.vector_store = None
        self.connected = False
        self.show_ui = show_ui  # Control UI visibility
        self.connect()

    def connect(self):
        """Connect to Neo4j database"""
        if not NEO4J_AVAILABLE:
            log_message("Neo4j functionality is not available", True)
            return False
            
        # Use a retry mechanism for connection
        max_retries = 5
        retry_delay = 2
        
        for attempt in range(max_retries):
            try:
                log_message(f"Connecting to Neo4j (attempt {attempt+1}/{max_retries})...")
                
                # Make sure we properly import Neo4j trust settings
                from neo4j import GraphDatabase, TrustAll, TRUST_ALL_CERTIFICATES
                
                # Create connection with more reliable settings for macOS
                self.driver = GraphDatabase.driver(
                    self.uri, 
                    auth=(self.username, self.password),
                    trusted_certificates=TrustAll(),  # Use proper TrustAll object
                    connection_timeout=15,           # Longer connection timeout
                    max_connection_lifetime=3600,    # Keep connections alive for 1 hour
                    keep_alive=True                  # Use TCP keepalive
                )
                
                # Test connection with retry
                for i in range(3):
                    try:
                        self.driver.verify_connectivity()
                        self.connected = True
                        log_message(f"Connected to Neo4j database at {self.uri}")
                        return True
                    except Exception as e:
                        if i < 2:  # Only log intermediate errors if we have more retries
                            log_message(f"Connectivity test attempt {i+1} failed: {str(e)}")
                            time.sleep(1)
                        else:
                            raise  # Re-raise the exception on the last attempt
                
            except Exception as e:
                log_message(f"Connection attempt {attempt+1} failed: {str(e)}", True)
                if self.driver:
                    try:
                        self.driver.close()
                    except:
                        pass
                    self.driver = None
                
                # If we still have retries left, wait and try again
                if attempt < max_retries - 1:
                    time.sleep(retry_delay)
                else:
                    log_message("All connection attempts to Neo4j failed", True)
                    self.connected = False
                    return False
    
    def close(self):
        """Close the Neo4j connection"""
        if self.driver:
            try:
                self.driver.close()
                log_message("Neo4j connection closed")
            except Exception as e:
                log_message(f"Error closing Neo4j connection: {str(e)}", True)
        self.connected = False

    def initialize_vector_store(self, embeddings):
        """Initialize the vector store with Neo4j integration"""
        try:
            if not LANGCHAIN_AVAILABLE or not self.connected:
                log_message("LangChain or Neo4j connection not available", True)
                return False
                
            # For Neo4j 4.x compatibility, we'll use a more basic approach
            log_message("Initializing Neo4j vector store (Neo4j 4.x compatible mode)")
            
            # Create the necessary schema using the Neo4j driver directly
            with self.driver.session() as session:
                # Create constraint if it doesn't exist (Neo4j 4.x syntax)
                try:
                    session.run("""
                    CREATE CONSTRAINT ON (d:Document) 
                    ASSERT d.document_id IS UNIQUE
                    """)
                    log_message("Created document ID constraint")
                except Exception as e:
                    # If the constraint already exists, just log and continue
                    if "EquivalentSchemaRuleAlreadyExists" in str(e):
                        log_message("Document ID constraint already exists", True)
                    else:
                        # If it's some other error, log it but continue
                        log_message(f"Warning creating constraint: {str(e)}", True)
                
                # Create a simple node to ensure we have at least one node in the graph
                # This helps when initializing from an existing graph
                try:
                    session.run("""
                    MERGE (d:Document {document_id: 'placeholder', content: 'Initial placeholder document'})
                    """)
                    log_message("Created placeholder document node")
                except Exception as e:
                    log_message(f"Warning creating placeholder: {str(e)}", True)
                
            # Create our custom Neo4j vector store implementation compatible with Neo4j 4.x
            from langchain_core.documents import Document
            from langchain_core.vectorstores import VectorStore
            
            class Neo4jVectorStore4x(VectorStore):
                """Custom Neo4j Vector Store for Neo4j 4.x compatibility"""
                
                def __init__(self, driver, embedding_function, database="neo4j"):
                    self.driver = driver
                    self.embedding_function = embedding_function
                    self.database = database
                
                def add_documents(self, documents):
                    """Add documents to the vector store with batch processing for better performance"""
                    try:
                        # Process documents in batches to improve performance
                        batch_size = 20  # Process more documents per batch for better throughput
                        
                        # Initialize a dictionary to track nodes by the last character of their ID
                        # This helps distribute documents to minimize lock contention
                        partition_buckets = {}
                        
                        # Assign documents to partitions based on document_id hash
                        for doc in documents:
                            doc_id = doc.metadata.get("document_id", "unknown")
                            # Simple partitioning by last character of the ID hash
                            partition_key = hash(doc_id) % 10
                            if partition_key not in partition_buckets:
                                partition_buckets[partition_key] = []
                            partition_buckets[partition_key].append(doc)
                        
                        # Process each partition bucket sequentially to avoid lock conflicts
                        for partition_key, partition_docs in partition_buckets.items():
                            log_message(f"Processing partition {partition_key} with {len(partition_docs)} documents")
                            
                            # Process in batches within each partition
                            for i in range(0, len(partition_docs), batch_size):
                                batch = partition_docs[i:i+batch_size]
                                
                                # Create all embeddings for the batch
                                embeddings_batch = []
                                metadata_batch = []
                                contents_batch = []
                                
                                # First collect all content to be embedded
                                for doc in batch:
                                    content = doc.page_content
                                    contents_batch.append(content)
                                    
                                    # Extract metadata
                                    metadata = {
                                        "document_id": doc.metadata.get("document_id", "unknown"),
                                        "title": doc.metadata.get("title", "Untitled"),
                                        "content": content,
                                        "raw_metadata": {k: v for k, v in doc.metadata.items() 
                                                       if isinstance(v, (str, int, float, bool))}
                                    }
                                    metadata_batch.append(metadata)
                                
                                # Now create embeddings more efficiently as a batch if supported
                                try:
                                    # Try to generate embeddings in one batch if the embedding function supports it
                                    if hasattr(self.embedding_function, "embed_documents"):
                                        embeddings_batch = self.embedding_function.embed_documents(contents_batch)
                                    else:
                                        # Fall back to individual embedding generation
                                        for content in contents_batch:
                                            try:
                                                embedding = self.embedding_function.embed_query(content)
                                                embeddings_batch.append(embedding)
                                            except Exception as e:
                                                log_message(f"Error generating embedding: {str(e)}", True)
                                                # Add None to keep indexes aligned
                                                embeddings_batch.append(None)
                                except Exception as e:
                                    log_message(f"Error in batch embedding: {str(e)}", True)
                                    # Fall back to individual embedding generation
                                    for content in contents_batch:
                                        try:
                                            embedding = self.embedding_function.embed_query(content)
                                            embeddings_batch.append(embedding)
                                        except Exception as e:
                                            log_message(f"Error generating embedding: {str(e)}", True)
                                            # Add None to keep indexes aligned
                                            embeddings_batch.append(None)
                                
                                # Now process the batch in one session
                                if embeddings_batch:
                                    with self.driver.session() as session:
                                        # Create a transaction function to process the entire batch
                                        def create_documents_tx(tx):
                                            successful = 0
                                            for idx, metadata in enumerate(metadata_batch):
                                                # Skip entries with failed embeddings
                                                if idx >= len(embeddings_batch) or embeddings_batch[idx] is None:
                                                    continue
                                                
                                                try:
                                                    # Create document node with embedding
                                                    tx.run("""
                                                    MERGE (d:Document {document_id: $doc_id})
                                                    SET d.content = $content,
                                                        d.title = $title,
                                                        d.embedding = $embedding
                                                    """, 
                                                    doc_id=metadata["document_id"], 
                                                    content=metadata["content"], 
                                                    title=metadata["title"], 
                                                    embedding=embeddings_batch[idx])
                                                    
                                                    # Add metadata as properties in a single query
                                                    if metadata["raw_metadata"]:
                                                        # Build SET clauses dynamically
                                                        set_clauses = []
                                                        params = {"doc_id": metadata["document_id"]}
                                                        
                                                        for meta_key, meta_value in metadata["raw_metadata"].items():
                                                            param_name = f"meta_{meta_key}_{idx}"
                                                            set_clauses.append(f"d.meta_{meta_key} = ${param_name}")
                                                            params[param_name] = meta_value
                                                        
                                                        if set_clauses:
                                                            meta_query = f"""
                                                            MATCH (d:Document {{document_id: $doc_id}})
                                                            SET {', '.join(set_clauses)}
                                                            """
                                                            tx.run(meta_query, **params)
                                                    successful += 1
                                                except Exception as e:
                                                    log_message(f"Error processing document {metadata['document_id']}: {str(e)}", True)
                                            return successful
                                        
                                        # Execute the batch transaction
                                        try:
                                            successful = session.execute_write(create_documents_tx)
                                            log_message(f"Added batch of {successful} documents to vector store")
                                        except Exception as e:
                                            # Use write_transaction for older Neo4j versions
                                            log_message(f"Falling back to write_transaction: {str(e)}", True)
                                            successful = session.write_transaction(create_documents_tx)
                                            log_message(f"Added batch of {successful} documents to vector store")
                    
                    except Exception as e:
                        log_message(f"Error in batch document processing: {str(e)}", True)
                        return False
                
                def similarity_search(self, query, k=5):
                    """Search for similar documents"""
                    # Get embedding for the query
                    try:
                        query_embedding = self.embedding_function.embed_query(query)
                    except Exception as e:
                        log_message(f"Error generating query embedding: {str(e)}", True)
                        return []
                    
                    # In Neo4j 4.x without GDS, we'll implement a basic vector similarity using Cypher
                    with self.driver.session() as session:
                        result = session.run("""
                        MATCH (d:Document) 
                        WHERE d.embedding IS NOT NULL
                        WITH d, 
                             REDUCE(dot = 0.0, i IN RANGE(0, SIZE(d.embedding) - 1) | 
                                dot + d.embedding[i] * $query_embedding[i]) /
                             (SQRT(REDUCE(norm = 0.0, i IN RANGE(0, SIZE(d.embedding) - 1) | 
                                norm + d.embedding[i] * d.embedding[i])) *
                              SQRT(REDUCE(norm = 0.0, i IN RANGE(0, SIZE($query_embedding) - 1) | 
                                norm + $query_embedding[i] * $query_embedding[i]))) AS score
                        ORDER BY score DESC 
                        LIMIT $k
                        RETURN d.content AS content, d.title AS title, d.document_id AS document_id, d.priority AS priority, score
                        """, query_embedding=query_embedding, k=k)
                        
                        # Convert results to Document objects
                        documents = []
                        for record in result:
                            # Create metadata from fields
                            metadata = {
                                "document_id": record["document_id"],
                                "title": record["title"],
                                "score": record["score"],
                                "priority": record["priority"] if record["priority"] is not None else "Medium"
                            }
                            
                            # Create Document object
                            doc = Document(
                                page_content=record["content"],
                                metadata=metadata
                            )
                            documents.append(doc)
                        
                        return documents
                
                def add_texts(self, texts, metadatas=None, **kwargs):
                    """Add texts to the vector store"""
                    if metadatas is None:
                        metadatas = [{} for _ in texts]
                    
                    # Convert to Document objects
                    documents = []
                    for i, text in enumerate(texts):
                        metadata = metadatas[i] if i < len(metadatas) else {}
                        doc = Document(page_content=text, metadata=metadata)
                        documents.append(doc)
                    
                    # Add documents
                    self.add_documents(documents)
                    
                    # Return empty list of IDs (not used in this implementation)
                    return []
                    
                def as_retriever(self, **kwargs):
                    """Return a retriever interface for the vector store"""
                    from langchain.retrievers import VectorStoreRetriever
                    return VectorStoreRetriever(vectorstore=self, **kwargs)
                
                @classmethod
                def from_texts(cls, texts, embedding, metadatas=None, **kwargs):
                    """Create a Neo4jVectorStore4x from a list of texts."""
                    # Process driver from kwargs
                    driver = kwargs.get("driver", None)
                    if driver is None:
                        url = kwargs.get("url")
                        username = kwargs.get("username", "")
                        password = kwargs.get("password", "")
                        database = kwargs.get("database", "neo4j")
                        
                        # Import here to avoid circular imports
                        from neo4j import GraphDatabase, TrustAll
                        
                        driver = GraphDatabase.driver(
                            url, 
                            auth=(username, password),
                            trusted_certificates=TrustAll()
                        )
                    
                    # Create Neo4jVectorStore instance
                    vs = cls(
                        driver=driver,
                        embedding_function=embedding,
                        database=kwargs.get("database", "neo4j")
                    )
                    
                    # Add texts if provided
                    if texts:
                        vs.add_texts(texts=texts, metadatas=metadatas)
                    
                    return vs
            
            # Initialize the vector store with our embedding function
            self.vector_store = Neo4jVectorStore4x(
                driver=self.driver,
                embedding_function=embeddings
            )
            
            log_message("Vector store initialized successfully (Neo4j 4.x compatibility mode)")
            return True
        except Exception as e:
            log_message(f"Error initializing vector store: {str(e)}", True)
            log_message(traceback.format_exc(), True)
            return False

    def add_document(self, document_id, title, content, metadata=None):
        """Add a document to both Neo4j graph database and vector store with optimized batch processing"""
        try:
            if not self.connected:
                log_message("Not connected to Neo4j", True)
                return False
            
            # Create safe metadata - ensure it contains only primitive types
            safe_metadata = {}
            if metadata:
                for k, v in metadata.items():
                    # Skip nested objects and complex types
                    if isinstance(v, (str, int, float, bool)) or (isinstance(v, list) and all(isinstance(i, (str, int, float, bool)) for i in v)):
                        safe_metadata[k] = v
                    else:
                        # Convert complex objects to string representation
                        safe_metadata[k] = str(v)
            
            # STEP 1: Add to Neo4j graph database with a single transaction
            with self.driver.session() as session:
                def create_document_tx(tx):
                    # Create the document node with all metadata in one operation
                    metadata_props = {f"meta_{k}": v for k, v in safe_metadata.items()}
                    all_props = {
                        "document_id": document_id,
                        "title": title,
                        "content": content,
                        "priority": safe_metadata.get("priority", "Medium"),
                        "updated_at": "datetime()",  # Will be evaluated by Neo4j
                        **metadata_props
                    }
                    
                    # Build a dynamic query for creating the document with all properties
                    properties_str = ", ".join([f"{k}: ${k}" for k in all_props.keys()])
                    query = f"""
                    MERGE (d:Document {{document_id: $document_id}})
                    SET d = {{ {properties_str} }}
                    """
                    
                    # Execute the query with all properties
                    tx.run(query, **all_props)
                    
                    log_message(f"Created document node with {len(metadata_props)} metadata properties")
                    return True
                
                # Execute the transaction
                session.write_transaction(create_document_tx)
            
            # STEP 2: Add to vector store if available - this uses the optimized batch processing already
            vector_store_success = False
            if LANGCHAIN_AVAILABLE and self.vector_store:
                try:
                    from langchain_core.documents import Document
                    from langchain.text_splitter import RecursiveCharacterTextSplitter
                    
                    # Create metadata for langchain document
                    lc_metadata = safe_metadata.copy()
                    lc_metadata["document_id"] = document_id
                    lc_metadata["title"] = title
                    
                    # Use more aggressive chunking for better vector embedding performance
                    text_splitter = RecursiveCharacterTextSplitter(
                        chunk_size=800,  # Larger chunks for better Neo4j performance
                        chunk_overlap=50  # Minimal overlap to reduce redundancy
                    )
                    
                    # Split document into chunks
                    docs = text_splitter.create_documents(
                        texts=[content],
                        metadatas=[lc_metadata]
                    )
                    
                    # Add to vector store - this now uses batching internally with our optimized method
                    start_time = time.time()
                    self.vector_store.add_documents(docs)
                    duration = time.time() - start_time
                    log_message(f"Added document {document_id} to vector store with {len(docs)} chunks in {duration:.2f} seconds")
                    vector_store_success = True
                except Exception as e:
                    log_message(f"Error adding document to vector store: {str(e)}", True)
                    # Continue even if vector store addition fails - we still have the graph data
            
            # STEP 3: Add to knowledge graph (optimized version)
            start_time = time.time()
            try:
                graph_success = self.add_document_to_knowledge_graph(document_id, title, content, safe_metadata)
                duration = time.time() - start_time
                log_message(f"Knowledge graph processing completed in {duration:.2f} seconds")
            except Exception as e:
                log_message(f"Error in knowledge graph processing: {str(e)}", True)
                graph_success = False
            
            # Return true if either component succeeded
            if vector_store_success or graph_success:
                log_message(f"Added document {document_id} to database")
                return True
            else:
                log_message(f"Failed to add document {document_id} to either vector store or knowledge graph", True)
                return False
        except Exception as e:
            log_message(f"Error adding document to database: {str(e)}", True)
            return False
    
    def remove_document(self, document_id):
        """Remove a document from both the Neo4j database and vector store"""
        try:
            if not self.connected:
                log_message("Neo4j connection not available", True)
                return False
                
            # Check if document exists
            with self.driver.session() as session:
                result = session.run(
                    "MATCH (d:Document {document_id: $id}) RETURN count(d) AS count",
                    id=document_id
                )
                doc_exists = result.single()["count"] > 0
                
            if not doc_exists:
                log_message(f"Document {document_id} not found in database")
                return False
            
            # STEP 1: Remove from Neo4j database
            with self.driver.session() as session:
                # Delete document node and relationships
                session.run(
                    """
                    MATCH (d:Document {document_id: $id})
                    DETACH DELETE d
                    """,
                    id=document_id
                )
                
            # STEP 2: If using vector store, remove from there too
            # Unfortunately, most vector stores don't support direct deletion by ID
            # So we'll need a workaround - we'll note that it's been removed
            # from the graph database which is the source of truth
            
            log_message(f"Document {document_id} removed from database")
            return True
        except Exception as e:
            log_message(f"Error removing document from database: {str(e)}", True)
            return False
    
    def delete_all_documents(self):
        """Delete all documents from both Neo4j and vector databases while preserving the databases themselves"""
        try:
            if not self.connected:
                log_message("Neo4j connection not available", True)
                return False

            # STEP 1: Delete all documents from Neo4j database
            with self.driver.session() as session:
                # Delete all document nodes and their relationships
                # Exclude the placeholder document
                session.run("""
                    MATCH (d:Document)
                    WHERE d.document_id <> 'placeholder'
                    DETACH DELETE d
                """)
                log_message("Deleted all documents from Neo4j database")

            # STEP 2: Reinitialize vector store to clear it
            if LANGCHAIN_AVAILABLE and self.vector_store:
                try:
                    # Reinitialize the vector store with empty state
                    self.vector_store = Neo4jVectorStore4x(
                        driver=self.driver,
                        embedding_function=self.vector_store.embedding_function
                    )
                    log_message("Vector store reinitialized and cleared")
                except Exception as e:
                    log_message(f"Error reinitializing vector store: {str(e)}", True)
                    # Continue even if vector store reinitialization fails

            log_message("All documents deleted successfully")
            return True
        except Exception as e:
            log_message(f"Error deleting all documents: {str(e)}", True)
            return False
    
    def get_document_list(self):
        """Get list of documents in the database"""
        try:
            with self.driver.session(database=self.database) as session:
                result = session.run("""
                    MATCH (d:Document)
                    WHERE d.title IS NOT NULL AND d.title <> 'None' AND d.document_id <> 'placeholder'
                    RETURN d.document_id as id, d.title as title, COUNT((d)-[:CONTAINS]->()) as chunks
                    ORDER BY d.title
                """)
                return [(record["id"], record["title"], record["chunks"]) for record in result]
        except Exception as e:
            log_message(f"Error getting document list: {str(e)}", True)
            return []
    
    def query_similar_text(self, query, limit=5, use_graph=False):
        """Query for similar text chunks using vector store or graph relationships
        
        Args:
            query: The text to find similar documents to
            limit: Maximum number of results to return
            use_graph: If True, use graph relationships for search; otherwise use vector embedding
        """
        try:
            if not self.connected:
                log_message("Neo4j connection not available", True)
                return []
            
            # METHOD 1: Using Vector Store for semantic search (standard RAG)
            if not use_graph and self.vector_store:
                try:
                    # Query for similar text using embeddings
                    results = self.vector_store.similarity_search(
                        query=query,
                        k=limit
                    )
                    
                    log_message(f"Found {len(results)} similar documents using vector search")
                    
                    # Apply priority sorting to the retrieved documents if running in an app context
                    if hasattr(wx.GetApp(), 'document_priorities'):
                        app = wx.GetApp()
                        priority_values = {"High": 3, "Medium": 2, "Low": 1}
                        
                        # Add priority score to each document's metadata
                        for doc in results:
                            # First check if priority is already in the metadata (from database)
                            if "priority" in doc.metadata:
                                priority = doc.metadata["priority"]
                            else:
                                # Get the title from metadata which should be the filename
                                title = doc.metadata.get("title", "")
                                
                                # Get priority from app's priorities or default to Medium
                                priority = app.document_priorities.get(title, "Medium")
                            
                            priority_value = priority_values.get(priority, 2)
                            doc.metadata["priority_value"] = priority_value
                            doc.metadata["priority"] = priority
                        
                        # Sort by priority (higher first), then by similarity score
                        results = sorted(results, 
                                         key=lambda doc: (doc.metadata.get("priority_value", 2), 
                                                         doc.metadata.get("score", 0.0)), 
                                         reverse=True)
                    
                    return results
                except Exception as e:
                    log_message(f"Error querying vector store: {str(e)}", True)
                    # Fall back to graph search if vector search fails
            
            # METHOD 2: Using Graph Relationships (GraphRAG approach)
            with self.driver.session() as session:
                if use_graph:
                    # For graph search, retrieve documents with relevant entities or matching content
                    # First retrieve all documents as a simpler approach
                    try:
                        # Simple query to get all documents except the placeholder
                        result = session.run("""
                        MATCH (d:Document)
                        WHERE d.document_id <> 'placeholder'
                        RETURN d.content AS content, d.title AS title, d.document_id AS document_id, d.priority AS priority, 1 AS score
                        LIMIT $limit
                        """, limit=limit)
                        
                        # Process results
                        try:
                            from langchain_core.documents import Document
                        except ImportError:
                            log_message("langchain_core.documents not available, using custom Document class", True)
                            # Define a simple Document class to mimic LangChain's Document
                            class Document:
                                def __init__(self, page_content, metadata=None):
                                    self.page_content = page_content
                                    self.metadata = metadata or {}
                        
                        documents = []
                        
                        for record in result:
                            # Create metadata
                            metadata = {
                                "document_id": record["document_id"],
                                "title": record["title"],
                                "score": record["score"],
                                "priority": record["priority"] if record["priority"] is not None else "Medium",
                                "source": "graph_search"
                            }
                            
                            # Create Document
                            doc = Document(
                                page_content=record["content"],
                                metadata=metadata
                            )
                            documents.append(doc)
                        
                        # Apply priority sorting to the retrieved documents if running in an app context
                        if hasattr(wx.GetApp(), 'document_priorities'):
                            app = wx.GetApp()
                            priority_values = {"High": 3, "Medium": 2, "Low": 1}
                            
                            # Add priority score to each document's metadata
                            for doc in documents:
                                # First check if priority is already in the metadata (from database)
                                if "priority" in doc.metadata:
                                    priority = doc.metadata["priority"]
                                else:
                                    # Get the title from metadata which should be the filename
                                    title = doc.metadata.get("title", "")
                                    
                                    # Get priority from app's priorities or default to Medium
                                    priority = app.document_priorities.get(title, "Medium")
                                
                                priority_value = priority_values.get(priority, 2)
                                doc.metadata["priority_value"] = priority_value
                                doc.metadata["priority"] = priority
                            
                            # Sort by priority (higher first), then by similarity score
                            documents = sorted(documents, 
                                             key=lambda doc: (doc.metadata.get("priority_value", 2), 
                                                             doc.metadata.get("score", 0.0)), 
                                             reverse=True)
                        
                        log_message(f"Found {len(documents)} similar documents using graph search")
                        return documents
                    except Exception as e:
                        log_message(f"Error in graph document retrieval: {str(e)}", True)
                        return []
                else:
                    # If vector store is not available but use_graph is False,
                    # run a basic full-text search in Neo4j as fallback
                    result = session.run("""
                    MATCH (d:Document)
                    WHERE d.document_id <> 'placeholder' AND toLower(d.content) CONTAINS toLower($query)
                    RETURN d.content AS content, d.title AS title, d.document_id AS document_id, d.priority AS priority, 1 AS score
                    LIMIT $limit
                    """, query=query, limit=limit)
                    
                    # Process results
                    try:
                        from langchain_core.documents import Document
                    except ImportError:
                        log_message("langchain_core.documents not available, using custom Document class", True)
                        # Define a simple Document class to mimic LangChain's Document
                        class Document:
                            def __init__(self, page_content, metadata=None):
                                self.page_content = page_content
                                self.metadata = metadata or {}
                    
                    documents = []
                    
                    for record in result:
                        # Create metadata
                        metadata = {
                            "document_id": record["document_id"],
                            "title": record["title"],
                            "score": record["score"],
                            "priority": record["priority"] if record["priority"] is not None else "Medium",
                            "source": "text_search"
                        }
                        
                        # Create Document
                        doc = Document(
                            page_content=record["content"],
                            metadata=metadata
                        )
                        documents.append(doc)
                    
                    # Apply priority sorting to the retrieved documents if running in an app context
                    if hasattr(wx.GetApp(), 'document_priorities'):
                        app = wx.GetApp()
                        priority_values = {"High": 3, "Medium": 2, "Low": 1}
                        
                        # Add priority score to each document's metadata
                        for doc in documents:
                            # First check if priority is already in the metadata (from database)
                            if "priority" in doc.metadata:
                                priority = doc.metadata["priority"]
                            else:
                                # Get the title from metadata which should be the filename
                                title = doc.metadata.get("title", "")
                                
                                # Get priority from app's priorities or default to Medium
                                priority = app.document_priorities.get(title, "Medium")
                            
                            priority_value = priority_values.get(priority, 2)
                            doc.metadata["priority_value"] = priority_value
                            doc.metadata["priority"] = priority
                        
                        # Sort by priority (higher first), then by similarity score
                        documents = sorted(documents, 
                                         key=lambda doc: (doc.metadata.get("priority_value", 2), 
                                                         doc.metadata.get("score", 0.0)), 
                                         reverse=True)
                    
                    log_message(f"Found {len(documents)} documents using text search (fallback)")
                    return documents
        except Exception as e:
            log_message(f"Error querying similar text: {str(e)}", True)
            return []
    
    def create_document_relationship(self, source_id, target_id, relationship_type, properties=None):
        """Create a relationship between two documents"""
        try:
            if not self.connected:
                log_message("Neo4j connection not available", True)
                return False
                
            if properties is None:
                properties = {}
                
            # Create relationship
            with self.driver.session() as session:
                result = session.run(
                    """
                    MATCH (source:Document {document_id: $source_id})
                    MATCH (target:Document {document_id: $target_id})
                    WITH source, target
                    MERGE (source)-[r:RELATED_TO]->(target)
                    SET r.type = $rel_type,
                        r += $properties
                    RETURN count(r) AS count
                    """,
                    source_id=source_id,
                    target_id=target_id,
                    rel_type=relationship_type,
                    properties=properties
                )
                
                created = result.single()["count"] > 0
                
            if created:
                log_message(f"Created {relationship_type} relationship between {source_id} and {target_id}")
                return True
            else:
                log_message(f"Failed to create relationship between {source_id} and {target_id}", True)
                return False
        except Exception as e:
            log_message(f"Error creating document relationship: {str(e)}", True)
            return False

    def add_document_to_knowledge_graph(self, document_id, title, content, metadata=None):
        """Add a document to the knowledge graph by extracting entities and relationships"""
        try:
            if not self.connected:
                log_message("Not connected to Neo4j", True)
                return False
                
            # Import required libraries
            required_libs = [
                "langchain_core.documents",
                "langchain.text_splitter",
                "langchain_experimental.graph_transformers",
                "langchain_openai"
            ]
            
            missing_libs = []
            imported_modules = {}
            
            # Check each required library
            for lib in required_libs:
                try:
                    module_name = lib.split(".")[-1]
                    imported_modules[module_name] = importlib.import_module(lib)
                except ImportError as e:
                    missing_libs.append(f"{lib}: {str(e)}")
            
            if missing_libs:
                log_message(f"Some required libraries for knowledge graph are missing: {', '.join(missing_libs)}", True)
                log_message("Knowledge graph features will be limited. Document will still be available for vector search.", True)
                # We'll still create the base document node so it can be found by vector search
                with self.driver.session() as session:
                    session.run("""
                    MERGE (d:Document {document_id: $doc_id})
                    SET d.title = $title,
                        d.content = $content,
                        d.updated_at = datetime()
                    """, doc_id=document_id, title=title, content=content)
                    log_message("Created placeholder document node")
                return False
                
            # We have all required libs - extract the imported modules
            Document = imported_modules.get("documents").Document
            RecursiveCharacterTextSplitter = imported_modules.get("text_splitter").RecursiveCharacterTextSplitter
            LLMGraphTransformer = imported_modules.get("graph_transformers").LLMGraphTransformer
            ChatOpenAI = imported_modules.get("langchain_openai").ChatOpenAI
            
            # Use LLM to extract entities and relationships
            # First, split document into manageable chunks
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=100
            )
            
            # Create LangChain document with metadata
            lc_metadata = {}
            if metadata:
                lc_metadata = {k: v for k, v in metadata.items() if isinstance(v, (str, int, float, bool))}
            lc_metadata["document_id"] = document_id
            lc_metadata["title"] = title
            
            # Create document chunks
            doc = Document(page_content=content, metadata=lc_metadata)
            chunks = text_splitter.split_documents([doc])
            
            # Initialize LLM and graph transformer
            llm = None
            try:
                api_key = os.environ.get("OPENAI_API_KEY")
                if api_key:
                    llm = ChatOpenAI(
                        temperature=0.2,  # Slightly higher temperature for more creative extraction
                        model="gpt-4o-mini",
                        max_tokens=4000   # Ensure enough output tokens
                    )
                    log_message("Using OpenAI GPT-4o-mini for knowledge graph extraction")
                else:
                    log_message("OpenAI API key not found", True)
                    raise ValueError("OpenAI API key not available")
            except Exception as e:
                log_message(f"Error initializing OpenAI LLM: {str(e)}", True)
                try:
                    from langchain_community.chat_models import ChatOllama
                    llm = ChatOllama(model="llama3", temperature=0)
                    log_message("Using Ollama for knowledge graph extraction")
                except Exception as ollama_err:
                    log_message(f"Error initializing Ollama LLM: {str(ollama_err)}", True)
                    log_message("No suitable LLM available for graph transformation", True)
                    # Still create the document node at minimum
                    with self.driver.session() as session:
                        session.run("""
                        MERGE (d:Document {document_id: $doc_id})
                        SET d.title = $title,
                            d.content = $content,
                            d.updated_at = datetime()
                        """, doc_id=document_id, title=title, content=content)
                        log_message("Created placeholder document node")
                    return False
            
            # Create graph transformer and convert documents
            if llm:
                try:
                    # Define graph schema with allowed nodes and relationships
                    allowed_nodes = ["Person", "Organization", "Location", "Concept", "Technology", "Event", 
                                   "Process", "Product", "Publication", "ResearchField", "Award", "Date", "Skill"]
                    
                    # Define relationships as source-type-target tuples for more precise structure
                    allowed_relationships = [
                        ("Person", "WORKS_FOR", "Organization"),
                        ("Person", "KNOWS", "Person"),
                        ("Person", "DEVELOPED", "Technology"),
                        ("Person", "AUTHOR_OF", "Publication"),
                        ("Person", "RESEARCHES", "ResearchField"),
                        ("Person", "BORN_IN", "Date"),
                        ("Person", "DIED_IN", "Date"),
                        ("Person", "WON", "Award"),
                        ("Person", "HAS_AWARD", "Award"),
                        ("Person", "PRODUCED", "Skill"),
                        ("Organization", "LOCATED_IN", "Location"),
                        ("Organization", "DEVELOPS", "Technology"),
                        ("Organization", "COLLABORATES_WITH", "Organization"),
                        ("Publication", "CITES", "Publication"),
                        ("Publication", "RELATES_TO", "Concept"),
                        ("Technology", "USED_IN", "Process"),
                        ("Product", "USES", "Technology"),
                        ("Concept", "RELATED_TO", "Concept"),
                        ("ResearchField", "SUBFIELD_OF", "ResearchField")
                    ]
                    
                    # Define node properties to capture
                    node_properties = ["name", "description", "date", "url", "role", "field", "type", 
                                     "category", "birth_date", "death_date", "nationality"]
                    
                    # Define relationship properties to capture
                    relationship_properties = ["start_date", "end_date", "strength", "source", "context"]
                    
                    # Create transformer with defined schema
                    llm_transformer = LLMGraphTransformer(
                        llm=llm,
                        allowed_nodes=allowed_nodes,
                        allowed_relationships=allowed_relationships,
                        node_properties=node_properties,
                        relationship_properties=relationship_properties,
                        strict_mode=False  # Allow more flexible extraction
                    )
                    
                    # Convert documents to graph format - optimize by limiting chunks
                    # Process only a subset of chunks to improve speed
                    max_chunks = min(len(chunks), 5)  # Process at most 5 chunks to improve speed while maintaining quality
                    log_message(f"Processing {max_chunks} chunks out of {len(chunks)} for knowledge graph extraction")
                    
                    # Process parallel in smaller batches for better performance
                    graph_documents = llm_transformer.convert_to_graph_documents(chunks[:max_chunks])
                    
                    # Debug: Log the structure of graph documents
                    for i, graph_doc in enumerate(graph_documents):
                        has_nodes = hasattr(graph_doc, 'nodes') and bool(graph_doc.nodes)
                        has_relationships = hasattr(graph_doc, 'relationships') and bool(graph_doc.relationships)
                        log_message(f"Graph document {i}: Has nodes: {has_nodes}, Has relationships: {has_relationships}")
                        if has_nodes:
                            log_message(f"  - Node count: {len(graph_doc.nodes)}")
                            for j, node in enumerate(graph_doc.nodes[:3]):  # Log first 3 nodes
                                log_message(f"  - Node {j}: Type: {node.type if hasattr(node, 'type') else 'Unknown'}, ID: {node.id if hasattr(node, 'id') else 'Unknown'}")
                        if has_relationships:
                            log_message(f"  - Relationship count: {len(graph_doc.relationships)}")
                            for j, rel in enumerate(graph_doc.relationships[:3]):  # Log first 3 relationships
                                log_message(f"  - Relationship {j}: Type: {rel.type}, Source: {rel.source.id if hasattr(rel.source, 'id') else 'Unknown'}, Target: {rel.target.id if hasattr(rel.target, 'id') else 'Unknown'}")
                    
                    # Add to graph database with batch processing
                    with self.driver.session() as session:
                        # First, create the base document node to connect everything
                        session.run("""
                        MERGE (d:Document {document_id: $doc_id})
                        SET d.title = $title,
                            d.content = $content,
                            d.updated_at = datetime()
                        """, doc_id=document_id, title=title, content=content)
                        
                        # Create relationships to other documents (simple ALL_DOCUMENTS relationship)
                        # This ensures that all documents are connected in the graph
                        session.run("""
                        MATCH (d:Document {document_id: $doc_id})
                        MATCH (other:Document)
                        WHERE other.document_id <> $doc_id AND other.document_id <> 'placeholder'
                        WITH d, other
                        LIMIT 100
                        MERGE (d)-[:ALL_DOCUMENTS]->(other)
                        """, doc_id=document_id)
                        
                        # Process graph documents - use transactions for batching
                        def create_kg_tx(tx, graph_doc, doc_id):
                            # Collection to store created node IDs
                            nodes_created = {}
                            created_kg_nodes = False
                            
                            # Check if the graph document has nodes
                            if hasattr(graph_doc, 'nodes') and graph_doc.nodes:
                                # Prepare all nodes in one batch
                                for node in graph_doc.nodes:
                                    # Create node with properties
                                    node_properties = {}
                                    if hasattr(node, 'properties'):
                                        for key, value in node.properties.items():
                                            if isinstance(value, (str, int, float, bool)):
                                                node_properties[key] = value
                                    
                                    # Get node type & ID
                                    node_type = node.type if hasattr(node, 'type') else "Entity"
                                    node_id = node.id if hasattr(node, 'id') else f"unknown_{uuid.uuid4()}"
                                    
                                    # Create entity node with simpler labels to avoid Neo4j 4.x limitations
                                    query = f"""
                                    MERGE (e:__Entity__ {{id: $node_id}})
                                    SET e.node_type = $node_type,
                                        e += $props
                                    WITH e
                                    MATCH (d:Document {{document_id: $doc_id}})
                                    MERGE (d)-[:CONTAINS]->(e)
                                    """
                                    
                                    tx.run(query, node_id=node_id, node_type=node_type,
                                         props=node_properties, doc_id=doc_id)
                                    
                                    # Track that we created this node
                                    nodes_created[node_id] = node_type
                                    created_kg_nodes = True
                                
                            # Create relationships in batch if nodes were created
                            if created_kg_nodes and hasattr(graph_doc, 'relationships') and graph_doc.relationships:
                                for rel in graph_doc.relationships:
                                    # Check if we have both source and target with valid IDs
                                    if (hasattr(rel, 'source') and hasattr(rel.source, 'id') and 
                                        hasattr(rel, 'target') and hasattr(rel.target, 'id')):
                                        
                                        source_id = rel.source.id
                                        target_id = rel.target.id
                                        rel_type = rel.type if hasattr(rel, 'type') else "RELATED_TO"
                                        
                                        # Only create relationships between nodes we've processed
                                        if source_id in nodes_created and target_id in nodes_created:
                                            # Get relationship properties
                                            rel_properties = {}
                                            if hasattr(rel, 'properties'):
                                                for key, value in rel.properties.items():
                                                    if isinstance(value, (str, int, float, bool)):
                                                        rel_properties[key] = value
                                            
                                            # Create relationship with a simpler syntax to avoid Neo4j 4.x limitations
                                            query = f"""
                                            MATCH (source:`__Entity__` {{id: $source_id}})
                                            MATCH (target:`__Entity__` {{id: $target_id}})
                                            WHERE source <> target
                                            MERGE (source)-[r:RELATED_TO]->(target)
                                            SET r.type = $rel_type,
                                                r += $props
                                            """
                                            
                                            tx.run(query, source_id=source_id, target_id=target_id, 
                                                 rel_type=rel_type, props=rel_properties)
                            
                            return created_kg_nodes
                        
                        # Process each graph document with its own transaction
                        for graph_doc in graph_documents:
                            # Skip empty graph documents
                            if not (hasattr(graph_doc, 'nodes') and graph_doc.nodes):
                                continue
                                
                            # Process with transaction
                            created = session.write_transaction(create_kg_tx, graph_doc, document_id)
                            if created:
                                log_message(f"Added knowledge graph elements for document {document_id}")
                
                    return True
                except Exception as e:
                    log_message(f"Error in knowledge graph processing: {str(e)}", True)
                    log_message(traceback.format_exc(), True)
                    # Create basic document node anyway
                    with self.driver.session() as session:
                        session.run("""
                        MERGE (d:Document {document_id: $doc_id})
                        SET d.title = $title,
                            d.content = $content,
                            d.updated_at = datetime()
                        """, doc_id=document_id, title=title, content=content)
                    return False
            
            return False
        except Exception as e:
            log_message(f"Error in knowledge graph processing: {str(e)}", True)
            log_message(traceback.format_exc(), True)
            return False

# Override IsDisplayAvailable to always return True
wx.PyApp.IsDisplayAvailable = lambda _: True

# Custom event for updating UI from threads
ResponseEvent, EVT_RESPONSE = wx.lib.newevent.NewEvent()
StreamEvent, EVT_STREAM = wx.lib.newevent.NewEvent()  # New event for streaming updates
DbInitEvent, EVT_DB_INIT = wx.lib.newevent.NewEvent()  # New event for DB initialization

# Class for managing the embedded Neo4j server
class EmbeddedNeo4jServer:
    NEO4J_VERSION = "4.4.30"  # Using a stable LTS version
    NEO4J_PORT = 7687  # Default Bolt port
    
    def __init__(self, base_path=None):
        self.base_path = base_path or APP_PATH
        self.server_dir = os.path.join(self.base_path, "Neo4jDB", "neo4j-server")
        self.data_dir = os.path.join(self.base_path, "Neo4jDB", "data")
        self.logs_dir = os.path.join(self.base_path, "Neo4jDB", "logs")
        self.process = None
        self.running = False
        
        # Create directories
        os.makedirs(self.server_dir, exist_ok=True)
        os.makedirs(self.data_dir, exist_ok=True)
        os.makedirs(self.logs_dir, exist_ok=True)
        
        # Register cleanup function
        atexit.register(self.stop)
        
    def check_java_version(self):
        """Check if compatible Java version is installed"""
        try:
            # First check for bundled JRE
            if self.check_bundled_jre():
                log_message("Found bundled JRE for Neo4j")
                return True
                
            # Check system Java
            try:
                java_cmd = "java"
                
                # Run java -version and capture output
                proc = subprocess.run(
                    [java_cmd, "-version"],
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    text=True
                )
                
                # Combine stdout and stderr since java -version outputs to stderr
                output = proc.stdout + proc.stderr
                
                if proc.returncode != 0:
                    log_message(f"Java check failed with return code {proc.returncode}", True)
                    log_message(f"Java error output: {output}", True)
                    return False
                
                # Extract version information
                log_message(f"Java version output: {output}")
                
                # Check if this is a compatible version (Java 11+)
                if "version" in output:
                    version_pattern = r'version "([^"]+)"'
                    matches = re.search(version_pattern, output)
                    
                    if matches:
                        version = matches.group(1)
                        log_message(f"Found Java version: {version}")
                        
                        # Check major version
                        # Look for patterns like "11.x.x" or "1.8.x"
                        major_version_pattern = r'(\d+)[\._]'
                        major_matches = re.search(major_version_pattern, version)
                        
                        if major_matches:
                            major_version = int(major_matches.group(1))
                            
                            # Neo4j 4.4 requires Java 11 or higher
                            if major_version >= 11:
                                log_message(f"Java version {major_version} is compatible with Neo4j")
                                return True
                            else:
                                log_message(f"Java version {major_version} is not compatible with Neo4j (needs 11+)", True)
                                return False
                
                log_message("Could not determine Java version, will attempt to use anyway")
                return True
                
            except Exception as e:
                log_message(f"Error checking Java version: {str(e)}", True)
                return False
                
        except Exception as e:
            log_message(f"Error in Java version check: {str(e)}", True)
            log_message(traceback.format_exc(), True)
            return False
    
    def install_java(self):
        """Download and install Java 11 for Neo4j"""
        try:
            log_message("Starting Java 11 installation for Neo4j...")
            
            # Create jre directory if it doesn't exist
            jre_dir = os.path.join(self.base_path, "jre")
            os.makedirs(jre_dir, exist_ok=True)
            
            # Determine download URL based on system architecture
            system = platform.system().lower()
            machine = platform.machine().lower()
            
            # Create a temporary directory for download
            with tempfile.TemporaryDirectory() as tmp_dir:
                if system == "darwin":  # macOS
                    if "arm" in machine or "aarch64" in machine:  # Apple Silicon
                        url = "https://github.com/adoptium/temurin11-binaries/releases/download/jdk-11.0.21%2B9/OpenJDK11U-jre_aarch64_mac_hotspot_11.0.21_9.tar.gz"
                    else:  # Intel Mac
                        url = "https://github.com/adoptium/temurin11-binaries/releases/download/jdk-11.0.21%2B9/OpenJDK11U-jre_x64_mac_hotspot_11.0.21_9.tar.gz"
                    
                    archive_path = os.path.join(tmp_dir, "openjdk-11-jre.tar.gz")
                elif system == "windows":  # Windows
                    url = "https://github.com/adoptium/temurin11-binaries/releases/download/jdk-11.0.21%2B9/OpenJDK11U-jre_x64_windows_hotspot_11.0.21_9.zip"
                    archive_path = os.path.join(tmp_dir, "openjdk-11-jre.zip")
                elif system == "linux":  # Linux
                    if "arm" in machine or "aarch64" in machine:  # ARM
                        url = "https://github.com/adoptium/temurin11-binaries/releases/download/jdk-11.0.21%2B9/OpenJDK11U-jre_aarch64_linux_hotspot_11.0.21_9.tar.gz"
                    else:  # x86_64
                        url = "https://github.com/adoptium/temurin11-binaries/releases/download/jdk-11.0.21%2B9/OpenJDK11U-jre_x64_linux_hotspot_11.0.21_9.tar.gz"
                    archive_path = os.path.join(tmp_dir, "openjdk-11-jre.tar.gz")
                else:
                    log_message(f"Unsupported platform: {system} {machine}", True)
                    return False
                
                # 1. Download the JRE
                log_message(f"Downloading Java 11 JRE from {url}")
                try:
                    response = requests.get(url, stream=True)
                    if response.status_code != 200:
                        log_message(f"Failed to download Java JRE: HTTP {response.status_code}", True)
                        return False
                    
                    with open(archive_path, 'wb') as f:
                        for chunk in response.iter_content(chunk_size=8192):
                            if chunk:
                                f.write(chunk)
                    
                    log_message("Java 11 JRE download complete")
                except Exception as e:
                    log_message(f"Error downloading Java JRE: {str(e)}", True)
                    return False
                
                # 2. Extract the archive
                log_message(f"Extracting Java 11 JRE to {jre_dir}")
                try:
                    # If jre directory already exists, remove it first
                    if os.path.exists(jre_dir):
                        shutil.rmtree(jre_dir)
                    os.makedirs(jre_dir, exist_ok=True)
                    
                    if archive_path.endswith('.tar.gz'):
                        with tarfile.open(archive_path, 'r:gz') as tar:
                            # Create a temporary extraction directory
                            extract_dir = os.path.join(tmp_dir, "extracted")
                            os.makedirs(extract_dir, exist_ok=True)
                            tar.extractall(extract_dir)
                            
                            # Find the JRE directory (usually there's a single top-level directory)
                            extracted_dirs = [d for d in os.listdir(extract_dir) if os.path.isdir(os.path.join(extract_dir, d))]
                            if not extracted_dirs:
                                log_message("No directories found in extracted archive", True)
                                return False
                            
                            # Move the contents to the Java directory
                            src_dir = os.path.join(extract_dir, extracted_dirs[0])
                            
                            # On macOS, preserve the special directory structure
                            if system == "darwin" and os.path.exists(os.path.join(src_dir, "Contents")):
                                # This is a macOS .jdk package structure
                                for item in os.listdir(src_dir):
                                    shutil.move(os.path.join(src_dir, item), os.path.join(jre_dir, item))
                            else:
                                # Standard structure
                                shutil.copytree(src_dir, jre_dir, dirs_exist_ok=True)
                    
                    elif archive_path.endswith('.zip'):
                        with zipfile.ZipFile(archive_path, 'r') as zip_ref:
                            # Create a temporary extraction directory
                            extract_dir = os.path.join(tmp_dir, "extracted")
                            os.makedirs(extract_dir, exist_ok=True)
                            zip_ref.extractall(extract_dir)
                            
                            # Find the JRE directory
                            extracted_dirs = [d for d in os.listdir(extract_dir) if os.path.isdir(os.path.join(extract_dir, d))]
                            if not extracted_dirs:
                                log_message("No directories found in extracted archive", True)
                                return False
                            
                            # Move the contents to the Java directory
                            src_dir = os.path.join(extract_dir, extracted_dirs[0])
                            shutil.copytree(src_dir, jre_dir, dirs_exist_ok=True)
                    
                    log_message("Java 11 JRE extracted successfully")
                except Exception as e:
                    log_message(f"Error extracting Java JRE: {str(e)}", True)
                    log_message(traceback.format_exc(), True)
                    return False
                
                # 3. Make executables executable on Unix systems
                if system != "windows":
                    try:
                        # Determine the bin directory
                        if system == "darwin" and os.path.exists(os.path.join(jre_dir, "Contents", "Home", "bin")):
                            bin_dir = os.path.join(jre_dir, "Contents", "Home", "bin")
                        else:
                            bin_dir = os.path.join(jre_dir, "bin")
                        
                        if os.path.exists(bin_dir):
                            for file in os.listdir(bin_dir):
                                file_path = os.path.join(bin_dir, file)
                                if os.path.isfile(file_path):
                                    # Make executable
                                    os.chmod(file_path, 0o755)
                            log_message(f"Made Java executables executable in {bin_dir}")
                        else:
                            log_message(f"Bin directory not found at {bin_dir}", True)
                    except Exception as e:
                        log_message(f"Error setting executable permissions: {str(e)}", True)
                
                # 4. Verify installation
                java_bin = None
                if system == "darwin" and os.path.exists(os.path.join(jre_dir, "Contents", "Home", "bin", "java")):
                    java_bin = os.path.join(jre_dir, "Contents", "Home", "bin", "java")
                elif os.path.exists(os.path.join(jre_dir, "bin", "java")):
                    java_bin = os.path.join(jre_dir, "bin", "java")
                elif system == "windows" and os.path.exists(os.path.join(jre_dir, "bin", "java.exe")):
                    java_bin = os.path.join(jre_dir, "bin", "java.exe")
                
                if not java_bin or not os.path.exists(java_bin):
                    log_message("Java executable not found in installed JRE", True)
                    return False
                
                # Test the Java installation
                try:
                    result = subprocess.run(
                        [java_bin, "-version"],
                        stdout=subprocess.PIPE,
                        stderr=subprocess.PIPE,
                        text=True
                    )
                    version_output = result.stdout + result.stderr
                    log_message(f"Installed Java version: {version_output.strip()}")
                    
                    if "11." not in version_output and "11]" not in version_output:
                        log_message("Warning: Installed Java does not appear to be version 11", True)
                    
                except Exception as e:
                    log_message(f"Error testing Java installation: {str(e)}", True)
                    return False
                
                # 5. Configure Neo4j to use the new Java
                try:
                    # Determine Java home directory
                    if system == "darwin" and os.path.exists(os.path.join(jre_dir, "Contents", "Home")):
                        java_home = os.path.join(jre_dir, "Contents", "Home")
                    else:
                        java_home = jre_dir
                    
                    # Set environment variables
                    os.environ["JAVA_HOME"] = java_home
                    path_separator = ";" if system == "windows" else ":"
                    os.environ["PATH"] = os.path.dirname(java_bin) + path_separator + os.environ.get("PATH", "")
                    
                    # Set up wrapper configuration
                    neo4j_dir = os.path.join(self.base_path, "Neo4jDB", "neo4j-server")
                    if os.path.exists(neo4j_dir):
                        # Create conf directory if it doesn't exist
                        conf_dir = os.path.join(neo4j_dir, "conf")
                        os.makedirs(conf_dir, exist_ok=True)
                        
                        # Create wrapper configuration file
                        wrapper_conf = os.path.join(conf_dir, "neo4j-wrapper.conf")
                        with open(wrapper_conf, 'w') as f:
                            f.write("# Neo4j wrapper configuration\n")
                            f.write("# Automatically configured by Neo4j Research Assistant\n\n")
                            # Escape backslashes on Windows
                            if system == "windows":
                                java_bin_esc = java_bin.replace("\\", "\\\\")
                            else:
                                java_bin_esc = java_bin
                            f.write(f"wrapper.java.command={java_bin_esc}\n")
                        
                        # Update startup scripts
                        if system != "windows":
                            # Unix-like system
                            neo4j_script = os.path.join(neo4j_dir, "bin", "neo4j")
                            if os.path.exists(neo4j_script):
                                # Read the script
                                with open(neo4j_script, 'r') as f:
                                    script_lines = f.readlines()
                                
                                # Look for JAVA_HOME setting
                                java_home_line = -1
                                for i, line in enumerate(script_lines):
                                    if "JAVA_HOME" in line and not line.strip().startswith('#'):
                                        java_home_line = i
                                        break
                                
                                if java_home_line >= 0:
                                    # Replace existing JAVA_HOME
                                    script_lines[java_home_line] = f'JAVA_HOME="{java_home}"\n'
                                else:
                                    # Add JAVA_HOME after shebang
                                    script_lines.insert(1, f'JAVA_HOME="{java_home}"\n')
                                    script_lines.insert(2, f'export JAVA_HOME\n')
                                    script_lines.insert(3, f'PATH="{os.path.dirname(java_bin)}:$PATH"\n')
                                    script_lines.insert(4, f'export PATH\n')
                                
                                # Write the updated script
                                with open(neo4j_script, 'w') as f:
                                    f.writelines(script_lines)
                                
                                # Make executable
                                os.chmod(neo4j_script, 0o755)
                        else:
                            # Windows
                            neo4j_bat = os.path.join(neo4j_dir, "bin", "neo4j.bat")
                            if os.path.exists(neo4j_bat):
                                # Read the script
                                with open(neo4j_bat, 'r') as f:
                                    script_lines = f.readlines()
                                
                                # Look for JAVA_HOME setting
                                java_home_line = -1
                                for i, line in enumerate(script_lines):
                                    if "JAVA_HOME" in line and not line.strip().lower().startswith('rem'):
                                        java_home_line = i
                                        break
                                
                                if java_home_line >= 0:
                                    # Replace existing JAVA_HOME
                                    script_lines[java_home_line] = f'set "JAVA_HOME={java_home}"\r\n'
                                else:
                                    # Add JAVA_HOME after initial comments
                                    insert_pos = 0
                                    for i, line in enumerate(script_lines):
                                        if not line.strip().lower().startswith('rem') and line.strip():
                                            insert_pos = i
                                            break
                                    
                                    script_lines.insert(insert_pos, f'set "JAVA_HOME={java_home}"\r\n')
                                    script_lines.insert(insert_pos + 1, f'set "PATH={os.path.dirname(java_bin)};%PATH%"\r\n')
                                
                                # Write the updated script
                                with open(neo4j_bat, 'w') as f:
                                    f.writelines(script_lines)
                        
                        log_message("Neo4j configured to use the installed Java")
                    else:
                        log_message("Neo4j server directory not found, skipping configuration", True)
                
                except Exception as e:
                    log_message(f"Error configuring Neo4j to use Java: {str(e)}", True)
                    log_message(traceback.format_exc(), True)
                    return False
            
            log_message("Java 11 installation completed successfully")
            return True
            
        except Exception as e:
            log_message(f"Error installing Java: {str(e)}", True)
            log_message(traceback.format_exc(), True)
            return False
    
    def download_if_needed(self):
        """Download Neo4j server if it doesn't exist"""
        try:
            # Check if Neo4j is already installed
            if self._is_neo4j_installed():
                log_message(f"Neo4j server already installed at {self.server_dir}")
                return True
                
            log_message(f"Downloading Neo4j {self.NEO4J_VERSION}...")
            
            # Determine download URL based on system
            system = platform.system().lower()
            if system == "windows":
                url = f"https://dist.neo4j.org/neo4j-community-{self.NEO4J_VERSION}-windows.zip"
                archive_path = os.path.join(self.base_path, "Neo4jDB", "neo4j.zip")
            else:  # Linux or macOS
                url = f"https://dist.neo4j.org/neo4j-community-{self.NEO4J_VERSION}-unix.tar.gz"
                archive_path = os.path.join(self.base_path, "Neo4jDB", "neo4j.tar.gz")
            
            # Download the archive
            response = requests.get(url, stream=True)
            if response.status_code != 200:
                log_message(f"Failed to download Neo4j: HTTP {response.status_code}", True)
                return False
                
            # Save the downloaded file
            with open(archive_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    if chunk:
                        f.write(chunk)
            
            log_message(f"Downloaded Neo4j to {archive_path}")
            
            # Extract the archive
            if system == "windows":
                with zipfile.ZipFile(archive_path, 'r') as zip_ref:
                    zip_ref.extractall(os.path.join(self.base_path, "Neo4jDB"))
            else:
                with tarfile.open(archive_path, 'r:gz') as tar_ref:
                    tar_ref.extractall(os.path.join(self.base_path, "Neo4jDB"))
            
            # Rename the extracted directory
            extracted_dir = os.path.join(self.base_path, "Neo4jDB", f"neo4j-community-{self.NEO4J_VERSION}")
            if os.path.exists(extracted_dir):
                if os.path.exists(self.server_dir):
                    shutil.rmtree(self.server_dir)
                shutil.move(extracted_dir, self.server_dir)
            
            # Remove the archive
            os.remove(archive_path)
            
            # Configure Neo4j
            self._configure_neo4j()
            
            log_message("Neo4j server installed successfully")
            return True
        except Exception as e:
            log_message(f"Error downloading Neo4j: {str(e)}", True)
            log_message(traceback.format_exc(), True)
            return False
    
    def _configure_neo4j(self):
        """Configure Neo4j settings"""
        try:
            config_path = os.path.join(self.server_dir, "conf", "neo4j.conf")
            if not os.path.exists(config_path):
                log_message(f"Neo4j config file not found at {config_path}", True)
                return False
            
            # Read the configuration file
            with open(config_path, 'r') as f:
                config_lines = f.readlines()
            
            # Modify the configuration
            new_config_lines = []
            for line in config_lines:
                # Set data directory
                if line.strip().startswith('#dbms.directories.data='):
                    line = f"dbms.directories.data={self.data_dir}\n"
                # Set logs directory
                elif line.strip().startswith('#dbms.directories.logs='):
                    line = f"dbms.directories.logs={self.logs_dir}\n"
                # Enable APOC
                elif line.strip().startswith('#dbms.security.procedures.unrestricted='):
                    line = "dbms.security.procedures.unrestricted=apoc.*\n"
                # Disable authentication - important for local use
                elif line.strip().startswith('#dbms.security.auth_enabled='):
                    line = "dbms.security.auth_enabled=false\n"
                # Set default listen address specifically to localhost to avoid network interface issues
                elif line.strip().startswith('#dbms.default_listen_address='):
                    line = "dbms.default_listen_address=localhost\n"
                # Configure the Bolt connector explicitly
                elif line.strip().startswith('dbms.connector.bolt.enabled='):
                    line = "dbms.connector.bolt.enabled=true\n"
                elif line.strip().startswith('#dbms.connector.bolt.tls_level='):
                    line = "dbms.connector.bolt.tls_level=DISABLED\n"
                elif line.strip().startswith('#dbms.connector.bolt.listen_address='):
                    line = f"dbms.connector.bolt.listen_address=localhost:{self.NEO4J_PORT}\n"
                # Disable HTTP connector to avoid port conflicts
                elif line.strip().startswith('#dbms.connector.http.enabled='):
                    line = "dbms.connector.http.enabled=false\n"
                # Configure memory settings for better vector operations
                elif line.strip().startswith('#dbms.memory.heap.initial_size='):
                    line = "dbms.memory.heap.initial_size=1g\n"
                elif line.strip().startswith('#dbms.memory.heap.max_size='):
                    line = "dbms.memory.heap.max_size=2g\n"
                elif line.strip().startswith('#dbms.memory.pagecache.size='):
                    line = "dbms.memory.pagecache.size=1g\n"
                
                new_config_lines.append(line)
            
            # Add additional configuration for APOC procedures and functions
            new_config_lines.append("\n# Custom configuration for Research Assistant\n")
            new_config_lines.append("dbms.security.procedures.allowlist=apoc.*\n")
            new_config_lines.append("dbms.security.procedures.unrestricted=apoc.*\n")
            
            # Add configuration for Cypher functionality required for vector similarity
            new_config_lines.append("\n# Enable REDUCE functionality for vector similarity calculations\n")
            new_config_lines.append("dbms.cypher.forbid_exhaustive_shortestpath=false\n")
            new_config_lines.append("dbms.cypher.min_replan_interval=0\n")
            
            # Add explicit configuration to ensure Bolt connector works reliably
            new_config_lines.append("\n# Ensure reliable connections\n")
            new_config_lines.append("dbms.connector.bolt.address=localhost:7687\n")
            
            # Add performance optimizations for bulk loading
            new_config_lines.append("\n# Performance optimizations for document uploads\n")
            new_config_lines.append("dbms.memory.pagecache.flush.buffer.enabled=true\n")
            new_config_lines.append("dbms.memory.pagecache.flush.buffer.size_in_pages=100\n")
            new_config_lines.append("dbms.tx_state.memory_allocation=ON_HEAP\n")
            new_config_lines.append("dbms.jvm.additional=-XX:+UseG1GC\n")
            
            # Add transaction and connection pool settings for faster batch operations
            new_config_lines.append("\n# Transaction and connection settings for faster batch operations\n")
            new_config_lines.append("dbms.transaction.concurrent.maximum=16\n")
            new_config_lines.append("dbms.memory.transaction.global_max_size=512m\n")
            new_config_lines.append("dbms.transaction.timeout=600s\n")  # 10 minute timeout for long transactions
            new_config_lines.append("dbms.connector.bolt.thread_pool_min_size=10\n")
            new_config_lines.append("dbms.connector.bolt.thread_pool_max_size=40\n")
            
            # Write the modified configuration
            with open(config_path, 'w') as f:
                f.writelines(new_config_lines)
            
            # Now configure wrapper.conf to use any available Java 11 or 17 if installed
            wrapper_path = os.path.join(self.server_dir, "conf", "neo4j-wrapper.conf")
            if os.path.exists(wrapper_path):
                try:
                    # Try to find Java 11 or 17 path using platform-specific commands
                    java_path = None
                    if platform.system().lower() == "darwin":  # macOS
                        try:
                            # Try to find Java 11
                            result = subprocess.run(
                                ["/usr/libexec/java_home", "-v", "11"], 
                                capture_output=True, 
                                text=True,
                                check=False
                            )
                            if result.returncode == 0:
                                java_path = result.stdout.strip()
                            else:
                                # Try to find Java 17
                                result = subprocess.run(
                                    ["/usr/libexec/java_home", "-v", "17"], 
                                    capture_output=True, 
                                    text=True,
                                    check=False
                                )
                                if result.returncode == 0:
                                    java_path = result.stdout.strip()
                        except Exception as e:
                            log_message(f"Error finding Java path: {str(e)}", True)
                    
                    if java_path:
                        log_message(f"Found Java at: {java_path}")
                        with open(wrapper_path, 'a') as f:
                            f.write(f"\n# Custom Java path\n")
                            f.write(f"wrapper.java.command={java_path}/bin/java\n")
                            # Add JVM heap settings for better performance with vectors
                            f.write(f"wrapper.java.additional=-Xms1g\n")
                            f.write(f"wrapper.java.additional=-Xmx2g\n")
                            # Add G1GC options for better GC performance
                            f.write(f"wrapper.java.additional=-XX:+UseG1GC\n")
                            f.write(f"wrapper.java.additional=-XX:G1HeapRegionSize=4m\n")
                            f.write(f"wrapper.java.additional=-XX:InitiatingHeapOccupancyPercent=70\n")
                            f.write(f"wrapper.java.additional=-XX:MaxGCPauseMillis=200\n")
                            f.write(f"wrapper.java.additional=-XX:G1ReservePercent=10\n")
                            f.write(f"wrapper.java.additional=-XX:+ParallelRefProcEnabled\n")
                except Exception as e:
                    log_message(f"Error configuring Java path: {str(e)}", True)
            
            log_message("Neo4j configured successfully with performance optimizations")
            return True
        except Exception as e:
            log_message(f"Error configuring Neo4j: {str(e)}", True)
            log_message(traceback.format_exc(), True)
            return False
    
    def start(self):
        """Start the Neo4j server process"""
        try:
            # If the server is already running, just return
            if self.running and self.process and self.process.poll() is None:
                log_message("Neo4j server is already running")
                return True
            
            # Check if Neo4j server is installed
            if not self._is_neo4j_installed():
                log_message("Neo4j server is not installed, downloading...")
                if not self.download_if_needed():
                    log_message("Failed to download Neo4j server", True)
                    return False
            
            # Check if Java is installed and compatible version
            if not self.check_java_version():
                log_message("Compatible Java version not found, attempting to install Java 11...")
                if not self.install_java():
                    log_message("Failed to install Java 11. Neo4j server cannot start", True)
                    return False
                log_message("Java 11 installed successfully, continuing with Neo4j startup")
            
            # Fix Neo4j startup script on macOS before starting
            if platform.system().lower() == "darwin":
                self._fix_neo4j_startup_script()
                
            # Kill any stale Java/Neo4j processes
            self.kill_stale_processes()
            
            # Check if we should preserve Neo4j data
            preserve_data = os.environ.get('PRESERVE_NEO4J_DATA', 'True').lower() == 'true'
            
            # Check for a preservation marker file
            preserve_marker = os.path.join(self.base_path, "Neo4jDB", ".preserve")
            if os.path.exists(preserve_marker):
                preserve_data = True
            
            # If there was a previous failure, try to clean up
            if hasattr(self, '_connection_failure_count') and self._connection_failure_count >= 3:
                log_message("Previous connection failures detected, attempting to reinstall Neo4j...", True)
                self.stop()  # Make sure any existing process is stopped
                
                # Remove existing Neo4j installation and try fresh (even if preservation is enabled)
                if os.path.exists(self.server_dir):
                    try:
                        shutil.rmtree(self.server_dir)
                        log_message("Removed existing Neo4j installation")
                    except Exception as e:
                        log_message(f"Error removing Neo4j directory: {str(e)}", True)
                
                # Re-download Neo4j
                if not self.download_if_needed():
                    log_message("Failed to re-download Neo4j server", True)
                    return False
                
                self._connection_failure_count = 0
            else:
                # Clean up database files only if not preserving data or only lock files if preserving
                self.cleanup_database_files()
            
            log_message("Starting Neo4j server...")
            
            # Prepare environment with correct Java path
            env = os.environ.copy()
            
            # Check for bundled JRE and set JAVA_HOME if available
            bundled_jre = self.check_bundled_jre()
            if bundled_jre:
                java_home = os.path.join(APP_PATH, "jre", "Contents", "Home")
                log_message(f"Setting JAVA_HOME to bundled JRE: {java_home}")
                env["JAVA_HOME"] = java_home
                env["PATH"] = f"{os.path.join(java_home, 'bin')}:{env.get('PATH', '')}"
            else:
                log_message("No bundled JRE found, using system Java")
            
            # Determine the correct start command based on platform
            if platform.system().lower() == "windows":
                cmd = [os.path.join(self.server_dir, "bin", "neo4j.bat"), "console"]
            else:
                cmd = [os.path.join(self.server_dir, "bin", "neo4j"), "console"]
            
            # Start the Neo4j process with the configured environment
            self.process = subprocess.Popen(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                env=env
            )
            
            # Set initial startup wait time based on platform
            # macOS needs more time to properly initialize Neo4j
            startup_wait = 15 if platform.system().lower() == "darwin" else 5
            log_message(f"Waiting {startup_wait} seconds for Neo4j to start...")
            
            # Capture output during startup
            output_lines = []
            error_lines = []
            start_time = time.time()
            
            # Collect output for diagnostics while waiting for startup
            while time.time() - start_time < startup_wait:
                # Check if the process has exited
                if self.process.poll() is not None:
                    log_message("Neo4j process exited prematurely during startup", True)
                    break
                
                # Read any available output/error with timeout to avoid blocking
                if self.process.stdout:
                    try:
                        # Set stdout to non-blocking mode (platform specific)
                        if FCNTL_AVAILABLE:
                            # Unix-like systems
                            fd = self.process.stdout.fileno()
                            fl = fcntl.fcntl(fd, fcntl.F_GETFL)
                            fcntl.fcntl(fd, fcntl.F_SETFL, fl | os.O_NONBLOCK)
                            line = self.process.stdout.readline()
                        else:
                            # Windows or other platforms - use polling
                            # Check if data is available without blocking
                            rlist, _, _ = select.select([self.process.stdout], [], [], 0.1)
                            if rlist:
                                line = self.process.stdout.readline()
                            else:
                                line = ""
                        
                        if line:
                            output_lines.append(line.strip())
                    except (IOError, OSError, ValueError) as e:
                        # No data available or other issues
                        pass
                
                if self.process.stderr:
                    try:
                        # Set stderr to non-blocking mode (platform specific)
                        if FCNTL_AVAILABLE:
                            # Unix-like systems
                            fd = self.process.stderr.fileno()
                            fl = fcntl.fcntl(fd, fcntl.F_GETFL)
                            fcntl.fcntl(fd, fcntl.F_SETFL, fl | os.O_NONBLOCK)
                            line = self.process.stderr.readline()
                        else:
                            # Windows or other platforms - use polling
                            # Check if data is available without blocking
                            rlist, _, _ = select.select([self.process.stderr], [], [], 0.1)
                            if rlist:
                                line = self.process.stderr.readline()
                            else:
                                line = ""
                        
                        if line:
                            error_lines.append(line.strip())
                    except (IOError, OSError, ValueError) as e:
                        # No data available or other issues
                        pass
                
                # Sleep a bit to avoid CPU spinning
                time.sleep(0.1)
            
            # Check if process is still running
            if self.process.poll() is None:
                self.running = True
                log_message("Neo4j server started successfully")
                
                # Log collected output for diagnostics
                if output_lines:
                    log_message("Neo4j startup output:")
                    for line in output_lines:
                        log_message(f"  {line}")
                        
                if error_lines:
                    log_message("Neo4j startup errors:")
                    for line in error_lines:
                        log_message(f"  {line}", True)
                
                # Check if we can connect to the server with overall timeout
                try:
                    from neo4j import GraphDatabase
                    
                    # Try to establish a connection with retry
                    max_retries = 5
                    retry_delay = 2
                    connection_timeout = 30  # 30 second overall timeout
                    connection_start_time = time.time()
                    
                    for retry in range(max_retries):
                        # Check if we've exceeded the overall timeout
                        if time.time() - connection_start_time > connection_timeout:
                            log_message("Connection timeout exceeded while trying to connect to Neo4j", True)
                            # Kill the process to avoid hanging
                            if self.process and self.process.poll() is None:
                                try:
                                    self.process.kill()
                                    log_message("Killed Neo4j process due to connection timeout", True)
                                except:
                                    pass
                            self.running = False
                            return False
                        
                        try:
                            # Use a short timeout for the driver connection
                            driver = GraphDatabase.driver(
                                f"bolt://localhost:{self.NEO4J_PORT}", 
                                auth=("neo4j", "neo4j"),
                                connection_timeout=5  # 5 second connection timeout
                            )
                            # Test the connection with a quick query
                            with driver.session() as session:
                                session.run("RETURN 1", timeout=5)  # 5 second query timeout
                            driver.close()
                            log_message("Connected to Neo4j server successfully")
                            # Reset connection failure count on success
                            self._connection_failure_count = 0
                            return True
                        except Exception as e:
                            log_message(f"Connection attempt {retry+1}/{max_retries} failed: {str(e)}")
                            
                            # Check if we've exceeded the overall timeout
                            if time.time() - connection_start_time > connection_timeout:
                                log_message("Connection timeout exceeded during retry", True)
                                break
                                
                            time.sleep(retry_delay)
                    
                    # Increment connection failure count
                    if not hasattr(self, '_connection_failure_count'):
                        self._connection_failure_count = 1
                    else:
                        self._connection_failure_count += 1
                    
                    log_message(f"Failed to connect to Neo4j server after multiple attempts (failure count: {self._connection_failure_count})", True)
                    
                    # Kill the process if we couldn't connect to avoid hanging
                    if self.process and self.process.poll() is None:
                        try:
                            self.process.kill()
                            log_message("Killed Neo4j process due to connection failure", True)
                        except:
                            pass
                    self.running = False
                    return False
                except Exception as e:
                    log_message(f"Error testing Neo4j connection: {str(e)}", True)
                    # Continue anyway as the server might still be usable
                    return True
            else:
                # Process exited prematurely - collect output
                if self.process.stdout:
                    for line in self.process.stdout:
                        output_lines.append(line.strip())
                if self.process.stderr:
                    for line in self.process.stderr:
                        error_lines.append(line.strip())
                
                # Log the output for diagnostics
                if output_lines:
                    log_message("Neo4j output before failure:")
                    for line in output_lines:
                        log_message(f"  {line}")
                
                if error_lines:
                    log_message("Neo4j errors:")
                    for line in error_lines:
                        log_message(f"  {line}", True)
                else:
                    log_message("No specific error output from Neo4j", True)
                
                # Log the exit code
                log_message(f"Neo4j process exited with code: {self.process.returncode}", True)
                
                self.process = None
                self.running = False
                return False
        except Exception as e:
            log_message(f"Error starting Neo4j server: {str(e)}", True)
            log_message(traceback.format_exc(), True)
            self.running = False
            return False
    
    def stop(self):
        """Stop the Neo4j server process"""
        try:
            if not self.running or not self.process:
                log_message("Neo4j server is not running")
                return True
            
            log_message("Stopping Neo4j server...")
            
            # First try a graceful shutdown
            if platform.system().lower() == "windows":
                stop_cmd = [os.path.join(self.server_dir, "bin", "neo4j.bat"), "stop"]
            else:
                stop_cmd = [os.path.join(self.server_dir, "bin", "neo4j"), "stop"]
            
            try:
                # Execute stop command with timeout
                subprocess.run(
                    stop_cmd,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    timeout=10
                )
                
                # Wait for process to exit
                for _ in range(10):
                    if self.process.poll() is not None:
                        log_message("Neo4j server stopped successfully")
                        self.process = None
                        self.running = False
                        return True
                    time.sleep(1)
            except Exception as e:
                log_message(f"Graceful shutdown failed: {str(e)}")
            
            # If graceful shutdown failed, terminate the process
            if self.process and self.process.poll() is None:
                log_message("Terminating Neo4j server process...")
                self.process.terminate()
                
                # Wait for process to terminate
                try:
                    self.process.wait(timeout=10)
                    log_message("Neo4j server terminated")
                except subprocess.TimeoutExpired:
                    # Force kill if termination timed out
                    log_message("Killing Neo4j server process...")
                    self.process.kill()
                    self.process.wait()
                    log_message("Neo4j server process killed")
            
            self.process = None
            self.running = False
            log_message("Neo4j server stopped")
            return True
        except Exception as e:
            log_message(f"Error stopping Neo4j server: {str(e)}", True)
            log_message(traceback.format_exc(), True)
            # Set to not running even if there was an error
            self.running = False
            self.process = None
            return False
    
    def check_bundled_jre(self):
        """Check if we have a bundled JRE"""
        try:
            jre_dir = os.path.join(self.base_path, "jre")
            
            # Check for Java binary based on platform
            if platform.system().lower() == "darwin":
                # macOS JRE has a different structure
                macos_java_bin = os.path.join(jre_dir, "Contents", "Home", "bin", "java")
                if os.path.exists(macos_java_bin):
                    return True
                java_bin = os.path.join(jre_dir, "bin", "java")
                return os.path.exists(java_bin)
            elif platform.system().lower() == "windows":
                java_bin = os.path.join(jre_dir, "bin", "java.exe")
                return os.path.exists(java_bin)
            else:
                java_bin = os.path.join(jre_dir, "bin", "java")
                return os.path.exists(java_bin)
        except Exception as e:
            log_message(f"Error checking bundled JRE: {str(e)}", True)
            return False
    
    def _is_neo4j_installed(self):
        """Check if Neo4j is already installed"""
        # Check for bin directory with neo4j executable
        if platform.system().lower() == "windows":
            return os.path.exists(os.path.join(self.server_dir, "bin", "neo4j.bat"))
        else:
            return os.path.exists(os.path.join(self.server_dir, "bin", "neo4j"))
    
    def cleanup_database_files(self):
        """Clean up database files to avoid corruption issues"""
        try:
            # Check if we should preserve Neo4j data
            preserve_data = os.environ.get('PRESERVE_NEO4J_DATA', 'True').lower() == 'true'
            
            # Check for a preservation marker file
            preserve_marker = os.path.join(self.base_path, "Neo4jDB", ".preserve")
            if os.path.exists(preserve_marker):
                preserve_data = True
                
            if preserve_data:
                log_message("Neo4j data preservation is enabled. Only removing lock files...")
                
                # Clean up store_lock files that can cause lock issues
                data_dir = os.path.join(self.base_path, "Neo4jDB", "data")
                if os.path.exists(data_dir):
                    # Remove only lock files that might be causing issues
                    for root, dirs, files in os.walk(data_dir):
                        for file in files:
                            if file == "store_lock" or file.endswith(".lock"):
                                lock_file = os.path.join(root, file)
                                try:
                                    os.remove(lock_file)
                                    log_message(f"Removed lock file: {lock_file}")
                                except Exception as e:
                                    log_message(f"Error removing lock file {lock_file}: {str(e)}", True)
                
                log_message("Lock files cleanup completed (database preserved)")
                return True
            
            # If not preserving data, proceed with full cleanup
            log_message("Cleaning up database files to resolve potential corruption...")
            
            # Clean up store_lock files that can cause lock issues
            data_dir = os.path.join(self.base_path, "Neo4jDB", "data")
            if os.path.exists(data_dir):
                # Remove store_lock files that might be causing issues
                for root, dirs, files in os.walk(data_dir):
                    for file in files:
                        if file == "store_lock" or file.endswith(".lock"):
                            lock_file = os.path.join(root, file)
                            try:
                                os.remove(lock_file)
                                log_message(f"Removed lock file: {lock_file}")
                            except Exception as e:
                                log_message(f"Error removing lock file {lock_file}: {str(e)}", True)
            
            # Remove databases directory
            db_dir = os.path.join(self.base_path, "Neo4jDB", "data", "databases")
            if os.path.exists(db_dir):
                try:
                    shutil.rmtree(db_dir)
                    log_message("Removed databases directory")
                except Exception as e:
                    log_message(f"Error removing databases directory: {str(e)}", True)
            
            # Remove transactions directory
            tx_dir = os.path.join(self.base_path, "Neo4jDB", "data", "transactions")
            if os.path.exists(tx_dir):
                try:
                    shutil.rmtree(tx_dir)
                    log_message("Removed transactions directory")
                except Exception as e:
                    log_message(f"Error removing transactions directory: {str(e)}", True)
                    
            # Wait a moment to ensure file system operations complete
            time.sleep(1)
            
            log_message("Database cleanup completed")
            return True
        except Exception as e:
            log_message(f"Error during database cleanup: {str(e)}", True)
            return False

    def kill_stale_processes(self):
        """Kill any running Neo4j Java processes that might interfere with startup"""
        try:
            log_message("Checking for stale Neo4j processes...")
            
            # Function to find and kill processes based on platform
            if platform.system().lower() == "darwin":  # macOS
                # First, check for the existence of any Neo4j Java processes
                try:
                    # Look for Java processes containing "neo4j"
                    check_cmd = ["pgrep", "-f", "neo4j"]
                    proc = subprocess.run(check_cmd, capture_output=True, text=True)
                    
                    if proc.stdout.strip():
                        log_message(f"Found {len(proc.stdout.strip().split())} potential Neo4j processes")
                        
                        # Get PIDs
                        pids = proc.stdout.strip().split()
                        for pid in pids:
                            try:
                                # Get more info about the process
                                ps_cmd = ["ps", "-p", pid, "-o", "command"]
                                proc_info = subprocess.run(ps_cmd, capture_output=True, text=True)
                                
                                # Only kill if it's definitely a Neo4j process
                                if "neo4j" in proc_info.stdout.lower():
                                    kill_cmd = ["kill", "-9", pid]
                                    subprocess.run(kill_cmd)
                                    log_message(f"Killed Neo4j process with PID {pid}")
                            except Exception as e:
                                log_message(f"Error processing PID {pid}: {str(e)}", True)
                    else:
                        log_message("No stale Neo4j processes found")
                        
                except Exception as e:
                    log_message(f"Error checking for Neo4j processes: {str(e)}", True)
                    
            elif platform.system().lower() == "windows":
                # For Windows, use taskkill to find and terminate Neo4j Java processes
                try:
                    # Find Java processes
                    check_cmd = ["tasklist", "/FI", "IMAGENAME eq java.exe", "/FO", "CSV"]
                    proc = subprocess.run(check_cmd, capture_output=True, text=True)
                    
                    # Check output for Neo4j related processes
                    if "java.exe" in proc.stdout:
                        log_message("Found Java processes, checking for Neo4j")
                        # Kill only the Neo4j related ones
                        kill_cmd = ["taskkill", "/F", "/FI", "IMAGENAME eq java.exe", "/FI", "WINDOWTITLE eq *neo4j*"]
                        subprocess.run(kill_cmd, capture_output=True)
                        log_message("Killed stale Neo4j Java processes")
                except Exception as e:
                    log_message(f"Error killing stale processes on Windows: {str(e)}", True)
            
            log_message("Process cleanup completed")
        except Exception as e:
            log_message(f"Error while killing stale processes: {str(e)}", True)

    def _fix_neo4j_startup_script(self):
        """Fix common issues in Neo4j startup scripts for macOS"""
        try:
            # Check if the startup script exists
            neo4j_script = os.path.join(self.server_dir, "bin", "neo4j")
            if not os.path.exists(neo4j_script):
                log_message("Neo4j startup script not found, skipping fix", True)
                return False
                
            log_message("Checking Neo4j startup script for issues...")
            
            # Read the script content
            with open(neo4j_script, 'r') as f:
                script_lines = f.readlines()
                
            # Find and fix syntax errors in the Darwin section
            fixed = False
            in_darwin_section = False
            darwin_section_start = -1
            darwin_section_end = -1
            
            # First pass: identify Darwin section
            for i, line in enumerate(script_lines):
                line = line.strip()
                if "Darwin*)" in line:
                    in_darwin_section = True
                    darwin_section_start = i
                elif in_darwin_section and line.endswith(";;"):
                    darwin_section_end = i
                    break
                    
            # Fix the Darwin section if found
            if darwin_section_start >= 0 and darwin_section_end >= 0:
                # Replace the Darwin section with a fixed version
                fixed_section = [
                    '  Darwin*) darwin=true\n',
                    '           if [ -z "$JAVA_VERSION" ] ; then\n',
                    '             JAVA_VERSION="CurrentJDK"\n',
                    '           else\n',
                    '             echo "Using Java version: $JAVA_VERSION"\n',
                    '           fi\n',
                    '           # Use the bundled JRE directly\n',
                    '           JAVA_HOME="' + os.path.join(APP_PATH, "jre", "Contents", "Home") + '"\n',
                    '           ;;\n'
                ]
                
                script_lines[darwin_section_start:darwin_section_end+1] = fixed_section
                fixed = True
                
                # Write the fixed script back
                with open(neo4j_script, 'w') as f:
                    f.writelines(script_lines)
                    
                # Make sure the script is executable
                os.chmod(neo4j_script, 0o755)
                log_message("Fixed Neo4j startup script")
                
            return fixed
        except Exception as e:
            log_message(f"Error fixing Neo4j startup script: {str(e)}", True)
            return False

# Dialog for managing document priorities
class DocumentPriorityDialog(wx.Dialog):
    def __init__(self, parent, documents, doc_priorities):
        super(DocumentPriorityDialog, self).__init__(parent, title="Document Priorities", size=(500, 400))
        
        self.documents = documents
        self.doc_priorities = doc_priorities.copy()
        
        # Create a panel and main sizer
        panel = wx.Panel(self)
        main_sizer = wx.BoxSizer(wx.VERTICAL)
        
        # Add instructions
        instructions = wx.StaticText(panel, label="Set priority for each document:")
        main_sizer.Add(instructions, 0, wx.ALL, 10)
        
        # Create a scrolled panel for the documents
        self.scroll_panel = scrolled.ScrolledPanel(panel)
        self.scroll_panel.SetAutoLayout(True)
        self.scroll_panel.SetupScrolling()
        
        scroll_sizer = wx.BoxSizer(wx.VERTICAL)
        
        # Priority choices
        self.priority_levels = ["Low", "Medium", "High"]
        self.priority_controls = {}
        
        # Add controls for each document
        for doc in self.documents:
            doc_sizer = wx.BoxSizer(wx.HORIZONTAL)
            
            # Document name
            doc_label = wx.StaticText(self.scroll_panel, label=doc)
            doc_sizer.Add(doc_label, 1, wx.ALIGN_CENTER_VERTICAL | wx.RIGHT, 10)
            
            # Priority dropdown
            priority_choice = wx.Choice(self.scroll_panel, choices=self.priority_levels)
            current_priority = self.doc_priorities.get(doc, "Medium")
            priority_choice.SetStringSelection(current_priority)
            doc_sizer.Add(priority_choice, 0)
            
            # Store the control reference
            self.priority_controls[doc] = priority_choice
            
            scroll_sizer.Add(doc_sizer, 0, wx.EXPAND | wx.ALL, 5)
        
        self.scroll_panel.SetSizer(scroll_sizer)
        main_sizer.Add(self.scroll_panel, 1, wx.EXPAND | wx.ALL, 10)
        
        # Add OK and Cancel buttons
        button_sizer = wx.StdDialogButtonSizer()
        ok_button = wx.Button(panel, wx.ID_OK)
        ok_button.SetDefault()
        button_sizer.AddButton(ok_button)
        cancel_button = wx.Button(panel, wx.ID_CANCEL)
        button_sizer.AddButton(cancel_button)
        button_sizer.Realize()
        
        main_sizer.Add(button_sizer, 0, wx.ALIGN_CENTER | wx.ALL, 10)
        
        panel.SetSizer(main_sizer)
        self.Centre()
    
    def get_priorities(self):
        priorities = {}
        for doc, control in self.priority_controls.items():
            priorities[doc] = control.GetStringSelection()
        return priorities

# Dialog for saving/loading prompts
class PromptLibraryDialog(wx.Dialog):
    def __init__(self, parent, mode="load", current_prompt=""):
        title = "Load Prompt" if mode == "load" else "Save Prompt"
        super(PromptLibraryDialog, self).__init__(parent, title=title, size=(500, 400))
        
        self.mode = mode
        self.current_prompt = current_prompt
        self.prompts_dir = os.path.join(APP_PATH, "Prompts")
        self.selected_prompt = None
        
        # Create a panel and main sizer
        panel = wx.Panel(self)
        main_sizer = wx.BoxSizer(wx.VERTICAL)
        
        # Add appropriate controls based on mode
        if mode == "load":
            # Prompt list for loading
            list_label = wx.StaticText(panel, label="Select a prompt to load:")
            main_sizer.Add(list_label, 0, wx.ALL, 10)
            
            self.prompt_list = wx.ListBox(panel, style=wx.LB_SINGLE)
            self.load_saved_prompts()
            main_sizer.Add(self.prompt_list, 1, wx.EXPAND | wx.LEFT | wx.RIGHT, 10)
            
            # Delete button for removing prompts
            delete_btn = wx.Button(panel, label="Delete Selected Prompt")
            delete_btn.Bind(wx.EVT_BUTTON, self.on_delete_prompt)
            main_sizer.Add(delete_btn, 0, wx.ALL, 10)
            
        else:  # Save mode
            # Prompt name field
            name_label = wx.StaticText(panel, label="Prompt Name:")
            main_sizer.Add(name_label, 0, wx.ALL, 10)
            
            self.name_field = wx.TextCtrl(panel)
            main_sizer.Add(self.name_field, 0, wx.EXPAND | wx.LEFT | wx.RIGHT | wx.BOTTOM, 10)
            
            # Preview of prompt to save
            preview_label = wx.StaticText(panel, label="Prompt Preview:")
            main_sizer.Add(preview_label, 0, wx.LEFT | wx.RIGHT | wx.TOP, 10)
            
            preview_text = wx.TextCtrl(panel, value=current_prompt[:200] + ("..." if len(current_prompt) > 200 else ""), 
                                     style=wx.TE_MULTILINE | wx.TE_READONLY)
            preview_text.SetMinSize((400, 100))
            main_sizer.Add(preview_text, 1, wx.EXPAND | wx.LEFT | wx.RIGHT | wx.BOTTOM, 10)
        
        # Add OK and Cancel buttons
        button_sizer = wx.StdDialogButtonSizer()
        ok_button = wx.Button(panel, wx.ID_OK)
        ok_button.SetDefault()
        button_sizer.AddButton(ok_button)
        cancel_button = wx.Button(panel, wx.ID_CANCEL)
        button_sizer.AddButton(cancel_button)
        button_sizer.Realize()
        
        main_sizer.Add(button_sizer, 0, wx.ALIGN_CENTER | wx.ALL, 10)
        
        panel.SetSizer(main_sizer)
        self.Centre()
    
    def load_saved_prompts(self):
        # Clear the list
        self.prompt_list.Clear()
        
        # Get all JSON and TXT files from the Prompts directory
        if os.path.exists(self.prompts_dir):
            prompt_files = [f for f in os.listdir(self.prompts_dir) 
                           if os.path.isfile(os.path.join(self.prompts_dir, f)) 
                           and (f.lower().endswith('.json') or f.lower().endswith('.txt'))]
            
            # Add each file to the list (without extension)
            for prompt_file in prompt_files:
                prompt_name = os.path.splitext(prompt_file)[0]
                self.prompt_list.Append(prompt_name)
    
    def on_delete_prompt(self, event):
        selected_idx = self.prompt_list.GetSelection()
        if selected_idx != wx.NOT_FOUND:
            prompt_name = self.prompt_list.GetString(selected_idx)
            
            # Confirm deletion
            dialog = wx.MessageDialog(self, 
                                     f"Are you sure you want to delete the prompt '{prompt_name}'?",
                                     "Confirm Deletion", 
                                     wx.YES_NO | wx.ICON_QUESTION)
            
            if dialog.ShowModal() == wx.ID_YES:
                # Create paths for both possible extensions
                json_path = os.path.join(self.prompts_dir, f"{prompt_name}.json") 
                txt_path = os.path.join(self.prompts_dir, f"{prompt_name}.txt")
                deleted = False
                
                # Try to delete JSON file if it exists
                if os.path.exists(json_path):
                    try:
                        os.remove(json_path)
                        deleted = True
                        log_message(f"Deleted JSON prompt: {prompt_name}")
                    except Exception as e:
                        log_message(f"Error deleting JSON prompt: {str(e)}", True)
                        
                # Try to delete TXT file if it exists
                if os.path.exists(txt_path):
                    try:
                        os.remove(txt_path)
                        deleted = True
                        log_message(f"Deleted TXT prompt: {prompt_name}")
                    except Exception as e:
                        log_message(f"Error deleting TXT prompt: {str(e)}", True)
                
                # Update the list if any file was deleted
                if deleted:
                    self.prompt_list.Delete(selected_idx)
                else:
                    wx.MessageBox(f"Error: Could not find or delete prompt files", "Error", wx.OK | wx.ICON_ERROR)
            
            dialog.Destroy()
    
    def get_prompt_content(self):
        if self.mode == "load":
            selected_idx = self.prompt_list.GetSelection()
            if selected_idx != wx.NOT_FOUND:
                prompt_name = self.prompt_list.GetString(selected_idx)
                self.selected_prompt = prompt_name
                
                # Try JSON format first
                json_path = os.path.join(self.prompts_dir, f"{prompt_name}.json")
                txt_path = os.path.join(self.prompts_dir, f"{prompt_name}.txt")
                
                # Check if JSON file exists
                if os.path.exists(json_path):
                    try:
                        with open(json_path, 'r', encoding='utf-8') as f:
                            prompt_data = json.load(f)
                            return prompt_data.get("content", "")
                    except Exception as e:
                        log_message(f"Error loading JSON prompt: {str(e)}", True)
                        return ""
                
                # If JSON doesn't exist or failed, try TXT format
                elif os.path.exists(txt_path):
                    try:
                        with open(txt_path, 'r', encoding='utf-8') as f:
                            return f.read()
                    except Exception as e:
                        log_message(f"Error loading TXT prompt: {str(e)}", True)
                        return ""
                        
                return ""
            return ""
        else:
            return self.current_prompt
    
    def get_prompt_name(self):
        if self.mode == "load":
            return self.selected_prompt
        else:
            return self.name_field.GetValue()

class MessageEditDialog(wx.Dialog):
    def __init__(self, parent, message, title="Edit Message"):
        super(MessageEditDialog, self).__init__(
            parent, title=title, size=(700, 500)
        )
        
        # Create main sizer
        main_sizer = wx.BoxSizer(wx.VERTICAL)
        
        # Message text control
        self.message_ctrl = wx.TextCtrl(
            self, style=wx.TE_MULTILINE | wx.TE_RICH2
        )
        self.message_ctrl.SetValue(message)
        main_sizer.Add(self.message_ctrl, 1, wx.EXPAND | wx.ALL, 10)
        
        # Buttons
        button_sizer = wx.StdDialogButtonSizer()
        
        self.save_button = wx.Button(self, wx.ID_OK, "Save")
        self.save_button.Bind(wx.EVT_BUTTON, self.on_save)
        button_sizer.AddButton(self.save_button)
        
        self.cancel_button = wx.Button(self, wx.ID_CANCEL, "Cancel")
        button_sizer.AddButton(self.cancel_button)
        
        button_sizer.Realize()
        main_sizer.Add(button_sizer, 0, wx.ALIGN_RIGHT | wx.ALL, 10)
        
        self.SetSizer(main_sizer)
        
    def on_save(self, event):
        self.EndModal(wx.ID_OK)
        
    def GetMessage(self):
        return self.message_ctrl.GetValue()

# Settings Dialog for API Keys
class SettingsDialog(wx.Dialog):
    def __init__(self, parent, config):
        super(SettingsDialog, self).__init__(
            parent, title="Settings", size=(500, 500),
            style=wx.DEFAULT_DIALOG_STYLE | wx.RESIZE_BORDER
        )
        
        self.config = config
        self.api_keys = {}
        
        # Load existing API keys from environment
        for model_key, model_config in self.config.get("models", {}).items():
            api_key_env = model_config.get("api_key_env", "")
            if api_key_env:
                self.api_keys[model_key] = os.environ.get(api_key_env, "")
        
        # Create main panel and sizer
        panel = wx.Panel(self)
        main_sizer = wx.BoxSizer(wx.VERTICAL)
        
        # Title
        title_text = wx.StaticText(panel, label="API Key Configuration")
        font = title_text.GetFont()
        font.SetPointSize(12)
        font.SetWeight(wx.FONTWEIGHT_BOLD)
        title_text.SetFont(font)
        main_sizer.Add(title_text, 0, wx.ALL, 10)
        
        # Description
        description = wx.StaticText(panel, label="Enter your API keys for the following models:")
        main_sizer.Add(description, 0, wx.LEFT | wx.RIGHT | wx.BOTTOM, 10)
        
        # Create a scrolled window for API key inputs
        scroll_window = wx.ScrolledWindow(panel, style=wx.VSCROLL)
        scroll_window.SetScrollRate(0, 10)
        scroll_sizer = wx.BoxSizer(wx.VERTICAL)
        
        # Add API key input fields
        self.api_key_ctrls = {}
        
        for model_key, model_config in sorted(self.config.get("models", {}).items()):
            if "api_key_env" in model_config:
                # Create a panel for this model
                model_panel = wx.Panel(scroll_window)
                model_sizer = wx.BoxSizer(wx.VERTICAL)
                
                # Model name label
                model_name = model_config.get("name", model_key)
                label = wx.StaticText(model_panel, label=f"{model_name} API Key:")
                model_sizer.Add(label, 0, wx.BOTTOM, 5)
                
                # API key input field
                api_key_ctrl = wx.TextCtrl(model_panel, style=wx.TE_PASSWORD)
                if model_key in self.api_keys:
                    api_key_ctrl.SetValue(self.api_keys[model_key])
                model_sizer.Add(api_key_ctrl, 0, wx.EXPAND | wx.BOTTOM, 10)
                
                # Environment variable name
                env_name = model_config.get("api_key_env", "")
                env_label = wx.StaticText(model_panel, label=f"Environment Variable: {env_name}")
                font = env_label.GetFont()
                font.SetPointSize(8)
                env_label.SetFont(font)
                model_sizer.Add(env_label, 0, wx.BOTTOM, 15)
                
                model_panel.SetSizer(model_sizer)
                scroll_sizer.Add(model_panel, 0, wx.EXPAND | wx.LEFT | wx.RIGHT, 10)
                
                # Store reference to the control
                self.api_key_ctrls[model_key] = api_key_ctrl
        
        scroll_window.SetSizer(scroll_sizer)
        main_sizer.Add(scroll_window, 1, wx.EXPAND | wx.ALL, 10)
        
        # Buttons
        button_sizer = wx.StdDialogButtonSizer()
        
        self.save_button = wx.Button(panel, wx.ID_OK, "Save Settings")
        self.save_button.Bind(wx.EVT_BUTTON, self.on_save)
        button_sizer.AddButton(self.save_button)
        
        self.cancel_button = wx.Button(panel, wx.ID_CANCEL, "Cancel")
        button_sizer.AddButton(self.cancel_button)
        
        button_sizer.Realize()
        main_sizer.Add(button_sizer, 0, wx.ALIGN_RIGHT | wx.ALL, 10)
        
        panel.SetSizer(main_sizer)
        
        # Set minimum size
        self.SetMinSize(wx.Size(400, 300))
    
    def on_save(self, event):
        # Save API keys to environment variables
        for model_key, api_key_ctrl in self.api_key_ctrls.items():
            api_key = api_key_ctrl.GetValue().strip()
            if model_key in self.config.get("models", {}):
                api_key_env = self.config["models"][model_key].get("api_key_env", "")
                if api_key_env and api_key:
                    os.environ[api_key_env] = api_key
                    log_message(f"Updated API key for {model_key}")
        
        # Create a .env file to store the API keys persistently
        try:
            env_path = os.path.join(APP_PATH, ".env")
            with open(env_path, 'w') as f:
                for model_key, api_key_ctrl in self.api_key_ctrls.items():
                    api_key = api_key_ctrl.GetValue().strip()
                    if model_key in self.config.get("models", {}):
                        api_key_env = self.config["models"][model_key].get("api_key_env", "")
                        if api_key_env and api_key:
                            f.write(f"{api_key_env}={api_key}\n")
            
            log_message("API keys saved to .env file")
        except Exception as e:
            log_message(f"Error saving API keys to .env file: {str(e)}", True)
        
        self.EndModal(wx.ID_OK)

# Main application class
class ResearchAssistantApp(wx.Frame):
    def __init__(self):
        super(ResearchAssistantApp, self).__init__(
            None, title="Neo4j Research Assistant", 
            size=(1200, 800)
        )
        
        # Initialize application state
        self.conversation_history = []
        self.documents = {}  # Filename -> Content
        self.document_priorities = {}  # Filename -> Priority (High, Medium, Low)
        self.message_positions = []
        self.db_initialized = False
        self.neo4j_manager = None
        self.neo4j_server = None
        self.rag_chain = None
        
        # Initialize streaming state
        self.current_streaming_response = None
        
        # Create and bind custom event types
        global StreamEvent
        StreamEvent, EVT_STREAM = wx.lib.newevent.NewEvent()
        self.Bind(EVT_STREAM, self.on_stream_event)
        
        # Create and bind custom database initialized event
        global DbInitEvent
        DbInitEvent, EVT_DB_INIT = wx.lib.newevent.NewEvent()
        self.Bind(EVT_DB_INIT, self.on_db_init_event)
        
        # Application path
        self.base_path = APP_PATH
        
        # Load configuration
        self.config = load_config() or create_default_config()
        
        # Initialize member variables
        self.conversation_dirty = False
        self.showing_loader = False
        
        # Setup UI
        self.setup_ui()
        
        # Center the window on the screen
        self.Centre()
        
        # Initialize the database in a background thread
        threading.Thread(target=self.initialize_database, daemon=True).start()
        
        # Register event handler for database initialization
        self.Bind(EVT_DB_INIT, self.on_db_init_event)
        self.Bind(EVT_STREAM, self.on_stream_event)
        
        # Register close handler
        self.Bind(wx.EVT_CLOSE, self.on_close)
        
        # Show the window
        self.Show()
    
    def initialize_database(self):
        """Initialize the Neo4j database connection"""
        try:
            # First check if Neo4j is available
            if not NEO4J_AVAILABLE:
                log_message("Neo4j functionality is not available", True)
                # Post event to main thread
                wx.PostEvent(self, DbInitEvent(success=False, error="Neo4j functionality is not available"))
                return
                
            # Initialize the embedded Neo4j server
            self.neo4j_server = EmbeddedNeo4jServer()
            
            # Check JRE if it's bundled
            jre_status = self.neo4j_server.check_bundled_jre()
            if jre_status:
                log_message("Found bundled JRE - will use for Neo4j")
            else:
                log_message("No bundled JRE found - will use system Java if available")
            
            # Start the server in a separate thread
            log_message("Starting Neo4j server...")
            server_thread = threading.Thread(target=self._start_neo4j_server_with_retry)
            server_thread.daemon = True
            server_thread.start()
        except Exception as e:
            log_message(f"Error initializing database: {str(e)}", True)
            log_message(traceback.format_exc(), True)
            # Post event to main thread
            wx.PostEvent(self, DbInitEvent(success=False, error=str(e)))
            
    def _start_neo4j_server_with_retry(self):
        """Start Neo4j server with retry logic"""
        max_retries = 3
        
        # Set timeout for Neo4j startup
        self.neo4j_startup_timed_out = False
        
        # Create timeout event
        def timeout_handler():
            log_message("Neo4j server startup timed out", True)
            self.neo4j_startup_timed_out = True
            
            # Post event to main thread
            if hasattr(self, 'db_init_complete_callback') and self.db_init_complete_callback:
                wx.PostEvent(self, DbInitEvent(success=False, error="Neo4j server startup timed out"))
        
        # Start timeout timer
        timeout = threading.Timer(120, timeout_handler)
        timeout.daemon = True
        timeout.start()
        
        try:
            # Check if we should preserve Neo4j data
            preserve_data = os.environ.get('PRESERVE_NEO4J_DATA', 'True').lower() == 'true'
            
            # Check for a preservation marker file
            preserve_marker = os.path.join(APP_PATH, "Neo4jDB", ".preserve")
            if os.path.exists(preserve_marker):
                preserve_data = True
                log_message("Neo4j data preservation marker file found")
                
            for attempt in range(max_retries):
                # Check if timeout occurred
                if self.neo4j_startup_timed_out:
                    log_message("Aborting Neo4j startup due to timeout", True)
                    return
                
                # For first attempt, do a complete reset only if not preserving data
                if attempt == 0 and not preserve_data:
                    log_message("Performing complete Neo4j server reset for version upgrade...")
                    self.neo4j_server.stop()
                    
                    # Remove existing Neo4j installation
                    if os.path.exists(self.neo4j_server.server_dir):
                        try:
                            shutil.rmtree(self.neo4j_server.server_dir)
                            log_message("Removed existing Neo4j installation")
                        except Exception as e:
                            log_message(f"Error removing Neo4j directory: {str(e)}", True)
                            
                    # Clean up any database files
                    self.neo4j_server.cleanup_database_files()
                elif attempt == 0 and preserve_data:
                    log_message("Neo4j data preservation is enabled, skipping server reset...")
                    # Just ensure the server is stopped
                    self.neo4j_server.stop()
                    # Only clean up lock files
                    self.neo4j_server.cleanup_database_files()
                
                # Start the Neo4j server
                if self.neo4j_server.start():
                    # Server started successfully, now initialize the database manager
                    try:
                        # Use a slight delay to ensure the server is ready
                        time.sleep(5)
                        
                        # Check if timeout occurred
                        if self.neo4j_startup_timed_out:
                            log_message("Aborting Neo4j initialization due to timeout", True)
                            return
                        
                        # Initialize database manager with auth disabled
                        log_message("Attempting to connect to Neo4j...")
                        self.db_manager = Neo4jDatabaseManager(
                            uri=f"bolt://localhost:{self.neo4j_server.NEO4J_PORT}",
                            username="neo4j",
                            password="neo4j",
                            database="neo4j"
                        )
                        
                        # Test connection
                        connection_retries = 3
                        for conn_attempt in range(connection_retries):
                            try:
                                if self.db_manager.connect():
                                    log_message("Connected to Neo4j database successfully")
                                    
                                    # Initialize embeddings
                                    self.initialize_embeddings()
                                    
                                    # Cancel timeout
                                    timeout.cancel()
                                    
                                    # Signal that initialization is complete
                                    if hasattr(self, 'db_init_complete_callback') and self.db_init_complete_callback:
                                        wx.PostEvent(self, DbInitEvent(success=True))
                                    
                                    return True
                                else:
                                    log_message(f"Database connection test failed on attempt {conn_attempt+1}/{connection_retries}", True)
                            except Exception as e:
                                log_message(f"Error testing database connection on attempt {conn_attempt+1}/{connection_retries}: {str(e)}", True)
                                
                            # Short delay before next connection attempt
                            time.sleep(2)
                    except Exception as e:
                        log_message(f"Error initializing database manager: {str(e)}", True)
                        log_message(traceback.format_exc(), True)
                
                # If we got here, there was a problem with this attempt
                log_message(f"Neo4j startup attempt {attempt+1}/{max_retries} failed, retrying...", True)
                
                # Stop any running server
                self.neo4j_server.stop()
                
                # Wait before next attempt
                time.sleep(3)
            
            # If we got here, all attempts failed
            log_message("All Neo4j startup attempts failed", True)
            
            # Signal failure
            if hasattr(self, 'db_init_complete_callback') and self.db_init_complete_callback:
                wx.PostEvent(self, DbInitEvent(success=False, error="Failed to start Neo4j server after multiple attempts"))
                
            # Cancel timeout
            timeout.cancel()
            
            return False
            
        except Exception as e:
            log_message(f"Error in Neo4j startup process: {str(e)}", True)
            log_message(traceback.format_exc(), True)
            
            # Signal failure
            if hasattr(self, 'db_init_complete_callback') and self.db_init_complete_callback:
                wx.PostEvent(self, DbInitEvent(success=False, error=f"Error: {str(e)}"))
                
            # Cancel timeout
            timeout.cancel()
            
            return False
    
    def initialize_embeddings(self):
        """Initialize embeddings and vector store"""
        try:
            if not NEO4J_AVAILABLE or not LANGCHAIN_AVAILABLE or not EMBEDDINGS_AVAILABLE:
                log_message("Required packages for Neo4j/RAG are not available", True)
                # Post event to main thread
                wx.PostEvent(self, DbInitEvent(success=False, error="Required packages not available"))
                return
            
            log_message("Initializing Neo4j database manager...")
            
            # Get embeddings provider
            embeddings = self.get_embeddings_provider()
            if not embeddings:
                log_message("Failed to initialize embeddings provider", True)
                # Post event to main thread
                wx.PostEvent(self, DbInitEvent(success=False, error="Failed to initialize embeddings"))
                return
            
            # Initialize Neo4j database manager
            self.neo4j_manager = Neo4jDatabaseManager(
                uri=f"bolt://localhost:{self.neo4j_server.NEO4J_PORT}",
                username="neo4j",
                password="neo4j",
                database="neo4j"
            )
            
            # Initialize vector store
            self.neo4j_manager.initialize_vector_store(embeddings)
            
            # Mark as initialized
            self.db_initialized = True
            log_message("Database initialization completed successfully")
            
            # Post success event to main thread
            wx.PostEvent(self, DbInitEvent(success=True))
        except Exception as e:
            log_message(f"Error initializing embeddings: {str(e)}", True)
            log_message(traceback.format_exc(), True)
            # Post failure event to main thread
            wx.PostEvent(self, DbInitEvent(success=False, error=str(e)))
    
    def get_embeddings_provider(self):
        """Initialize and return an embeddings provider"""
        try:
            from langchain.embeddings import OpenAIEmbeddings
            from langchain_community.embeddings import FakeEmbeddings
            
            # Try to use OpenAI embeddings if the API key is available
            openai_key = os.environ.get("OPENAI_API_KEY", "")
            
            if openai_key:
                log_message("Using OpenAI embeddings")
                return OpenAIEmbeddings(model="text-embedding-3-small")
            else:
                # Use fake embeddings for testing if no API key
                log_message("Using fake embeddings (no OpenAI API key found)")
                return FakeEmbeddings(size=1536)
        except Exception as e:
            log_message(f"Error initializing embeddings provider: {str(e)}", True)
            return None
    
    def force_java_config(self):
        """Run the force_java_config.py script"""
        try:
            log_message("Running force_java_config.py script...")
            
            # Show a status dialog
            status_dialog = wx.MessageDialog(
                self,
                "Fixing Neo4j Java configuration...\n"
                "The application will restart when complete.",
                "Fixing Configuration",
                wx.OK | wx.ICON_INFORMATION
            )
            status_dialog.ShowModal()
            status_dialog.Destroy()
            
            # Run the force_java_config.py script
            force_config_path = os.path.join(self.base_path, "force_java_config.py")
            if not os.path.exists(force_config_path):
                log_message("force_java_config.py script not found", True)
                wx.MessageBox(
                    "Could not find the force_java_config.py script. "
                    "Please ensure the script is in the application directory.",
                    "Script Not Found",
                    wx.OK | wx.ICON_ERROR
                )
                return
            
            # Make sure it's executable
            if platform.system().lower() != "windows":
                os.chmod(force_config_path, 0o755)
            
            # Run the script and wait for completion
            process = subprocess.run(
                [sys.executable, force_config_path],
                capture_output=True,
                text=True,
                check=False
            )
            
            if process.returncode == 0:
                log_message("Java configuration fixed successfully")
                wx.MessageBox(
                    "Neo4j Java configuration has been updated successfully.\n"
                    "The application will now restart.",
                    "Configuration Updated",
                    wx.OK | wx.ICON_INFORMATION
                )
                self.restart_application()
            else:
                log_message(f"Error running force_java_config.py: {process.stderr}", True)
                wx.MessageBox(
                    f"Failed to update Neo4j Java configuration.\n\n"
                    f"Error: {process.stderr}\n\n"
                    "You can continue using the application without Neo4j features.",
                    "Configuration Failed",
                    wx.OK | wx.ICON_ERROR
                )
        except Exception as e:
            log_message(f"Error running force_java_config.py: {str(e)}", True)
            log_message(traceback.format_exc(), True)
            wx.MessageBox(
                f"Error running force_java_config.py: {str(e)}\n\n"
                "You can continue using the application without Neo4j features.",
                "Error",
                wx.OK | wx.ICON_ERROR
            )
    
    def install_java_for_neo4j(self):
        """Install Java 11 for Neo4j - safe for UI threading"""
        try:
            # Start Java installation in a background thread
            def run_installer():
                log_message("Starting automatic Java 11 installation...")
                success = self.neo4j_server.install_java()
                
                # Use wx.CallAfter to update UI from the main thread
                if success:
                    log_message("Java installation completed successfully")
                    wx.CallAfter(self.restart_application)
                else:
                    log_message("Failed to install Java 11", True)
                    wx.CallAfter(
                        wx.MessageBox,
                        "Failed to install Java 11.\n"
                        "You can continue using the application without Neo4j features,\n"
                        "or install Java 11 manually and restart the application.",
                        "Installation Failed",
                        wx.OK | wx.ICON_ERROR
                    )
            
            # Start the thread
            install_thread = threading.Thread(target=run_installer, daemon=True)
            install_thread.start()
            
        except Exception as e:
            log_message(f"Error initiating Java installation: {str(e)}", True)
            wx.MessageBox(
                f"Error starting Java installation: {str(e)}\n"
                "You can continue using the application without Neo4j features.",
                "Installation Error",
                wx.OK | wx.ICON_ERROR
            )
            
    def restart_application(self):
        """Restart the application after Java installation"""
        log_message("Restarting application after Java installation")
        try:
            # Stop Neo4j if it's running
            if hasattr(self, 'neo4j_server'):
                self.neo4j_server.stop()
            
            # Close Neo4j connection if open
            if self.neo4j_manager:
                self.neo4j_manager.close()
            
            # Get the path to the current executable
            if getattr(sys, 'frozen', False):
                # Running as compiled app
                app_path = sys.executable
                subprocess.Popen([app_path])
            else:
                # Running from script
                app_path = sys.executable
                script_path = os.path.abspath(__file__)
                subprocess.Popen([app_path, script_path])
            
            # Exit the current instance
            self.Destroy()
            sys.exit(0)
        except Exception as e:
            log_message(f"Error restarting application: {str(e)}", True)
            wx.MessageBox(
                "Failed to restart application. Please restart manually.",
                "Restart Failed",
                wx.OK | wx.ICON_ERROR
            )
    
    def on_db_init_event(self, event):
        """Handle database initialization event"""
        if event.success:
            log_message("Database initialization successful")
            self.db_initialized = True
            
            # Initialize RAG chains
            self.initialize_rag_chains()
        else:
            log_message(f"Database initialization failed: {event.error}", True)
            
            # Even though Neo4j failed, we'll set up a minimal UI to allow document functionality
            # without the RAG/GraphRAG features
            wx.MessageBox(
                f"Neo4j database initialization failed: {event.error}\n\n"
                "You can still use basic document functionality but RAG features will be disabled.",
                "Database Error",
                wx.OK | wx.ICON_WARNING
            )
            
            # Create an empty placeholder for the DB manager if needed
            if not hasattr(self, 'db_manager') or self.db_manager is None:
                class DummyManager:
                    def __init__(self):
                        self.connected = False
                        
                    def close(self):
                        pass
                        
                self.db_manager = DummyManager()
            
            # Setup basic UI without database features
            if hasattr(self, 'rag_toggle'):
                self.rag_toggle.Enable(False)
            
            # Create a basic handler for the rag chain to avoid errors
            def dummy_rag_chain(query):
                return "RAG features are disabled because Neo4j database failed to initialize."
                
            self.rag_chain = dummy_rag_chain
    
    def initialize_rag_chains(self):
        """Initialize RAG chains for document querying"""
        try:
            if not self.neo4j_manager or not self.neo4j_manager.vector_store:
                log_message("RAG chains not initialized - vector store not available", True)
                return False

            log_message("Initializing combined RAG chain")
            
            # Get a client to use for the chain
            llm_client = self.get_llm_client()
            if not llm_client:
                log_message("RAG chain not initialized - LLM client not available", True)
                return False
                
            # Create combined RAG chain
            self.rag_chain = create_rag_chain(self.neo4j_manager, llm_client)
            if not self.rag_chain:
                log_message("Combined RAG chain initialization failed", True)
                wx.CallAfter(lambda: self.db_status.SetLabel("Database: Error"))
                return False
            else:
                log_message("Combined RAG chain initialized successfully")
                wx.CallAfter(lambda: self.db_status.SetLabel("Database: Ready"))
                return True
                
        except Exception as e:
            log_message(f"Error initializing RAG chain: {str(e)}", True)
            log_message(traceback.format_exc(), True)
            wx.CallAfter(lambda: self.db_status.SetLabel("Database: Error"))
            return False
    
    def on_rag_toggle(self, event):
        """Handle the RAG toggle checkbox"""
        try:
            is_enabled = self.rag_toggle.GetValue()
            log_message(f"RAG toggle set to: {is_enabled}")
            
            # Check if Neo4j is available before allowing RAG to be enabled
            if is_enabled and not self.db_initialized:
                log_message("Warning: Cannot enable RAG because Neo4j database is not available", True)
                wx.MessageBox(
                    "RAG features require a working Neo4j database, which is not available.\n\n"
                    "Possible reasons:\n"
                    "1. Java 11 or later is not installed (Neo4j requirement)\n"
                    "2. Neo4j server failed to start\n"
                    "3. Required packages are missing\n\n"
                    "To use RAG features, please install Java 11 or higher and restart the application.",
                    "RAG Unavailable",
                    wx.OK | wx.ICON_WARNING
                )
                # Turn off the toggle
                self.rag_toggle.SetValue(False)
                return
            
            if is_enabled and not self.neo4j_manager:
                log_message("Warning: RAG enabled but database not available", True)
                wx.MessageBox("RAG is enabled but the database is not available.\n"
                              "Document context will still be used, but RAG retrieval will be limited.",
                              "Warning", wx.OK | wx.ICON_WARNING)
        except Exception as e:
            log_message(f"Error in RAG toggle handler: {str(e)}", True)
            log_message(traceback.format_exc(), True)
    
    def on_close(self, event):
        """Handle window close event - stop Neo4j server"""
        try:
            # Stop the Neo4j server if it's running
            if hasattr(self, 'neo4j_server'):
                self.neo4j_server.stop()
            
            # Close Neo4j connection if open
            if self.neo4j_manager:
                self.neo4j_manager.close()
        except Exception as e:
            log_message(f"Error during shutdown: {str(e)}", True)
        
        # Destroy the window
        self.Destroy()
    
    def setup_ui(self):
        try:
            log_message("Setting up wxPython UI components")
            
            # Create status bar
            self.CreateStatusBar()
            self.SetStatusText("Ready")
            
            # Main panel
            panel = wx.Panel(self)
            
            # Main sizer
            main_sizer = wx.BoxSizer(wx.HORIZONTAL)
            
            # Left panel (Documents) - 1/3 of width
            left_panel = wx.Panel(panel)
            left_sizer = wx.BoxSizer(wx.VERTICAL)
            
            # Document section title
            doc_title = wx.StaticText(left_panel, label="Documents")
            font = doc_title.GetFont()
            font.SetPointSize(14)
            font.SetWeight(wx.FONTWEIGHT_BOLD)
            doc_title.SetFont(font)
            left_sizer.Add(doc_title, 0, wx.ALL, 10)
            
            # Upload button
            upload_btn = wx.Button(left_panel, label="Upload Document")
            upload_btn.Bind(wx.EVT_BUTTON, self.on_upload_document)
            left_sizer.Add(upload_btn, 0, wx.EXPAND | wx.LEFT | wx.RIGHT | wx.BOTTOM, 10)
            
            # Set document priorities button
            priority_btn = wx.Button(left_panel, label="Set Document Priorities")
            priority_btn.Bind(wx.EVT_BUTTON, self.on_set_priorities)
            left_sizer.Add(priority_btn, 0, wx.EXPAND | wx.LEFT | wx.RIGHT | wx.BOTTOM, 10)
            
            # RAG Controls Section (new)
            rag_panel = wx.Panel(left_panel)
            rag_sizer = wx.BoxSizer(wx.VERTICAL)
            
            # RAG title
            rag_title = wx.StaticText(rag_panel, label="Knowledge Base (RAG)")
            rag_title.SetFont(font)  # Reuse font from above
            rag_sizer.Add(rag_title, 0, wx.ALL, 10)
            
            # RAG toggle
            self.rag_toggle = wx.CheckBox(rag_panel, label="Enable RAG/GraphRAG")
            self.rag_toggle.SetToolTip("Enable Retrieval Augmented Generation with document database")
            self.rag_toggle.Bind(wx.EVT_CHECKBOX, self.on_rag_toggle)
            rag_sizer.Add(self.rag_toggle, 0, wx.LEFT | wx.RIGHT | wx.BOTTOM, 10)
            
            # Database actions panel
            db_actions_panel = wx.Panel(rag_panel)
            db_actions_sizer = wx.BoxSizer(wx.HORIZONTAL)
            
            # Upload to Database button
            upload_to_db_btn = wx.Button(db_actions_panel, label="Upload to Database")
            upload_to_db_btn.SetToolTip("Upload selected documents to database")
            upload_to_db_btn.Bind(wx.EVT_BUTTON, self.on_upload_to_database)
            db_actions_sizer.Add(upload_to_db_btn, 1, wx.RIGHT, 5)
            
            # Delete from Database button
            delete_from_db_btn = wx.Button(db_actions_panel, label="Delete from Database")
            delete_from_db_btn.SetToolTip("Delete selected documents from database")
            delete_from_db_btn.Bind(wx.EVT_BUTTON, self.on_delete_from_database)
            db_actions_sizer.Add(delete_from_db_btn, 1, wx.RIGHT, 5)

            # Delete All Documents button
            delete_all_btn = wx.Button(db_actions_panel, label="Delete All Documents")
            delete_all_btn.SetToolTip("Delete all documents from both Neo4j and vector databases")
            delete_all_btn.Bind(wx.EVT_BUTTON, self.on_delete_all_documents)
            db_actions_sizer.Add(delete_all_btn, 1)
            
            db_actions_panel.SetSizer(db_actions_sizer)
            rag_sizer.Add(db_actions_panel, 0, wx.EXPAND | wx.LEFT | wx.RIGHT | wx.BOTTOM, 10)

            # Show Database Documents button
            show_db_docs_btn = wx.Button(rag_panel, label="Show Database Documents")
            show_db_docs_btn.SetToolTip("Show documents in vector database and Neo4j database")
            show_db_docs_btn.Bind(wx.EVT_BUTTON, self.on_show_database_documents)
            rag_sizer.Add(show_db_docs_btn, 0, wx.EXPAND | wx.LEFT | wx.RIGHT | wx.BOTTOM, 10)
            
            # GraphRAG/RAG selection
            rag_type_panel = wx.Panel(rag_panel)
            rag_type_sizer = wx.BoxSizer(wx.HORIZONTAL)
            
            # Add an explanatory text instead
            rag_info = wx.StaticText(rag_type_panel, label="Documents will be added to both vector database and knowledge graph.")
            rag_type_sizer.Add(rag_info, 1, wx.EXPAND)
            
            rag_type_panel.SetSizer(rag_type_sizer)
            rag_sizer.Add(rag_type_panel, 0, wx.EXPAND | wx.LEFT | wx.RIGHT | wx.BOTTOM, 10)
            
            # Database status
            self.db_status = wx.StaticText(rag_panel, label="Database: Initializing...")
            rag_sizer.Add(self.db_status, 0, wx.LEFT | wx.RIGHT | wx.BOTTOM, 10)
            
            rag_panel.SetSizer(rag_sizer)
            left_sizer.Add(rag_panel, 0, wx.EXPAND | wx.ALL, 10)
            
            # Document list in a scrolled panel
            self.doc_panel = scrolled.ScrolledPanel(left_panel, style=wx.SUNKEN_BORDER)
            self.doc_panel.SetAutoLayout(True)
            self.doc_panel.SetupScrolling()
            
            self.doc_sizer = wx.BoxSizer(wx.VERTICAL)
            self.doc_panel.SetSizer(self.doc_sizer)
            left_sizer.Add(self.doc_panel, 1, wx.EXPAND | wx.LEFT | wx.RIGHT | wx.BOTTOM, 10)
            
            left_panel.SetSizer(left_sizer)
            main_sizer.Add(left_panel, 1, wx.EXPAND | wx.ALL, 10)
            
            # Right panel (Chat) - 2/3 of width
            right_panel = wx.Panel(panel)
            right_sizer = wx.BoxSizer(wx.VERTICAL)
            
            # Chat section title
            chat_title = wx.StaticText(right_panel, label="Research Assistant Chat")
            chat_title.SetFont(font)  # Reuse font from above
            right_sizer.Add(chat_title, 0, wx.ALL, 10)
            
            # Model selection
            model_panel = wx.Panel(right_panel)
            model_sizer = wx.BoxSizer(wx.HORIZONTAL)
            
            model_label = wx.StaticText(model_panel, label="Model:")
            model_sizer.Add(model_label, 0, wx.ALIGN_CENTER_VERTICAL | wx.RIGHT, 5)
            
            # Get list of model names for the dropdown
            model_names = [model_info["name"] for model_key, model_info in self.config["models"].items()]
            
            # Set default model
            default_model_key = self.config.get("default_model", "openai")
            default_model_name = ""
            if default_model_key in self.config["models"]:
                default_model_name = self.config["models"][default_model_key]["name"]
            elif model_names:
                default_model_name = model_names[0]
            
            self.model_choice = wx.Choice(model_panel, choices=model_names)
            if default_model_name in model_names:
                self.model_choice.SetStringSelection(default_model_name)
            model_sizer.Add(self.model_choice, 1, wx.EXPAND)
            
            # Add settings button
            settings_btn = wx.Button(model_panel, label="Settings")
            settings_btn.SetToolTip("Configure API keys and other settings")
            settings_btn.Bind(wx.EVT_BUTTON, self.on_open_settings)
            model_sizer.Add(settings_btn, 0, wx.LEFT, 5)
            
            model_panel.SetSizer(model_sizer)
            right_sizer.Add(model_panel, 0, wx.EXPAND | wx.LEFT | wx.RIGHT | wx.BOTTOM, 10)
            
            # Chat management buttons
            chat_management_panel = wx.Panel(right_panel)
            chat_management_sizer = wx.BoxSizer(wx.HORIZONTAL)

            clear_all_btn = wx.Button(chat_management_panel, label="Clear All Chat")
            clear_all_btn.Bind(wx.EVT_BUTTON, self.on_clear_all_chat)
            chat_management_sizer.Add(clear_all_btn, 1, wx.RIGHT, 5)

            clear_last_btn = wx.Button(chat_management_panel, label="Clear Last Exchange")
            clear_last_btn.Bind(wx.EVT_BUTTON, self.on_clear_last_exchange)
            chat_management_sizer.Add(clear_last_btn, 1, wx.RIGHT, 5)

            edit_msg_btn = wx.Button(chat_management_panel, label="Edit Message")
            edit_msg_btn.Bind(wx.EVT_BUTTON, self.on_edit_message)
            chat_management_sizer.Add(edit_msg_btn, 1)

            chat_management_panel.SetSizer(chat_management_sizer)
            right_sizer.Add(chat_management_panel, 0, wx.EXPAND | wx.LEFT | wx.RIGHT | wx.BOTTOM, 10)
            
            # Prompt library buttons (new)
            prompt_panel = wx.Panel(right_panel)
            prompt_sizer = wx.BoxSizer(wx.HORIZONTAL)
            
            save_prompt_btn = wx.Button(prompt_panel, label="Save Prompt")
            save_prompt_btn.Bind(wx.EVT_BUTTON, self.on_save_prompt)
            prompt_sizer.Add(save_prompt_btn, 1, wx.RIGHT, 5)
            
            load_prompt_btn = wx.Button(prompt_panel, label="Load Prompt")
            load_prompt_btn.Bind(wx.EVT_BUTTON, self.on_load_prompt)
            prompt_sizer.Add(load_prompt_btn, 1)
            
            prompt_panel.SetSizer(prompt_sizer)
            right_sizer.Add(prompt_panel, 0, wx.EXPAND | wx.LEFT | wx.RIGHT | wx.BOTTOM, 10)
            
            # Chat display
            self.chat_display = wx.TextCtrl(right_panel, style=wx.TE_MULTILINE | wx.TE_READONLY | wx.TE_RICH2 | wx.BORDER_SUNKEN)
            right_sizer.Add(self.chat_display, 2, wx.EXPAND | wx.LEFT | wx.RIGHT, 10)
            
            # User input label
            input_label = wx.StaticText(right_panel, label="Your message:")
            right_sizer.Add(input_label, 0, wx.LEFT | wx.RIGHT | wx.TOP, 10)
            
            # User input area
            self.user_input = wx.TextCtrl(right_panel, style=wx.TE_MULTILINE | wx.BORDER_SUNKEN)
            right_sizer.Add(self.user_input, 1, wx.EXPAND | wx.LEFT | wx.RIGHT, 10)
            
            # Send button
            send_btn = wx.Button(right_panel, label="Send")
            send_btn.Bind(wx.EVT_BUTTON, self.on_send_message)
            right_sizer.Add(send_btn, 0, wx.EXPAND | wx.ALL, 10)
            
            right_panel.SetSizer(right_sizer)
            main_sizer.Add(right_panel, 2, wx.EXPAND | wx.ALL, 10)
            
            panel.SetSizer(main_sizer)
            
            log_message("wxPython UI setup complete")
        except Exception as e:
            log_message(f"Error setting up UI: {str(e)}", True)
            log_message(traceback.format_exc(), True)
            wx.MessageBox(f"Error setting up UI: {str(e)}", "Error", wx.OK | wx.ICON_ERROR)
    
    def on_upload_document(self, event):
        """Handle document upload button event"""
        try:
            # Create and show the file dialog
            wildcard = "All files (*.*)|*.*|PDF Files (*.pdf)|*.pdf|DOCX Files (*.docx)|*.docx|Text Files (*.txt)|*.txt|Markdown Files (*.md)|*.md"
            dialog = wx.FileDialog(
                self, message="Choose a file to upload",
                defaultDir=os.getcwd(),
                defaultFile="",
                wildcard=wildcard,
                style=wx.FD_OPEN | wx.FD_FILE_MUST_EXIST | wx.FD_MULTIPLE
            )
            
            if dialog.ShowModal() == wx.ID_CANCEL:
                return
            
            # Get the selected path(s)
            paths = dialog.GetPaths()
            dialog.Destroy()
            
            # Documents directory
            documents_dir = os.path.join(self.base_path, "Documents")
            os.makedirs(documents_dir, exist_ok=True)
            
            # Process each selected file
            for path in paths:
                try:
                    # Get the filename
                    filename = os.path.basename(path)
                    
                    # Check if document with this name is already in the list
                    if filename in self.documents:
                        # Ask user if they want to replace it
                        msg = f"Document '{filename}' already exists. Replace it?"
                        dlg = wx.MessageDialog(self, msg, "Confirm Replace", wx.YES_NO | wx.ICON_QUESTION)
                        result = dlg.ShowModal()
                        dlg.Destroy()
                        
                        if result != wx.ID_YES:
                            continue
                    
                    # Copy the file to the Documents directory (if not already there)
                    dest_path = os.path.join(documents_dir, filename)
                    if path != dest_path:
                        shutil.copy2(path, dest_path)
                    
                    # Load the document into memory (simplified for now)
                    file_extension = os.path.splitext(filename)[1].lower()
                    content = ""
                    
                    # Use appropriate method to read file based on extension
                    if file_extension == '.pdf':
                        try:
                            # Try loading with pypdf first
                            import pypdf
                            with open(dest_path, 'rb') as f:
                                pdf_reader = pypdf.PdfReader(f)
                                for page in pdf_reader.pages:
                                    content += page.extract_text() + "\n"
                            
                            # If content is too short, try alternate extraction methods
                            if len(content.strip()) < 100:
                                log_message(f"Limited text extracted from PDF {filename}, trying alternate methods")
                                
                                # Try alternative PDF extraction if available
                                try:
                                    # Try to import langchain PDF loader
                                    from langchain_community.document_loaders import PyPDFLoader
                                    pdf_doc = PyPDFLoader(dest_path).load()
                                    content = ""
                                    for page in pdf_doc:
                                        content += page.page_content + "\n"
                                    log_message(f"Used langchain PyPDFLoader for {filename}")
                                except ImportError:
                                    log_message(f"langchain PyPDFLoader not available", True)
                        except ImportError:
                            # Try langchain if pypdf is not installed
                            try:
                                from langchain_community.document_loaders import PyPDFLoader
                                pdf_doc = PyPDFLoader(dest_path).load()
                                content = ""
                                for page in pdf_doc:
                                    content += page.page_content + "\n"
                                log_message(f"Used langchain PyPDFLoader for {filename} (pypdf not available)")
                            except ImportError:
                                # Fallback if langchain is not installed
                                with open(dest_path, 'rb') as f:
                                    content = f"[PDF content not extracted - required modules not available]\n\nFile: {filename}"
                        
                        # If content is still very limited, warn the user
                        if len(content.strip()) < 100:
                            log_message(f"Warning: Limited text extracted from PDF {filename}. PDF may be scanned or have restricted permissions.", True)
                            content += "\n[Note: Limited text could be extracted from this PDF. It may be a scanned document or have security restrictions.]"
                    elif file_extension == '.docx':
                        try:
                            import docx
                            doc = docx.Document(dest_path)
                            for para in doc.paragraphs:
                                content += para.text + "\n"
                        except ImportError:
                            # Fallback if python-docx is not installed
                            content = f"[DOCX content not extracted - docx module not available]\n\nFile: {filename}"
                    else:
                        # Plain text or markdown
                        try:
                            with open(dest_path, 'r', encoding='utf-8') as f:
                                content = f.read()
                        except UnicodeDecodeError:
                            # Try another encoding if UTF-8 fails
                            try:
                                with open(dest_path, 'r', encoding='latin-1') as f:
                                    content = f.read()
                            except Exception as e:
                                content = f"[Error reading file: {str(e)}]\n\nFile: {filename}"
                    
                    # Note: We're not automatically adding documents to the database anymore.
                    # Users should select documents and click the "Upload to Database" button to add them.
                    
                    # Store document with default priority
                    self.documents[filename] = content
                    if filename not in self.document_priorities:
                        self.document_priorities[filename] = "Medium"
                    
                    # Add document to the UI
                    self.add_document_to_ui(filename)
                    
                    log_message(f"Document uploaded: {filename}")
                except Exception as e:
                    log_message(f"Error processing document '{path}': {str(e)}", True)
                    log_message(traceback.format_exc(), True)
                    wx.MessageBox(
                        f"Error processing document '{os.path.basename(path)}':\n{str(e)}",
                        "Upload Error",
                        wx.OK | wx.ICON_ERROR
                    )
        except Exception as e:
            log_message(f"Error in document upload: {str(e)}", True)
            log_message(traceback.format_exc(), True)
            wx.MessageBox(
                f"Error uploading document:\n{str(e)}",
                "Upload Error",
                wx.OK | wx.ICON_ERROR
            )
    
    def add_document_to_ui(self, filename):
        """Add a document to the UI list"""
        try:
            # Create a panel for this document
            doc_item_panel = wx.Panel(self.doc_panel)
            doc_item_sizer = wx.BoxSizer(wx.HORIZONTAL)
            
            # Document checkbox
            doc_check = wx.CheckBox(doc_item_panel, label=filename)
            doc_item_sizer.Add(doc_check, 1, wx.EXPAND | wx.RIGHT, 5)
            
            # Delete button
            delete_btn = wx.Button(doc_item_panel, label="X", size=(25, 25))
            delete_btn.SetToolTip("Remove document")
            delete_btn.Bind(wx.EVT_BUTTON, lambda evt, name=filename: self.on_delete_document(evt, name))
            doc_item_sizer.Add(delete_btn, 0)
            
            doc_item_panel.SetSizer(doc_item_sizer)
            
            # Add to documents panel
            self.doc_sizer.Add(doc_item_panel, 0, wx.EXPAND | wx.ALL, 5)
            
            # Refresh the UI
            self.doc_panel.Layout()
            self.doc_sizer.Layout()
        except Exception as e:
            log_message(f"Error adding document to UI: {str(e)}", True)
    
    def on_delete_document(self, event, filename):
        """Handle document deletion"""
        try:
            # Confirm deletion
            dlg = wx.MessageDialog(
                self, 
                f"Are you sure you want to remove '{filename}'?", 
                "Confirm Deletion", 
                wx.YES_NO | wx.ICON_QUESTION
            )
            result = dlg.ShowModal()
            dlg.Destroy()
            
            if result != wx.ID_YES:
                return
            
            # Remove from memory
            if filename in self.documents:
                del self.documents[filename]
            
            if filename in self.document_priorities:
                del self.document_priorities[filename]
            
            # Remove from database if RAG is enabled
            if self.db_initialized and self.neo4j_manager:
                # Create a document ID consistent with upload
                import hashlib
                doc_id = hashlib.md5(filename.encode()).hexdigest()
                self.neo4j_manager.remove_document(doc_id)
            
            # Refresh the document list
            self.refresh_document_list()
            
            log_message(f"Document removed: {filename}")
        except Exception as e:
            log_message(f"Error deleting document: {str(e)}", True)
            log_message(traceback.format_exc(), True)
    
    def on_upload_to_database(self, event):
        """Upload selected documents to the database without deleting them first"""
        try:
            # Check if database is ready
            if not self.db_initialized or not self.neo4j_manager:
                wx.MessageBox(
                    "Database is not initialized. Please make sure Neo4j is running and try again.",
                    "Database Not Ready",
                    wx.OK | wx.ICON_ERROR
                )
                return
            
            # Get selected documents
            selected_docs = []
            for child in self.doc_panel.GetChildren():
                for grandchild in child.GetChildren():
                    if isinstance(grandchild, wx.CheckBox) and grandchild.GetValue():
                        filename = grandchild.GetLabel()
                        if filename in self.documents:
                            selected_docs.append(filename)
            
            if not selected_docs:
                wx.MessageBox(
                    "No documents selected. Please select documents to upload to the database.",
                    "No Selection",
                    wx.OK | wx.ICON_INFORMATION
                )
                return
            
            # Show progress dialog
            progress_dlg = wx.ProgressDialog(
                "Uploading Documents",
                f"Uploading {len(selected_docs)} documents to database...",
                maximum=len(selected_docs),
                parent=self,
                style=wx.PD_APP_MODAL | wx.PD_AUTO_HIDE
            )
            
            # Upload each selected document
            successful = 0
            for i, filename in enumerate(selected_docs):
                progress_dlg.Update(i, f"Uploading {filename}...")
                content = self.documents[filename]
                
                # Skip if content is empty or just a placeholder message
                if not content or content.startswith("[PDF content not extracted"):
                    log_message(f"Skipping '{filename}' - no content could be extracted", True)
                    wx.MessageBox(
                        f"Cannot upload '{filename}' because no content could be extracted from it.\n\nPlease try a different PDF reader software to extract the text first, then save it as a text file and upload that instead.",
                        "Upload Error",
                        wx.OK | wx.ICON_ERROR
                    )
                    continue
                
                # Create a unique document ID
                import hashlib
                doc_id = hashlib.md5(filename.encode()).hexdigest()
                
                # Add to database
                log_message(f"Adding document {doc_id} to database: {filename}")
                
                # Get file extension to check if it's a PDF
                file_extension = os.path.splitext(filename)[1].lower()
                if file_extension == '.pdf':
                    log_message(f"Processing PDF document: {filename}")
                
                # Add to vector store
                success = self.neo4j_manager.add_document(
                    document_id=doc_id,
                    title=filename,
                    content=content,
                    metadata={
                        "filename": filename, 
                        "file_type": file_extension[1:] if file_extension else "unknown",
                        "priority": self.document_priorities.get(filename, "Medium")
                    }
                )
                
                if success:
                    successful += 1
                    log_message(f"Added document '{filename}' to database")
                else:
                    log_message(f"Failed to add document '{filename}' to database", True)
            
            progress_dlg.Destroy()
            
            # Show result message
            if successful > 0:
                wx.MessageBox(
                    f"Successfully uploaded {successful} of {len(selected_docs)} documents to the database.",
                    "Upload Complete",
                    wx.OK | wx.ICON_INFORMATION
                )
            else:
                wx.MessageBox(
                    "Failed to upload any documents to the database. Please check the logs for details.",
                    "Upload Failed",
                    wx.OK | wx.ICON_ERROR
                )
        except Exception as e:
            log_message(f"Error uploading documents to database: {str(e)}", True)
            log_message(traceback.format_exc(), True)
            wx.MessageBox(
                f"Error uploading documents to database: {str(e)}",
                "Upload Error",
                wx.OK | wx.ICON_ERROR
            )
    
    def on_delete_from_database(self, event):
        """Delete selected documents from the database without removing them from the UI"""
        try:
            # Check if database is ready
            if not self.db_initialized or not self.neo4j_manager:
                wx.MessageBox(
                    "Database is not initialized. Please make sure Neo4j is running and try again.",
                    "Database Not Ready",
                    wx.OK | wx.ICON_ERROR
                )
                return
            
            # Get selected documents
            selected_docs = []
            for child in self.doc_panel.GetChildren():
                for grandchild in child.GetChildren():
                    if isinstance(grandchild, wx.CheckBox) and grandchild.GetValue():
                        filename = grandchild.GetLabel()
                        if filename in self.documents:
                            selected_docs.append(filename)
            
            if not selected_docs:
                wx.MessageBox(
                    "No documents selected. Please select documents to delete from the database.",
                    "No Selection",
                    wx.OK | wx.ICON_INFORMATION
                )
                return
            
            # Confirm deletion
            dlg = wx.MessageDialog(
                self, 
                f"Are you sure you want to delete {len(selected_docs)} documents from the database? This will not remove them from your local documents list.",
                "Confirm Database Deletion", 
                wx.YES_NO | wx.ICON_QUESTION
            )
            result = dlg.ShowModal()
            dlg.Destroy()
            
            if result != wx.ID_YES:
                return
            
            # Show progress dialog
            progress_dlg = wx.ProgressDialog(
                "Deleting Documents",
                f"Deleting {len(selected_docs)} documents from database...",
                maximum=len(selected_docs),
                parent=self,
                style=wx.PD_APP_MODAL | wx.PD_AUTO_HIDE
            )
            
            # Delete each selected document from the database
            successful = 0
            for i, filename in enumerate(selected_docs):
                progress_dlg.Update(i, f"Deleting {filename}...")
                
                # Create a document ID consistent with upload
                import hashlib
                doc_id = hashlib.md5(filename.encode()).hexdigest()
                
                if self.neo4j_manager.remove_document(doc_id):
                    successful += 1
                    log_message(f"Removed document '{filename}' from database")
                else:
                    log_message(f"Failed to remove document '{filename}' from database", True)
            
            progress_dlg.Destroy()
            
            # Show result message
            if successful > 0:
                wx.MessageBox(
                    f"Successfully deleted {successful} of {len(selected_docs)} documents from the database.",
                    "Deletion Complete",
                    wx.OK | wx.ICON_INFORMATION
                )
            else:
                wx.MessageBox(
                    "Failed to delete any documents from the database. Please check the logs for details.",
                    "Deletion Failed",
                    wx.OK | wx.ICON_ERROR
                )
        except Exception as e:
            log_message(f"Error deleting documents from database: {str(e)}", True)
            log_message(traceback.format_exc(), True)
            wx.MessageBox(
                f"Error deleting documents from database: {str(e)}",
                "Deletion Error",
                wx.OK | wx.ICON_ERROR
            )
    
    def on_delete_all_documents(self, event):
        """Handle delete all documents button event"""
        try:
            if not self.neo4j_manager or not self.neo4j_manager.connected:
                wx.MessageBox(
                    "Not connected to Neo4j database. Please ensure the database is running.",
                    "Connection Error",
                    wx.OK | wx.ICON_ERROR
                )
                return

            # Show confirmation dialog
            confirm = wx.MessageBox(
                "Are you sure you want to delete ALL documents from both Neo4j and vector databases?\n\n"
                "This action cannot be undone, but you can upload the documents again later.",
                "Confirm Delete All",
                wx.YES_NO | wx.ICON_WARNING
            )

            if confirm == wx.YES:
                # Delete all documents
                success = self.neo4j_manager.delete_all_documents()
                
                if success:
                    wx.MessageBox(
                        "All documents have been successfully deleted from both databases.",
                        "Deletion Complete",
                        wx.OK | wx.ICON_INFORMATION
                    )
                    # Refresh the document list
                    self.refresh_document_list()
                else:
                    wx.MessageBox(
                        "Failed to delete all documents. Please check the logs for details.",
                        "Deletion Failed",
                        wx.OK | wx.ICON_ERROR
                    )
        except Exception as e:
            log_message(f"Error deleting all documents: {str(e)}", True)
            log_message(traceback.format_exc(), True)
            wx.MessageBox(
                f"Error deleting all documents: {str(e)}",
                "Deletion Error",
                wx.OK | wx.ICON_ERROR
            )
    
    def refresh_document_list(self):
        """Refresh the document list in the UI"""
        try:
            # Clear the current document list
            self.doc_sizer.Clear(True)
            
            # Re-add all documents
            for filename in sorted(self.documents.keys()):
                self.add_document_to_ui(filename)
            
            # Refresh the UI
            self.doc_panel.Layout()
            self.doc_sizer.Layout()
        except Exception as e:
            log_message(f"Error refreshing document list: {str(e)}", True)
    
    def on_set_priorities(self, event):
        """Handle setting document priorities"""
        try:
            if not self.documents:
                wx.MessageBox("No documents to prioritize", "No Documents", wx.OK | wx.ICON_INFORMATION)
                return
            
            # Show the dialog
            dialog = DocumentPriorityDialog(self, list(self.documents.keys()), self.document_priorities)
            result = dialog.ShowModal()
            
            if result == wx.ID_OK:
                # Get the updated priorities
                self.document_priorities = dialog.get_priorities()
                log_message("Document priorities updated")
            
            dialog.Destroy()
        except Exception as e:
            log_message(f"Error setting document priorities: {str(e)}", True)
            log_message(traceback.format_exc(), True)
            wx.MessageBox(
                f"Error setting document priorities:\n{str(e)}",
                "Priority Error",
                wx.OK | wx.ICON_ERROR
            )
    
    def on_save_prompt(self, event):
        """Handle saving a prompt to the library"""
        try:
            # Get the current prompt text
            current_prompt = self.user_input.GetValue().strip()
            
            if not current_prompt:
                wx.MessageBox("Please enter a prompt to save", "Empty Prompt", wx.OK | wx.ICON_INFORMATION)
                return
            
            # Show the dialog
            dialog = PromptLibraryDialog(self, mode="save", current_prompt=current_prompt)
            result = dialog.ShowModal()
            
            if result == wx.ID_OK:
                # Get the prompt name and save it
                prompt_name = dialog.get_prompt_name()
                if prompt_name:
                    # Ensure prompts directory exists
                    prompts_dir = os.path.join(self.base_path, "Prompts")
                    os.makedirs(prompts_dir, exist_ok=True)
                    
                    # Save the prompt as JSON
                    prompt_path = os.path.join(prompts_dir, f"{prompt_name}.json")
                    prompt_data = {"content": current_prompt}
                    with open(prompt_path, 'w', encoding='utf-8') as f:
                        json.dump(prompt_data, f, ensure_ascii=False, indent=2)
                    
                    log_message(f"Prompt saved as: {prompt_name}")
                    wx.MessageBox(f"Prompt saved as: {prompt_name}", "Prompt Saved", wx.OK | wx.ICON_INFORMATION)
            
            dialog.Destroy()
        except Exception as e:
            log_message(f"Error saving prompt: {str(e)}", True)
            log_message(traceback.format_exc(), True)
            wx.MessageBox(
                f"Error saving prompt:\n{str(e)}",
                "Save Error",
                wx.OK | wx.ICON_ERROR
            )
    
    def on_load_prompt(self, event):
        # Open prompt library dialog in load mode
        dialog = PromptLibraryDialog(self, mode="load")
        if dialog.ShowModal() == wx.ID_OK:
            prompt_content = dialog.get_prompt_content()
            if prompt_content:
                self.user_input.SetValue(prompt_content)
        dialog.Destroy()
        
    def on_open_settings(self, event):
        """Open the settings dialog"""
        try:
            dialog = SettingsDialog(self, self.config)
            if dialog.ShowModal() == wx.ID_OK:
                # Refresh the UI or apply changes as needed
                log_message("Settings saved successfully")
                
                # Optionally reload environment variables to ensure they're available
                load_env_variables()
            dialog.Destroy()
        except Exception as e:
            log_message(f"Error opening settings dialog: {str(e)}", True)
            wx.MessageBox(
                f"Error opening settings dialog: {str(e)}",
                "Settings Error",
                wx.OK | wx.ICON_ERROR
            )
    
    def on_clear_all_chat(self, event):
        """Clear all chat history"""
        try:
            # Confirm before clearing
            dlg = wx.MessageDialog(
                self, 
                "Are you sure you want to clear the entire conversation?", 
                "Confirm Clear All", 
                wx.YES_NO | wx.ICON_QUESTION
            )
            result = dlg.ShowModal()
            dlg.Destroy()
            
            if result != wx.ID_YES:
                return
                
            # Clear conversation history
            self.conversation_history = []
            
            # Clear chat display
            self.chat_display.Clear()
            self.message_positions = []
            
            log_message("Chat history cleared")
        except Exception as e:
            log_message(f"Error clearing chat: {str(e)}", True)
            log_message(traceback.format_exc(), True)
            wx.MessageBox(
                f"Error clearing chat:\n{str(e)}",
                "Clear Error",
                wx.OK | wx.ICON_ERROR
            )
    
    def on_clear_last_exchange(self, event):
        try:
            if len(self.conversation_history) >= 2:
                # Remove last assistant and user messages (one exchange)
                self.conversation_history = self.conversation_history[:-2]
                
                # Redraw chat display
                self.chat_display.Clear()
                self.message_positions = []
                
                for msg in self.conversation_history:
                    sender = "You" if msg["role"] == "user" else "Assistant"
                    self.append_to_chat(msg["content"], sender)
                
                log_message("Last exchange cleared")
                self.SetStatusText("Last exchange cleared")
            else:
                self.SetStatusText("No complete exchanges to clear")
        except Exception as e:
            error_msg = f"Error clearing last exchange: {str(e)}"
            log_message(error_msg, True)
            log_message(traceback.format_exc(), True)
            self.SetStatusText(error_msg)

    def on_edit_message(self, event):
        try:
            # First, store the current position of each message to enable selection
            if not self.message_positions:
                # If we haven't tracked positions yet, we need to rebuild this list
                self.rebuild_message_positions()
            
            # Create a dialog to select which message to edit
            messages = []
            for i, msg in enumerate(self.conversation_history):
                role = "You" if msg["role"] == "user" else "Assistant"
                # Truncate message for display in selection dialog
                content = msg["content"]
                if len(content) > 50:
                    content = content[:47] + "..."
                messages.append(f"{i+1}. {role}: {content}")
            
            # Show message selection dialog
            dialog = wx.SingleChoiceDialog(
                self, "Select a message to edit:", "Edit Message", messages)
            
            if dialog.ShowModal() == wx.ID_OK:
                selected_index = dialog.GetSelection()
                dialog.Destroy()
                
                # Now show edit dialog for the selected message
                edit_dialog = MessageEditDialog(
                    self, self.conversation_history[selected_index]["content"])
                
                if edit_dialog.ShowModal() == wx.ID_OK:
                    # Update the message with edited content
                    edited_content = edit_dialog.GetMessage()
                    self.conversation_history[selected_index]["content"] = edited_content
                    
                    # Store original message role
                    edited_message_role = self.conversation_history[selected_index]["role"]
                    
                    # Redraw chat display up to the edited message
                    self.chat_display.Clear()
                    self.message_positions = []
                    
                    # Display all messages up to and including the edited one
                    for i, msg in enumerate(self.conversation_history[:selected_index+1]):
                        sender = "You" if msg["role"] == "user" else "Assistant"
                        self.append_to_chat(msg["content"], sender)
                    
                    # Remove all messages after the edited one
                    self.conversation_history = self.conversation_history[:selected_index+1]
                    
                    # Only process the message if it's a user message
                    if edited_message_role == "user":
                        # Disable input during processing
                        self.user_input.Disable()
                        self.SetStatusText("Processing edited message...")
                        
                        # Create a copy of the edited message content for processing
                        message_to_process = edited_content
                        
                        # Process the edited message in a separate thread to avoid UI freezing
                        threading.Thread(
                            target=self.process_edited_message,
                            args=(message_to_process, selected_index),
                            daemon=True
                        ).start()
                    else:
                        # If editing an assistant message, just update the UI
                        log_message("Assistant message edited, not reprocessing")
                        self.SetStatusText("Message edited successfully")
                
                edit_dialog.Destroy()
            else:
                dialog.Destroy()
        except Exception as e:
            error_msg = f"Error editing message: {str(e)}"
            log_message(error_msg, True)
            log_message(traceback.format_exc(), True)
            self.SetStatusText(error_msg)
    
    def process_edited_message(self, message_content, message_index):
        """Process an edited message with proper error handling for all LLM types"""
        try:
            # Get LLM client instance
            llm_client = self.get_llm_client()
            if not llm_client:
                error_msg = "Could not initialize language model. Please check your configuration."
                wx.CallAfter(self.handle_edit_error, error_msg)
                return
            
            # Initialize streaming response container
            self.current_streaming_response = ""
            
            # Create a system prompt with conversation context (shortened version)
            system_prompt = self.config.get("system_prompt", "You are a helpful research assistant.")
            
            # Process streaming response
            response_chunks = []
            
            # Define the callback for streaming updates
            def update_response_callback(chunk):
                # Post event to main thread for UI update
                # Make sure we're passing a string to the StreamEvent
                if isinstance(chunk, str):
                    evt = StreamEvent(text=chunk)
                    wx.PostEvent(self, evt)
                else:
                    # If not a string, convert it
                    evt = StreamEvent(text=str(chunk))
                    wx.PostEvent(self, evt)
            
            # Check if RAG is enabled - mirroring the process_message function logic
            if hasattr(self, 'rag_toggle') and self.rag_toggle.GetValue():
                if self.rag_chain:
                    log_message("Using RAG for edited message response")
                    try:
                        # For RAG responses, we need to use streaming if available
                        # or explicitly handle the response
                        response = self.rag_chain(message_content)
                        # If it's not a streaming response, manually handle it
                        wx.CallAfter(self.handle_edited_response, response)
                        return
                    except Exception as rag_error:
                        log_message(f"Error using RAG chain: {str(rag_error)}", True)
                        # Fall back to direct query if RAG fails
            
            # If RAG isn't enabled or failed, use document context approach
            # Get selected documents
            selected_docs = []
            for child in self.doc_panel.GetChildren():
                for grandchild in child.GetChildren():
                    if isinstance(grandchild, wx.CheckBox) and grandchild.GetValue():
                        filename = grandchild.GetLabel()
                        if filename in self.documents:
                            content = self.documents[filename]
                            # Get priority and sort accordingly
                            priority = self.document_priorities.get(filename, "Medium")
                            priority_value = {"High": 3, "Medium": 2, "Low": 1}.get(priority, 2)
                            selected_docs.append((filename, content, priority_value))
            
            # Sort docs by priority (highest first)
            selected_docs.sort(key=lambda x: x[2], reverse=True)
            
            # Build document context
            doc_context = ""
            if selected_docs:
                doc_context = "Here are the relevant documents:\n\n"
                for filename, content, _ in selected_docs:
                    doc_context += f"[{filename}]\n{content}\n\n"
            
            # Build full prompt with document context
            full_prompt = f"{system_prompt}\n\n{doc_context}User: {message_content}\n\nAssistant: I'll combine information from the documents with my knowledge to give you the most helpful answer."
            
            # Try to process with the appropriate model
            try:
                for chunk in llm_client.generate_streaming(full_prompt, callback=update_response_callback):
                    response_chunks.append(chunk)
                
                # Combine chunks for the final response
                response = "".join(response_chunks)
                
                # Use CallAfter to update UI from the thread
                wx.CallAfter(self.handle_edited_response, response)
                
            except Exception as model_error:
                error_msg = f"Error processing with selected model: {str(model_error)}"
                log_message(error_msg, True)
                
                # Try a fallback approach if necessary
                wx.CallAfter(self.handle_edit_error, error_msg)
                
        except Exception as e:
            error_msg = f"Error processing edited message: {str(e)}"
            log_message(error_msg, True)
            log_message(traceback.format_exc(), True)
            wx.CallAfter(self.handle_edit_error, error_msg)
    
    def handle_edited_response(self, response):
        """Handle the response from the language model after an edit"""
        try:
            # Handle AIMessage objects from LangChain
            if hasattr(response, 'content') and callable(getattr(response, 'content', None)):
                # This is for older versions of langchain
                response = response.content()
            elif hasattr(response, 'content') and not callable(getattr(response, 'content', None)):
                # This is for newer versions of langchain
                response = response.content
                
            # Only append to chat if it's not already displayed via streaming
            already_streamed = (hasattr(self, 'current_streaming_response') and 
                               self.current_streaming_response is not None and 
                               self.current_streaming_response != "")
            
            if not already_streamed:
                # No streaming happened, so we need to add the full response
                self.append_to_chat(response, "Assistant")
                
            # Reset streaming state regardless
            self.current_streaming_response = None
            
            # Add to conversation history
            self.conversation_history.append({"role": "assistant", "content": response})
            
            # Re-enable input and update status
            self.user_input.Enable()
            self.SetStatusText("Message edited successfully")
            
            # Log completion of edit
            log_message("Edited message response handled successfully")
            
        except Exception as e:
            error_msg = f"Error handling edited response: {str(e)}"
            log_message(error_msg, True)
            log_message(traceback.format_exc(), True)
            self.handle_edit_error(error_msg)
    
    def handle_edit_error(self, error_msg):
        """Handle errors during the edit process"""
        # Re-enable input
        self.user_input.Enable()
        # Update status
        self.SetStatusText(f"Edit error: {error_msg}")
        # Log the error
        log_message(f"Edit error: {error_msg}", True)
    
    def rebuild_message_positions(self):
        try:
            self.message_positions = []
            current_position = 0
            
            # Get the text of the chat display
            text = self.chat_display.GetValue()
            lines = text.split('\n')
            
            # Rebuild positions based on message markers in text
            for i, line in enumerate(lines):
                # Check if this line starts a message (contains "You: " or "Assistant: ")
                if line.startswith("You: ") or line.startswith("Assistant: "):
                    # Calculate position in characters
                    position = sum(len(lines[j]) + 1 for j in range(i))
                    self.message_positions.append(position)
            
            log_message(f"Rebuilt message positions: {len(self.message_positions)} found")
        except Exception as e:
            log_message(f"Error rebuilding message positions: {str(e)}", True)
            log_message(traceback.format_exc(), True)

    def edit_conversation_history(self, event):
        try:
            # Create a dialog to edit conversation history
            dialog = ConversationHistoryDialog(self, self.conversation_history)
            if dialog.ShowModal() == wx.ID_OK:
                # Update conversation history
                self.conversation_history = dialog.get_updated_history()
                
                # Update chat display
                self.chat_display.Clear()
                for msg in self.conversation_history:
                    sender = "You" if msg["role"] == "user" else "Assistant"
                    self.append_to_chat(msg["content"], sender)
                
                log_message("Conversation history updated")
                self.SetStatusText("Conversation history updated")
            
            dialog.Destroy()
        except Exception as e:
            error_msg = f"Error editing conversation history: {str(e)}"
            log_message(error_msg, True)
            log_message(traceback.format_exc(), True)
            self.SetStatusText(error_msg)
    
    def on_send_message(self, event):
        try:
            user_message = self.user_input.GetValue().strip()
            
            if not user_message:
                return
            
            log_message(f"Sending user message: {user_message[:50]}...")
            
            # Clear input
            self.user_input.Clear()
            
            # Add to chat display
            self.append_to_chat(user_message, "You")
            
            # Add to conversation history
            self.conversation_history.append({"role": "user", "content": user_message})
            
            # Disable input during processing
            self.user_input.Disable()
            self.SetStatusText("Processing...")
            
            # Process in a thread and handle the response
            def process_and_handle_response():
                try:
                    response = self.process_message(user_message)
                    # Use CallAfter to update UI from the thread
                    wx.CallAfter(self.handle_response, response)
                except Exception as e:
                    error_msg = f"Error processing message: {str(e)}"
                    log_message(error_msg, True)
                    log_message(traceback.format_exc(), True)
                    wx.CallAfter(self.handle_response, f"Error: {error_msg}")
            
            # Start the processing thread
            threading.Thread(target=process_and_handle_response, daemon=True).start()
        except Exception as e:
            error_msg = f"Error sending message: {str(e)}"
            log_message(error_msg, True)
            log_message(traceback.format_exc(), True)
            self.SetStatusText(error_msg)
            self.user_input.Enable()

    def on_stream_event(self, event):
        """Handle streaming chunks, including those from edited messages"""
        try:
            # Append the chunk to the chat display
            self.append_streaming_chunk(event.text)
            
            # Add to the current streaming response
            if hasattr(self, 'current_streaming_response'):
                self.current_streaming_response += event.text
            else:
                self.current_streaming_response = event.text
                
            # Ensure UI is updated immediately
            wx.GetApp().Yield()
        except Exception as e:
            log_message(f"Error handling stream event: {str(e)}", True)
            log_message(traceback.format_exc(), True)
    
    def append_to_chat(self, message, sender):
        """Append a message to the chat display"""
        try:
            # Handle AIMessage objects from LangChain
            if hasattr(message, 'content') and callable(getattr(message, 'content', None)):
                # This is for older versions of langchain
                message = message.content()
            elif hasattr(message, 'content') and not callable(getattr(message, 'content', None)):
                # This is for newer versions of langchain
                message = message.content
            
            # Get current text position for indexing/editing
            current_position = len(self.chat_display.GetValue())
            if current_position > 0 and not self.chat_display.GetValue().endswith("\n\n"):
                self.chat_display.AppendText("\n\n")
                current_position += 2
            elif current_position > 0:
                self.chat_display.AppendText("\n")
                current_position += 1
                
            # Add sender with bold formatting
            self.chat_display.AppendText(f"{sender}: ")
            
            # Set the style for the sender (bold)
            end_pos = len(self.chat_display.GetValue())
            self.chat_display.SetStyle(current_position, end_pos, wx.TextAttr(wx.NullColour, wx.NullColour, wx.Font(wx.NORMAL_FONT.GetPointSize(), wx.FONTFAMILY_DEFAULT, wx.FONTSTYLE_NORMAL, wx.FONTWEIGHT_BOLD)))
            
            # Add the message text
            self.chat_display.AppendText(str(message))
            
            # Store the position of this message for editing purposes
            self.message_positions.append(current_position)
            
            # Scroll to the bottom
            self.chat_display.ShowPosition(self.chat_display.GetLastPosition())
        except Exception as e:
            log_message(f"Error appending to chat: {str(e)}", True)
    
    def append_streaming_chunk(self, chunk):
        """Append a streaming chunk to the chat display"""
        try:
            # Handle AIMessage objects from LangChain
            if hasattr(chunk, 'content') and callable(getattr(chunk, 'content', None)):
                # This is for older versions of langchain
                chunk = chunk.content()
            elif hasattr(chunk, 'content') and not callable(getattr(chunk, 'content', None)):
                # This is for newer versions of langchain
                chunk = chunk.content
            
            # Ensure chunk is a string
            chunk = str(chunk)
            
            # If this is the first chunk, add a new message
            if not hasattr(self, 'current_streaming_response') or self.current_streaming_response is None or self.current_streaming_response == "":
                # Init streaming container if needed
                if not hasattr(self, 'current_streaming_response'):
                    self.current_streaming_response = ""
                
                # Add a new message from the assistant
                current_position = len(self.chat_display.GetValue())
                if current_position > 0 and not self.chat_display.GetValue().endswith("\n\n"):
                    self.chat_display.AppendText("\n\n")
                elif current_position > 0:
                    self.chat_display.AppendText("\n")
                
                # Add sender with bold formatting
                self.chat_display.AppendText("Assistant: ")
                
                # Store the position of this message
                self.message_positions.append(current_position)
            
            # Append the new chunk
            self.chat_display.AppendText(chunk)
            
            # Scroll to the bottom
            self.chat_display.ShowPosition(self.chat_display.GetLastPosition())
            
            # Force UI update
            wx.GetApp().Yield()
        except Exception as e:
            log_message(f"Error appending streaming chunk: {str(e)}", True)
            log_message(traceback.format_exc(), True)
    
    def get_llm_client(self):
        """Get the LLM client based on the selected model"""
        try:
            # Get the selected model from combobox
            model_name = self.model_choice.GetStringSelection()
            
            # Get model config from the config
            model_config = None
            model_key = None
            for key, config in self.config.get("models", {}).items():
                if config.get("name") == model_name:
                    model_config = config
                    model_key = key
                    break
            
            if not model_config:
                log_message(f"No configuration found for model: {model_name}", True)
                return None
            
            # Get the API key from environment
            api_key_env = model_config.get("api_key_env")
            api_key = os.environ.get(api_key_env) if api_key_env else None
            
            # Get the model ID
            model_id = model_config.get("model_name", "")
            
            if not api_key:
                log_message(f"API key environment variable {api_key_env} not set", True)
                return None
            
            # Create the appropriate client
            client_class = LLMClient(api_key, model_id, model_key=model_key)
            return client_class
            
        except Exception as e:
            log_message(f"Error getting LLM client: {str(e)}", True)
            log_message(traceback.format_exc(), True)
            return None
            
    def process_message(self, user_message):
        """Process a user message and generate a response"""
        try:
            # Get LLM client instance
            llm_client = self.get_llm_client()
            if not llm_client:
                return "Error: Could not initialize language model. Please check your configuration."
            
            # Define a callback function for streaming updates
            def update_response_callback(chunk):
                # Post event to main thread for UI update
                evt = StreamEvent(text=chunk)
                wx.PostEvent(self, evt)
                
            # Check if RAG is enabled
            if hasattr(self, 'rag_toggle') and self.rag_toggle.GetValue():
                if self.rag_chain:
                    log_message("Using RAG for response")
                    response = self.rag_chain(user_message)
                else:
                    log_message("No RAG chain available, falling back to direct query", True)
                    # Fall back to direct query
                    # Initialize streaming response
                    self.current_streaming_response = ""
                    
                    # Use streaming response with callback
                    system_prompt = self.config.get("system_prompt", "You are a helpful research assistant.")
                    full_prompt = f"{system_prompt}\n\nUser: {user_message}\n\nAssistant: I'll combine information from the documents with my knowledge to give you the most helpful answer."
                    
                    # Process streaming response
                    response_chunks = []
                    for chunk in llm_client.generate_streaming(full_prompt, callback=update_response_callback):
                        response_chunks.append(chunk)
                        
                    response = "".join(response_chunks)
            else:
                # Regular prompt-based approach when RAG is not enabled
                log_message("Processing message with standard context")
                
                # Get selected documents
                selected_docs = []
                for child in self.doc_panel.GetChildren():
                    for grandchild in child.GetChildren():
                        if isinstance(grandchild, wx.CheckBox) and grandchild.GetValue():
                            filename = grandchild.GetLabel()
                            if filename in self.documents:
                                content = self.documents[filename]
                                # Get priority and sort accordingly
                                priority = self.document_priorities.get(filename, "Medium")
                                priority_value = {"High": 3, "Medium": 2, "Low": 1}.get(priority, 2)
                                selected_docs.append((filename, content, priority_value))
                
                # Sort docs by priority (highest first)
                selected_docs.sort(key=lambda x: x[2], reverse=True)
                
                # Build document context
                doc_context = ""
                if selected_docs:
                    doc_context = "Here are the relevant documents:\n\n"
                    for filename, content, _ in selected_docs:
                        doc_context += f"[{filename}]\n{content}\n\n"
                
                # Build full prompt
                system_prompt = self.config.get("system_prompt", "You are a helpful research assistant.")
                full_prompt = f"{system_prompt}\n\n{doc_context}User: {user_message}\n\nAssistant: I'll combine information from the documents with my knowledge to give you the most helpful answer."
                
                # Initialize streaming response
                self.current_streaming_response = ""
                
                # Use streaming response with callback
                response_chunks = []
                for chunk in llm_client.generate_streaming(full_prompt, callback=update_response_callback):
                    response_chunks.append(chunk)
                    
                response = "".join(response_chunks)
            
            return response
        except Exception as e:
            log_message(f"Error processing message: {str(e)}", True)
            log_message(traceback.format_exc(), True)
            return f"Error: {str(e)}"

    def handle_response(self, response):
        """Handle the response from the language model"""
        try:
            # Handle AIMessage objects from LangChain
            original_response = response
            if hasattr(response, 'content') and callable(getattr(response, 'content', None)):
                # This is for older versions of langchain
                response = response.content()
            elif hasattr(response, 'content') and not callable(getattr(response, 'content', None)):
                # This is for newer versions of langchain
                response = response.content
                
            # Only append to chat if it's not a streaming response
            if not hasattr(self, 'current_streaming_response') or self.current_streaming_response is None:
                self.append_to_chat(response, "Assistant")
            else:
                # Reset streaming state
                self.current_streaming_response = None
            
            # Add to conversation history - preserve the original structure if it was an AIMessage
            if hasattr(original_response, 'content'):
                self.conversation_history.append({"role": "assistant", "content": response})
            else:
                self.conversation_history.append({"role": "assistant", "content": original_response})
            
            # Re-enable input and update status
            self.user_input.Enable()
            self.SetStatusText("Ready")
        except Exception as e:
            log_message(f"Error handling response: {str(e)}", True)
            self.user_input.Enable()
            self.SetStatusText(f"Error: {str(e)}")

    def on_show_database_documents(self, event):
        """Show documents in both vector database and Neo4j database"""
        try:
            if not self.db_initialized or not self.neo4j_manager:
                wx.MessageBox(
                    "Database not initialized. Please wait for initialization to complete.",
                    "Database Error",
                    wx.OK | wx.ICON_ERROR
                )
                return

            # Get documents from Neo4j database
            docs = self.neo4j_manager.get_document_list()
            
            if not docs:
                wx.MessageBox(
                    "No documents found in the database.",
                    "No Documents",
                    wx.OK | wx.ICON_INFORMATION
                )
                return

            # Create a dialog to display the documents
            dialog = wx.Dialog(self, title="Documents in Database", size=(800, 600))
            panel = wx.Panel(dialog)
            sizer = wx.BoxSizer(wx.VERTICAL)

            # Create a text control with a scrollbar
            text_ctrl = wx.TextCtrl(panel, style=wx.TE_MULTILINE | wx.TE_READONLY | wx.TE_RICH2 | wx.HSCROLL)
            
            # Format the document list
            doc_text = f"Documents in Database ({len(docs)} total):\n\n"
            for doc_id, title, chunks in docs:
                doc_text += f"Title: {title}\n"
                doc_text += f"ID: {doc_id}\n"
                doc_text += "-" * 80 + "\n"

            text_ctrl.SetValue(doc_text)
            sizer.Add(text_ctrl, 1, wx.EXPAND | wx.ALL, 10)

            # Add a close button
            close_btn = wx.Button(panel, label="Close")
            close_btn.Bind(wx.EVT_BUTTON, lambda e: dialog.EndModal(wx.ID_OK))
            sizer.Add(close_btn, 0, wx.ALIGN_CENTER | wx.ALL, 10)

            panel.SetSizer(sizer)
            dialog.ShowModal()
            dialog.Destroy()

        except Exception as e:
            log_message(f"Error showing database documents: {str(e)}", True)
            wx.MessageBox(
                f"Error showing database documents:\n{str(e)}",
                "Error",
                wx.OK | wx.ICON_ERROR
            )

# Dialog for editing conversation history
class ConversationHistoryDialog(wx.Dialog):
    def __init__(self, parent, conversation_history):
        super(ConversationHistoryDialog, self).__init__(
            parent, title="Edit Conversation History", size=(800, 600)
        )
        
        self.conversation_history = conversation_history.copy()
        
        # Create main sizer
        main_sizer = wx.BoxSizer(wx.VERTICAL)
        
        # Instructions
        instructions = wx.StaticText(self, label="Edit or delete conversation messages:")
        main_sizer.Add(instructions, 0, wx.ALL, 10)
        
        # Create a scrolled panel for the messages
        self.panel = scrolled.ScrolledPanel(self)
        self.panel.SetAutoLayout(True)
        self.panel.SetupScrolling()
        
        self.panel_sizer = wx.BoxSizer(wx.VERTICAL)
        self.message_editors = []
        
        # Add each message to the panel
        for i, msg in enumerate(self.conversation_history):
            self.add_message_editor(i, msg)
        
        self.panel.SetSizer(self.panel_sizer)
        main_sizer.Add(self.panel, 1, wx.EXPAND | wx.ALL, 10)
        
        # Buttons
        button_sizer = wx.BoxSizer(wx.HORIZONTAL)
        
        ok_button = wx.Button(self, wx.ID_OK, "Save Changes")
        cancel_button = wx.Button(self, wx.ID_CANCEL, "Cancel")
        
        button_sizer.Add(ok_button, 0, wx.ALL, 5)
        button_sizer.Add(cancel_button, 0, wx.ALL, 5)
        
        main_sizer.Add(button_sizer, 0, wx.ALIGN_RIGHT | wx.ALL, 10)
        
        self.SetSizer(main_sizer)
        self.Centre()
    
    def add_message_editor(self, index, message):
        # Create a panel for this message
        msg_panel = wx.Panel(self.panel)
        msg_sizer = wx.BoxSizer(wx.VERTICAL)
        
        # Role selection
        role_sizer = wx.BoxSizer(wx.HORIZONTAL)
        role_label = wx.StaticText(msg_panel, label="Role:")
        role_choices = ["user", "assistant"]
        role_choice = wx.Choice(msg_panel, choices=["User", "Assistant"])
        role_choice.SetSelection(0 if message["role"] == "user" else 1)
        
        role_sizer.Add(role_label, 0, wx.ALIGN_CENTER_VERTICAL | wx.RIGHT, 5)
        role_sizer.Add(role_choice, 0)
        
        # Delete button
        delete_btn = wx.Button(msg_panel, label="Delete")
        delete_btn.index = index  # Store the index
        delete_btn.Bind(wx.EVT_BUTTON, self.on_delete_message)
        
        role_sizer.Add(delete_btn, 0, wx.LEFT, 10)
        
        msg_sizer.Add(role_sizer, 0, wx.EXPAND | wx.ALL, 5)
        
        # Message content
        content_label = wx.StaticText(msg_panel, label="Content:")
        msg_sizer.Add(content_label, 0, wx.ALL, 5)
        
        content_text = wx.TextCtrl(msg_panel, value=message["content"], style=wx.TE_MULTILINE)
        msg_sizer.Add(content_text, 0, wx.EXPAND | wx.ALL, 5)
        
        msg_panel.SetSizer(msg_sizer)
        self.panel_sizer.Add(msg_panel, 0, wx.EXPAND | wx.ALL | wx.BOTTOM, 10)
        
        # Add separator line
        line = wx.StaticLine(self.panel, style=wx.LI_HORIZONTAL)
        self.panel_sizer.Add(line, 0, wx.EXPAND | wx.ALL, 5)
        
        # Store references to UI elements
        self.message_editors.append({
            "panel": msg_panel,
            "role": role_choice,
            "content": content_text,
            "index": index
        })
    
    def on_delete_message(self, event):
        # Get the index from the button that triggered the event
        index = event.GetEventObject().index
        
        # Find the corresponding editor
        editor = None
        for ed in self.message_editors:
            if ed["index"] == index:
                editor = ed
                break
        
        if editor:
            # Remove the panel and its separator line from the sizer
            self.panel_sizer.Remove(editor["panel"])
            editor["panel"].Destroy()
            
            # Get the index of the editor in the list
            list_index = self.message_editors.index(editor)
            
            # If not the last item, remove the separator line
            if list_index < len(self.message_editors) - 1:
                # The separator line is right after the panel
                next_item = self.panel_sizer.GetItem(list_index * 2 + 1).GetWindow()
                if next_item:
                    self.panel_sizer.Remove(next_item)
                    next_item.Destroy()
            
            # Remove from our list
            self.message_editors.remove(editor)
            
            # Update the UI
            self.panel.Layout()
            self.panel.SetupScrolling()
    
    def get_updated_history(self):
        # Create a new history from the current state of editors
        updated_history = []
        
        for editor in self.message_editors:
            role = "user" if editor["role"].GetSelection() == 0 else "assistant"
            content = editor["content"].GetValue()
            
            updated_history.append({
                "role": role,
                "content": content
            })
        
        return updated_history

# Set up basic error logging
def setup_error_logging():
    try:
        base_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
        log_path = os.path.join(base_dir, "error_log.txt")
        
        with open(log_path, 'a', encoding='utf-8') as f:
            f.write(f"Application starting at {os.path.abspath(sys.argv[0])}\n")
        
        log_message(f"Application starting. Errors will be logged to {log_path}")
        return log_path
    except Exception as e:
        print(f"Error setting up logging: {str(e)}")
        return None

def load_config():
    """Load configuration from config.json file"""
    try:
        config_path = os.path.join(APP_PATH, "config.json")
        if os.path.exists(config_path):
            with open(config_path, 'r') as f:
                config = json.load(f)
            log_message("Configuration loaded successfully")
            return config
        else:
            log_message("Configuration file not found, using default configuration", True)
            return None
    except Exception as e:
        log_message(f"Error loading configuration: {str(e)}", True)
        return None

def create_default_config():
    """Create default configuration"""
    try:
        default_config = {
            "models": {
                "openai": {
                    "name": "OpenAI GPT-4",
                    "api_key_env": "OPENAI_API_KEY",
                    "model_name": "gpt-4o-mini"
                },
                "anthropic": {
                    "name": "Anthropic Claude",
                    "api_key_env": "ANTHROPIC_API_KEY",
                    "model_name": "claude-3-7-sonnet-20250219"
                },
                "gemini": {
                    "name": "Google Gemini-2.0-Flash",
                    "api_key_env": "GOOGLE_API_KEY",
                    "model_name": "gemini-2.0-flash"
                },
            },
            "default_model": "openai",
            "max_tokens": 8000,
            "system_prompt": "You are a helpful AI research assistant. Your goal is to help researchers write new papers or expand work-in-progress papers based on the provided documents and instructions."
        }
        
        # Save the default configuration
        config_path = os.path.join(APP_PATH, "config.json")
        with open(config_path, 'w') as f:
            json.dump(default_config, f, indent=2)
        
        log_message("Default configuration created and saved")
        return default_config
    except Exception as e:
        log_message(f"Error creating default configuration: {str(e)}", True)
        return {
            "models": {"openai": {"name": "OpenAI GPT-4", "api_key_env": "OPENAI_API_KEY", "model_name": "gpt-4"}},
            "default_model": "openai",
            "max_tokens": 8000,
            "system_prompt": "You are a helpful AI research assistant."
        }

def create_rag_chain(neo4j_manager, llm_client):
    """Create a combined RAG chain using both vector store and knowledge graph"""
    try:
        # Check if the requirements are met
        if not LANGCHAIN_AVAILABLE:
            log_message("LangChain is not available, RAG chain cannot be created", True)
            return None
        
        if not neo4j_manager or not hasattr(neo4j_manager, 'connected') or not neo4j_manager.connected:
            log_message("Neo4j manager is not connected, RAG chain cannot be created", True)
            return None
            
        # Import necessary langchain components
        from langchain_core.prompts import ChatPromptTemplate
        from langchain_core.output_parsers import StrOutputParser
        
        # Create prompt template
        prompt = ChatPromptTemplate.from_template("""
        You are a research assistant helping with scholarly papers and documents.
        Below are some relevant passages from academic documents that might help answer the user's question.
        The documents are connected in a knowledge graph, which provides additional context about the relationships.
        
        Relevant passages:
        {context}
        
        Knowledge graph relationships:
        {graph_info}
        
        User's question: {query}
        
        Please provide a comprehensive and accurate response based on the context provided.
        Combine the information from the database with your own knowledge to give the most helpful answer.
        Use your general knowledge to add context, explanations, or examples when relevant, but prioritize the information from the provided documents.
        Cite specific parts of the documents when appropriate by mentioning document titles.
        Leverage the connections between documents and entities to provide a more holistic view.
        If the information isn't in the context, you can still try to answer based on your knowledge but make it clear what comes from the documents and what comes from your general knowledge.
        """)
        
        # Define the combined RAG chain
        def combined_rag_chain(query):
            # Check connection
            if not neo4j_manager or not hasattr(neo4j_manager, 'connected') or not neo4j_manager.connected:
                return "Database connection is not available. Please try again later."
                
            # Get similar documents using vector search first
            vector_docs = []
            try:
                vector_docs = neo4j_manager.query_similar_text(query, limit=5, use_graph=False)
                log_message(f"Found {len(vector_docs)} similar documents using vector search")
            except Exception as e:
                log_message(f"Error in vector search: {str(e)}", True)
            
            # Get related documents using graph-based search
            graph_docs = []
            try:
                graph_docs = neo4j_manager.query_similar_text(query, limit=5, use_graph=True)
                log_message(f"Found {len(graph_docs)} similar documents using graph search")
            except Exception as e:
                log_message(f"Error in graph search: {str(e)}", True)
                # Continue with just vector results if graph search fails
            
            # Combine and deduplicate documents based on document_id
            unique_docs = {}
            for doc in vector_docs + graph_docs:
                doc_id = doc.metadata.get("document_id", "unknown")
                if doc_id not in unique_docs:
                    unique_docs[doc_id] = doc
                # If we already have this document but the current one has higher priority, replace it
                elif doc.metadata.get("priority_value", 0) > unique_docs[doc_id].metadata.get("priority_value", 0):
                    unique_docs[doc_id] = doc
            
            # Get the final list of documents
            relevant_docs = list(unique_docs.values())
            
            # Sort final list by priority value
            relevant_docs = sorted(relevant_docs, 
                                key=lambda doc: doc.metadata.get("priority_value", 2),
                                reverse=True)
            
            # Log information about the retrieved documents
            log_message(f"Retrieved {len(relevant_docs)} total unique documents:")
            for doc in relevant_docs:
                doc_id = doc.metadata.get("document_id", "unknown")
                title = doc.metadata.get("title", "Untitled")
                priority = doc.metadata.get("priority", "Medium")
                log_message(f"  - Document '{title}' (ID: {doc_id}, Priority: {priority})")
            
            # If no documents were found, return a simple response
            if not relevant_docs:
                return "I couldn't find any relevant information in the database for your query. Please try a different question or add more documents to the database."
                
            # Format context string from documents
            context_parts = []
            doc_ids = []
            for i, doc in enumerate(relevant_docs):
                title = doc.metadata.get("title", f"Document {i+1}")
                doc_id = doc.metadata.get("document_id", "unknown")
                doc_ids.append(doc_id)
                context_parts.append(f"[{title}]: {doc.page_content}")
            
            context = "\n\n".join(context_parts)
            
            # Get document relationships from the graph
            graph_info_parts = []
            with neo4j_manager.driver.session() as session:
                # Look for direct document relationships
                for doc_id in doc_ids:
                    result = session.run("""
                    MATCH (d:Document {document_id: $doc_id})-[r]-(other)
                    WHERE other:Document
                    RETURN d.title AS title, type(r) AS relation, other.title AS other_title
                    LIMIT 5
                    """, doc_id=doc_id)
                    
                    for record in result:
                        graph_info_parts.append(
                            f"Document '{record['title']}' {record['relation']} '{record['other_title']}'"
                        )
                
                # Look for entities in the documents and their relationships
                for doc_id in doc_ids:
                    result = session.run("""
                    MATCH (d:Document {document_id: $doc_id})-[:CONTAINS]->(e)
                    RETURN d.title AS doc_title, labels(e)[0] AS entity_type, e.id AS entity_id, e.name AS entity_name
                    LIMIT 10
                    """, doc_id=doc_id)
                    
                    for record in result:
                        entity_name = record['entity_name'] if record['entity_name'] else record['entity_id']
                        graph_info_parts.append(
                            f"Document '{record['doc_title']}' contains {record['entity_type']} '{entity_name}'"
                        )
                    
                    # Find relationships between entities
                    result = session.run("""
                    MATCH (d:Document {document_id: $doc_id})-[:CONTAINS]->(e1)-[r]->(e2)<-[:CONTAINS]-(d)
                    RETURN labels(e1)[0] AS type1, e1.name AS name1, e1.id AS id1, 
                           type(r) AS relation,
                           labels(e2)[0] AS type2, e2.name AS name2, e2.id AS id2
                    LIMIT 10
                    """, doc_id=doc_id)
                    
                    for record in result:
                        name1 = record['name1'] if record['name1'] else record['id1']
                        name2 = record['name2'] if record['name2'] else record['id2']
                        graph_info_parts.append(
                            f"{record['type1']} '{name1}' {record['relation']} {record['type2']} '{name2}'"
                        )
            
            # If no relationships found, extract key terms
            if not graph_info_parts:
                # Extract key terms from documents
                try:
                    import re
                    from collections import Counter
                    
                    # Get all document content
                    all_content = " ".join([doc.page_content for doc in relevant_docs])
                    
                    # Remove stop words and extract key terms
                    stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'with', 'by', 'about', 'of'}
                    words = re.findall(r'\b\w+\b', all_content.lower())
                    key_terms = [w for w in words if w not in stop_words and len(w) > 3]
                    
                    # Get most common terms
                    term_counts = Counter(key_terms)
                    top_terms = term_counts.most_common(5)
                    
                    for term, count in top_terms:
                        graph_info_parts.append(f"Key term '{term}' appears {count} times in the documents")
                except Exception as e:
                    log_message(f"Error extracting key terms: {str(e)}", True)
            
            graph_info = "\n".join(graph_info_parts) if graph_info_parts else "No explicit relationships found between documents."
            
            # Execute prompt with LLM
            model = llm_client.get_model()
            if model:
                formatted_prompt = prompt.format(context=context, graph_info=graph_info, query=query)
                response = model.invoke(formatted_prompt)
                
                # Convert AIMessage to string if needed
                if hasattr(response, 'content'):
                    if callable(getattr(response, 'content', None)):
                        return response.content()
                    else:
                        return response.content
                return response
            else:
                return "Language model is not available. Please check your API configuration."
                
        log_message("Combined RAG chain initialized successfully")
        return combined_rag_chain
    except Exception as e:
        log_message(f"Error creating combined RAG chain: {str(e)}", True)
        log_message(traceback.format_exc(), True)
        return None

# Add the LLMClient class back since it's still being referenced
class LLMClient:
    """Unified interface for different LLM providers"""
    
    def __init__(self, api_key, model_id, model_key=None):
        self.api_key = api_key
        self.model_id = model_id
        self.model_key = model_key or "openai"  # Default to OpenAI
        self.client = None
        self._initialize_client()
        
    def _initialize_client(self):
        """Initialize the appropriate client based on the model key"""
        try:
            if self.model_key == "openai":
                from openai import OpenAI
                self.client = OpenAI(api_key=self.api_key)
            elif self.model_key == "anthropic":
                if check_package_installed("anthropic"):
                    from anthropic import Anthropic
                    self.client = Anthropic(api_key=self.api_key)
                else:
                    log_message("Anthropic package not installed. Installing required packages...", True)
                    subprocess.check_call([sys.executable, "-m", "pip", "install", "anthropic"])
                    log_message("Anthropic package installed successfully. Initializing client...")
                    from anthropic import Anthropic
                    self.client = Anthropic(api_key=self.api_key)
            elif self.model_key == "gemini":
                # Google package installation requires special handling
                try:
                    # First try to import directly
                    log_message("Attempting to import Google packages...")
                    import google.generativeai as genai
                    genai.configure(api_key=self.api_key)
                    self.client = genai
                    log_message("Successfully imported Google Generative AI package")
                except ImportError:
                    # If import fails, install packages one by one
                    log_message("Google packages not found or incomplete. Installing required packages...", True)
                    packages = [
                        "google-api-python-client",
                        "google-api-core",
                        "google-cloud-core",
                        "google-cloud",
                        "google-generativeai"
                    ]
                    
                    # Install each package individually and report success/failure
                    for package in packages:
                        try:
                            log_message(f"Installing {package}...")
                            subprocess.check_call([sys.executable, "-m", "pip", "install", "--upgrade", package])
                            log_message(f"Successfully installed {package}")
                        except Exception as e:
                            log_message(f"Error installing {package}: {str(e)}", True)
                    
                    # Force Python to reload modules
                    log_message("Reloading modules...")
                    if 'google' in sys.modules:
                        del sys.modules['google']
                    if 'google.generativeai' in sys.modules:
                        del sys.modules['google.generativeai']
                    
                    # Try import again after installation
                    try:
                        import google.generativeai as genai
                        genai.configure(api_key=self.api_key)
                        self.client = genai
                        log_message("Successfully imported Google Generative AI package after installation")
                    except ImportError as e:
                        log_message(f"Failed to import Google packages after installation: {str(e)}", True)
                        raise ValueError(f"Could not initialize Google Generative AI. Error: {str(e)}")
            else:
                log_message(f"Unsupported model provider: {self.model_key}", True)
                self.client = None
        except Exception as e:
            log_message(f"Error initializing LLM client: {str(e)}", True)
            self.client = None
            
    def get_model(self):
        """Get a unified model interface for langchain"""
        try:
            if not self.client:
                return None
                
            # Create the appropriate model based on the provider
            if self.model_key == "openai":
                try:
                    from langchain_openai import ChatOpenAI
                    return ChatOpenAI(
                        model=self.model_id,
                        openai_api_key=self.api_key,
                        temperature=0.7
                    )
                except ImportError:
                    log_message("langchain_openai not installed. Installing required package...", True)
                    subprocess.check_call([sys.executable, "-m", "pip", "install", "langchain_openai"])
                    from langchain_openai import ChatOpenAI
                    return ChatOpenAI(
                        model=self.model_id,
                        openai_api_key=self.api_key,
                        temperature=0.7
                    )
            elif self.model_key == "anthropic":
                try:
                    from langchain_anthropic import ChatAnthropic
                    return ChatAnthropic(
                        model=self.model_id,
                        anthropic_api_key=self.api_key,
                        temperature=0.7
                    )
                except ImportError:
                    log_message("langchain_anthropic not installed. Installing required package...", True)
                    subprocess.check_call([sys.executable, "-m", "pip", "install", "langchain_anthropic", "anthropic"])
                    from langchain_anthropic import ChatAnthropic
                    return ChatAnthropic(
                        model=self.model_id,
                        anthropic_api_key=self.api_key,
                        temperature=0.7
                    )
            elif self.model_key == "gemini":
                try:
                    # First check if the Google integration is available
                    log_message("Checking for langchain_google_genai integration...")
                    import importlib.util
                    if importlib.util.find_spec("langchain_google_genai") is None:
                        log_message("langchain_google_genai not found. Installing required packages...", True)
                        required_packages = [
                            "langchain_google_genai",
                            "google-generativeai>=0.3.0",
                            "google-api-python-client",
                            "google-api-core"
                        ]
                        # Install each package individually
                        for package in required_packages:
                            try:
                                log_message(f"Installing {package}...")
                                subprocess.check_call([sys.executable, "-m", "pip", "install", "--upgrade", package])
                                log_message(f"Successfully installed {package}")
                            except Exception as e:
                                log_message(f"Error installing {package}: {str(e)}", True)
                    
                    # Force reload of modules
                    if 'langchain_google_genai' in sys.modules:
                        del sys.modules['langchain_google_genai']
                    if 'google' in sys.modules:
                        del sys.modules['google']
                    if 'google.generativeai' in sys.modules:
                        del sys.modules['google.generativeai']
                    
                    # Try importing and creating the model
                    log_message("Importing langchain_google_genai...")
                    from langchain_google_genai import ChatGoogleGenerativeAI
                    log_message(f"Creating Google model with ID: {self.model_id}")
                    return ChatGoogleGenerativeAI(
                        model=self.model_id,
                        google_api_key=self.api_key,
                        temperature=0.7
                    )
                except ImportError as e:
                    log_message(f"Error importing langchain_google_genai: {str(e)}", True)
                    log_message("Attempting alternative installation methods...", True)
                    try:
                        # Try direct pip install with system call
                        subprocess.check_call([sys.executable, "-m", "pip", "install", "--upgrade", 
                                             "langchain-google-genai", "google-generativeai"])
                        
                        # Import after installation
                        from langchain_google_genai import ChatGoogleGenerativeAI
                        return ChatGoogleGenerativeAI(
                            model=self.model_id,
                            google_api_key=self.api_key,
                            temperature=0.7
                        )
                    except Exception as e2:
                        log_message(f"Failed to install and import langchain_google_genai: {str(e2)}", True)
                        raise ImportError(f"Could not setup Google Generative AI integration: {str(e2)}")
                except Exception as e:
                    log_message(f"Error creating Google Generative AI model: {str(e)}", True)
                    raise ValueError(f"Error creating Google model: {str(e)}")
            else:
                return None
        except Exception as e:
            log_message(f"Error creating langchain model: {str(e)}", True)
            return None
            
    def generate_streaming(self, prompt, callback=None):
        """Generate streaming response from the LLM"""
        try:
            if not self.client:
                raise ValueError("LLM client not initialized")
                
            # Handle different client types
            if self.model_key == "openai":
                # OpenAI streaming
                response = self.client.chat.completions.create(
                    model=self.model_id,
                    messages=[{"role": "user", "content": prompt}],
                    stream=True,
                    max_tokens=4000
                )
                
                # Process the streaming response
                for chunk in response:
                    if chunk.choices and chunk.choices[0].delta.content:
                        content = chunk.choices[0].delta.content
                        if callback:
                            callback(content)
                        yield content
                        
            elif self.model_key == "anthropic":
                # Anthropic streaming
                try:
                    response = self.client.messages.create(
                        model=self.model_id,
                        max_tokens=4000,
                        messages=[{"role": "user", "content": prompt}],
                        stream=True
                    )
                    
                    # Process the streaming response
                    for chunk in response:
                        if chunk.type == 'content_block_delta' and hasattr(chunk, 'delta') and hasattr(chunk.delta, 'text'):
                            content = chunk.delta.text
                            if callback:
                                callback(content)
                            yield content
                except AttributeError as e:
                    # Anthropic API might have changed
                    log_message(f"Error with Anthropic streaming: {str(e)}. Trying alternative API...", True)
                    response = self.client.completions.create(
                        prompt=f"\n\nHuman: {prompt}\n\nAssistant:",
                        model=self.model_id,
                        max_tokens_to_sample=4000,
                        stream=True
                    )
                    for completion in response:
                        if hasattr(completion, 'completion'):
                            content = completion.completion
                            if callback:
                                callback(content)
                            yield content
                        
            elif self.model_key == "gemini":
                # First ensure the necessary packages are installed
                log_message("Ensuring Google packages are installed for streaming...")
                
                if not check_package_installed("google.generativeai"):
                    log_message("Google Generative AI package not found. Installing...", True)
                    try:
                        # Install the package
                        subprocess.check_call([sys.executable, "-m", "pip", "install", "--upgrade", 
                                             "google-generativeai>=0.3.0", 
                                             "google-api-python-client"])
                        
                        # Force reload of modules
                        if 'google' in sys.modules:
                            del sys.modules['google']
                        if 'google.generativeai' in sys.modules:
                            del sys.modules['google.generativeai']
                        
                        # Import the package
                        import google.generativeai as genai
                        genai.configure(api_key=self.api_key)
                        self.client = genai
                    except Exception as e:
                        error_msg = f"Failed to install Google packages: {str(e)}"
                        log_message(error_msg, True)
                        if callback:
                            callback(f"\n\nError: {error_msg}")
                        yield f"\n\nError: {error_msg}"
                        return
                
                # Try to use the client for streaming
                try:
                    # Log the model being used
                    log_message(f"Using Gemini model: {self.model_id}")
                    
                    # Create model instance
                    model = self.client.GenerativeModel(self.model_id)
                    
                    # Configure generation settings
                    generation_config = {
                        "temperature": 0.7,
                        "top_p": 0.95,
                        "top_k": 40,
                        "max_output_tokens": 4000,
                    }
                    
                    # Generate content with streaming
                    response = model.generate_content(
                        prompt,
                        generation_config=generation_config,
                        stream=True
                    )
                    
                    # Process the streaming response
                    for chunk in response:
                        if hasattr(chunk, 'text') and chunk.text:
                            content = chunk.text
                            if callback:
                                callback(content)
                            yield content
                        elif hasattr(chunk, 'parts') and len(chunk.parts) > 0:
                            # Alternative response format
                            for part in chunk.parts:
                                if hasattr(part, 'text') and part.text:
                                    content = part.text
                                    if callback:
                                        callback(content)
                                    yield content
                
                except AttributeError as e:
                    # The API might have changed, try alternative method
                    log_message(f"Attribute error with Gemini streaming: {str(e)}. Trying alternative method...", True)
                    try:
                        # Alternative API approach
                        response = model.start_chat().send_message(prompt, stream=True)
                        for chunk in response:
                            if hasattr(chunk, 'text'):
                                content = chunk.text
                                if callback:
                                    callback(content)
                                yield content
                    except Exception as e2:
                        error_msg = f"Error with alternative Gemini streaming: {str(e2)}"
                        log_message(error_msg, True)
                        if callback:
                            callback(f"\n\nError: {error_msg}")
                        yield f"\n\nError: {error_msg}"
                
                except Exception as e:
                    error_msg = f"Error with Gemini streaming: {str(e)}"
                    log_message(error_msg, True)
                    if callback:
                        callback(f"\n\nError: {error_msg}")
                    yield f"\n\nError: {error_msg}"
            else:
                raise ValueError(f"Unsupported model provider: {self.model_key}")
                
        except Exception as e:
            error_msg = f"Error generating streaming response: {str(e)}"
            log_message(error_msg, True)
            if callback:
                callback(f"\n\nError: {error_msg}")
            yield f"\n\nError: {error_msg}"

# Main entry point
if __name__ == "__main__":
    try:
        # Set up logging
        log_file = setup_error_logging()
        
        # Check for Google packages first - most common issue
        log_message("Checking Google packages installation...")
        if not check_package_installed("google.generativeai"):
            log_message("Google packages not found. Installing them now...", True)
            ensure_google_packages()
        
        # Check for LLM packages and install if missing
        required_packages = {
            "openai": ["openai", "langchain_openai"],
            "anthropic": ["anthropic", "langchain_anthropic"],
            "gemini": [
                "google-api-python-client", 
                "google-api-core", 
                "google-cloud-core", 
                "google-generativeai>=0.3.0", 
                "langchain_google_genai",
                "langchain-google-genai",
                "protobuf>=4.23.0"
            ]
        }
        
        # Check config for default model
        try:
            config = load_config()
            default_model = config.get("default_model", "openai")
            log_message(f"Default model is: {default_model}")
            
            # Install packages for the default model
            if default_model in required_packages:
                missing = []
                for pkg in required_packages[default_model]:
                    # Strip version specifier for checking
                    base_pkg = pkg.split('>=')[0] if '>=' in pkg else pkg
                    base_pkg = base_pkg.replace('-', '_').split('.')[0]
                    if not check_package_installed(base_pkg):
                        missing.append(pkg)
                
                if missing:
                    log_message(f"Installing required packages for {default_model}: {', '.join(missing)}")
                    try:
                        for pkg in missing:
                            # Install each package individually for better error reporting
                            try:
                                log_message(f"Installing {pkg}...")
                                subprocess.check_call([sys.executable, "-m", "pip", "install", "--upgrade", pkg])
                                log_message(f"Successfully installed {pkg}")
                            except Exception as e:
                                log_message(f"Error installing {pkg}: {str(e)}", True)
                    except Exception as e:
                        log_message(f"Error during package installation: {str(e)}", True)
        except Exception as e:
            log_message(f"Error checking/installing packages: {str(e)}", True)
        
        # Initialize wx app
        app = wx.App()
        
        # Create main window
        frame = ResearchAssistantApp()
        
        # Start the main loop
        app.MainLoop()
    except Exception as e:
        error_message = f"Critical error: {str(e)}"
        print(error_message)
        log_message(error_message, True)
        log_message(traceback.format_exc(), True)
        
        # Show error in dialog
        try:
            wx.MessageBox(f"Application error: {str(e)}\nCheck error_log.txt for details.", "Error", 
                         wx.OK | wx.ICON_ERROR)
        except:
            pass